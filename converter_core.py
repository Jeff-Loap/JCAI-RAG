import os
import re
import sys
import shutil
import subprocess
import tempfile
import io
import json
import hashlib
import importlib.util
from zipfile import ZipFile
from html import escape
from pathlib import Path
from typing import Optional, List, Tuple, Dict, Any
from urllib.parse import urlparse, parse_qs, urlencode, urlunparse

from PIL import Image, ImageOps, ImageDraw
import numpy as np

from PySide6.QtCore import QThread, Signal, Qt, QSize, QUrl
from PySide6.QtGui import QPixmap, QFont, QIcon
from PySide6.QtWidgets import (
    QApplication,
    QWidget,
    QLabel,
    QPushButton,
    QFileDialog,
    QLineEdit,
    QTextEdit,
    QComboBox,
    QCheckBox,
    QMessageBox,
    QGridLayout,
    QHBoxLayout,
    QVBoxLayout,
    QProgressBar,
    QSizePolicy,
    QListWidget,
    QListWidgetItem,
    QGroupBox,
    QSplitter,
    QFrame,
    QStackedWidget,
    QScrollArea,
    QSlider,
    QSpinBox,
)

IMAGE_EXTENSIONS = {
    ".jpg", ".jpeg", ".png", ".bmp", ".gif", ".webp", ".tif", ".tiff", ".ico"
}

AUDIO_EXTENSIONS = {
    ".mp3", ".wav", ".flac", ".aac", ".ogg", ".m4a", ".wma"
}

VIDEO_EXTENSIONS = {
    ".mp4", ".avi", ".mov", ".mkv", ".wmv", ".flv", ".webm", ".m4v"
}

DOCUMENT_EXTENSIONS = {
    ".doc", ".docx", ".odt", ".rtf", ".txt",
    ".xls", ".xlsx", ".ods", ".csv",
    ".ppt", ".pptx", ".odp", ".pdf"
}

CATEGORY_TO_SOURCE_EXTENSIONS = {
    "image": IMAGE_EXTENSIONS,
    "audio": AUDIO_EXTENSIONS,
    "video": VIDEO_EXTENSIONS,
    "document": DOCUMENT_EXTENSIONS,
}

CATEGORY_TO_TARGET_EXTENSIONS = {
    "image": ["jpg", "jpeg", "png", "bmp", "gif", "webp", "tif", "tiff", "ico", "pdf"],
    "audio": ["mp3", "wav", "flac", "aac", "ogg", "m4a"],
    "video": ["mp4", "avi", "mov", "mkv", "webm"],
    "document": ["pdf", "docx", "xlsx", "pptx", "txt", "html", "odt", "ods", "odp"],
}

CATEGORY_NAME_MAP = {
    "image": "图片",
    "audio": "音频",
    "video": "视频",
    "document": "文档",
}

EXTENSION_TO_CATEGORY = {}
for _category_name, _extensions in CATEGORY_TO_SOURCE_EXTENSIONS.items():
    for _ext in _extensions:
        EXTENSION_TO_CATEGORY[_ext] = _category_name

URL_EXTRACT_PATTERN = re.compile(r"https?://[^\s<>\"']+", re.IGNORECASE)
TRAILING_URL_PUNCTUATION = ".,;:!?)]}>，。；：！？、】）］」』"
FFMPEG_ENCODER_CACHE: Dict[str, set] = {}
CERTIFICATE_INFO_KEYWORDS = [
    "姓名", "性别", "民族", "出生", "住址", "公民身份号码", "身份证",
    "签发机关", "有效期限", "居民身份证", "证号", "护照", "passport",
    "驾驶证", "行驶证", "统一社会信用代码", "法定代表人", "营业执照",
]
CERTIFICATE_ROTATION_ANGLES = [0, 90, 180, 270]
CERTIFICATE_TESSERACT_PSMS = [6, 11, 12]
PDF2DOCX_MODULE_NAME = "pdf2docx"
PDF2DOCX_ENGINE_LABEL = "pdf2docx"
OCRMYPDF_MODULE_NAME = "ocrmypdf"
OCRMYPDF_ENGINE_LABEL = "OCRmyPDF + pdf2docx"
FREEP2W_ENGINE_LABEL = "FreeP2W"
LIBREOFFICE_ENGINE_LABEL = "LibreOffice"
RESUME_SECTION_TITLES = {
    "教育背景",
    "核心技能",
    "项目经历",
    "工作/实习经历",
    "自我评价",
}
RESUME_TITLE_COLOR = "1F4E79"
RESUME_TEXT_COLOR = "222222"
RESUME_META_COLOR = "555555"
PDF_ROUTE_ANALYSIS_CACHE: Dict[Tuple[str, int, int], Dict[str, Any]] = {}

if sys.platform.startswith("win"):
    LIBREOFFICE_BINARY_CANDIDATES = [
        "soffice.com",
        "soffice.exe",
        "soffice",
        r"C:\Program Files\LibreOffice\program\soffice.com",
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.com",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    FFMPEG_BINARY_CANDIDATES = [
        "ffmpeg.exe",
        "ffmpeg",
        r"C:\ffmpeg\bin\ffmpeg.exe",
        r"C:\Program Files\ffmpeg\bin\ffmpeg.exe",
        r"C:\Program Files (x86)\ffmpeg\bin\ffmpeg.exe",
    ]
    TESSERACT_BINARY_CANDIDATES = [
        "tesseract.exe",
        "tesseract",
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    ]
    YT_DLP_BINARY_CANDIDATES = [
        "yt-dlp.exe",
        "yt-dlp",
        r"C:\Program Files\yt-dlp\yt-dlp.exe",
    ]
else:
    LIBREOFFICE_BINARY_CANDIDATES = [
        "libreoffice",
        "soffice",
    ]
    FFMPEG_BINARY_CANDIDATES = [
        "ffmpeg",
    ]
    TESSERACT_BINARY_CANDIDATES = [
        "tesseract",
    ]
    YT_DLP_BINARY_CANDIDATES = [
        "yt-dlp",
    ]

SCENE_CONFIG = {
    "student": {
        "name": "学生",
        "description": "作业、论文、课件转换",
        "icon": "📚",
        "features": ["format_convert", "scan_to_doc", "pdf_watermark_remove"],
    },
    "office": {
        "name": "职场",
        "description": "文档处理、格式转换",
        "icon": "💼",
        "features": ["format_convert", "pdf_watermark_remove", "compress"],
    },
    "media": {
        "name": "自媒体",
        "description": "图片视频压缩、格式转换",
        "icon": "🎬",
        "features": ["format_convert", "compress", "video_download"],
    },
}


def find_executable(candidates: List[str]) -> Optional[str]:
    for candidate in candidates:
        candidate_path = Path(candidate)
        if candidate_path.is_file():
            return str(candidate_path)

        found = shutil.which(candidate)
        if found:
            return found

    return None


def normalize_ext(ext: str) -> str:
    ext = ext.strip().lower()
    if ext.startswith("."):
        ext = ext[1:]
    return ext


def format_file_size(size: int) -> str:
    units = ["B", "KB", "MB", "GB", "TB"]
    value = float(size)
    for unit in units:
        if value < 1024 or unit == units[-1]:
            return f"{value:.2f} {unit}"
        value /= 1024
    return f"{size} B"


def detect_category_by_file(file_path: Path) -> Optional[str]:
    return EXTENSION_TO_CATEGORY.get(file_path.suffix.lower())


def get_supported_targets_for_file(file_path: Path) -> List[str]:
    category = detect_category_by_file(file_path)
    if not category:
        return []

    suffix = file_path.suffix.lower()
    if category == "image":
        return ["jpg", "jpeg", "png", "bmp", "gif", "webp", "tif", "tiff", "ico", "pdf"]
    if category == "audio":
        return ["mp3", "wav", "flac", "aac", "ogg", "m4a"]
    if category == "video":
        return ["mp4", "avi", "mov", "mkv", "webm"]

    word_like = {".doc", ".docx", ".odt", ".rtf", ".txt", ".pdf"}
    sheet_like = {".xls", ".xlsx", ".ods", ".csv"}
    slide_like = {".ppt", ".pptx", ".odp"}

    if suffix in word_like:
        return ["pdf", "docx", "txt", "html", "odt"]
    if suffix in sheet_like:
        return ["pdf", "xlsx", "txt", "html", "ods"]
    if suffix in slide_like:
        return ["pdf", "pptx", "html", "odp"]

    return CATEGORY_TO_TARGET_EXTENSIONS.get(category, [])


def has_pdf2docx_engine() -> bool:
    return importlib.util.find_spec(PDF2DOCX_MODULE_NAME) is not None


def has_ocrmypdf_engine() -> bool:
    return importlib.util.find_spec(OCRMYPDF_MODULE_NAME) is not None or shutil.which("ocrmypdf") is not None


def has_freep2w_engine() -> bool:
    return (
        importlib.util.find_spec("freep2w") is not None
        or importlib.util.find_spec("FreeP2W") is not None
        or shutil.which("freep2w") is not None
    )


def analyze_pdf_to_docx_route(src_file: Path) -> Dict[str, Any]:
    try:
        stat_result = src_file.stat()
        cache_key = (str(src_file.resolve()), stat_result.st_mtime_ns, stat_result.st_size)
    except Exception:
        cache_key = (str(src_file), 0, 0)

    cached = PDF_ROUTE_ANALYSIS_CACHE.get(cache_key)
    if cached is not None:
        return dict(cached)

    route = {
        "engine": PDF2DOCX_ENGINE_LABEL,
        "engine_label": PDF2DOCX_ENGINE_LABEL,
        "is_scanned": False,
        "is_complex": False,
        "page_sample_count": 0,
        "avg_words_per_page": 0.0,
        "image_page_ratio": 0.0,
        "reason": "默认走 pdf2docx",
    }

    if src_file.suffix.lower() != ".pdf":
        PDF_ROUTE_ANALYSIS_CACHE[cache_key] = dict(route)
        return route

    try:
        import fitz

        pdf_doc = fitz.open(str(src_file))
        sample_page_count = min(5, len(pdf_doc))
        if sample_page_count <= 0:
            PDF_ROUTE_ANALYSIS_CACHE[cache_key] = dict(route)
            return route

        total_words = 0
        image_heavy_pages = 0
        dense_image_pages = 0
        formula_like_chars = 0
        for page_index in range(sample_page_count):
            page = pdf_doc[page_index]
            words = page.get_text("words")
            word_count = len(words)
            total_words += word_count
            image_count = len(page.get_images(full=True))
            text_length = len(page.get_text().strip())
            if image_count > 0 and word_count <= 25 and text_length <= 600:
                image_heavy_pages += 1
            if image_count >= 2:
                dense_image_pages += 1
            formula_like_chars += sum(page.get_text().count(ch) for ch in ("∑", "∫", "√", "≤", "≥", "≈", "×", "÷"))

        avg_words_per_page = total_words / sample_page_count
        image_page_ratio = image_heavy_pages / sample_page_count
        route["page_sample_count"] = sample_page_count
        route["avg_words_per_page"] = round(avg_words_per_page, 1)
        route["image_page_ratio"] = round(image_page_ratio, 2)

        is_scanned = image_heavy_pages >= max(1, sample_page_count // 2) and avg_words_per_page < 40
        is_complex = dense_image_pages >= max(1, sample_page_count // 3) or formula_like_chars >= 8

        route["is_scanned"] = is_scanned
        route["is_complex"] = is_complex
        if is_scanned:
            route["engine"] = OCRMYPDF_ENGINE_LABEL
            route["engine_label"] = OCRMYPDF_ENGINE_LABEL
            route["reason"] = "检测为扫描件或图片型 PDF，需先做 OCR"
        elif is_complex and has_freep2w_engine():
            route["engine"] = FREEP2W_ENGINE_LABEL
            route["engine_label"] = FREEP2W_ENGINE_LABEL
            route["reason"] = "检测为复杂版面 PDF，优先尝试 FreeP2W"
        else:
            route["reason"] = "检测为普通文本型 PDF，优先走 pdf2docx"
    except Exception:
        pass

    PDF_ROUTE_ANALYSIS_CACHE[cache_key] = dict(route)
    return route


def uses_pdf2docx_engine(src_file: Path, target_ext: str) -> bool:
    if src_file.suffix.lower() != ".pdf" or normalize_ext(target_ext) != "docx":
        return False
    return analyze_pdf_to_docx_route(src_file).get("engine") == PDF2DOCX_ENGINE_LABEL


def uses_ocrmypdf_engine(src_file: Path, target_ext: str) -> bool:
    if src_file.suffix.lower() != ".pdf" or normalize_ext(target_ext) != "docx":
        return False
    return analyze_pdf_to_docx_route(src_file).get("engine") == OCRMYPDF_ENGINE_LABEL


def uses_freep2w_engine(src_file: Path, target_ext: str) -> bool:
    if src_file.suffix.lower() != ".pdf" or normalize_ext(target_ext) != "docx":
        return False
    return analyze_pdf_to_docx_route(src_file).get("engine") == FREEP2W_ENGINE_LABEL


def uses_libreoffice_engine(src_file: Path, target_ext: str) -> bool:
    normalized_target = normalize_ext(target_ext)
    category = detect_category_by_file(src_file)

    if src_file.suffix.lower() == ".pdf" and normalized_target == "docx":
        return False

    if category == "image" and normalized_target == "docx":
        return True

    if category == "document":
        return not uses_pdf2docx_engine(src_file, normalized_target)

    return False


def get_document_conversion_engine_label(src_file: Path, target_ext: str) -> str:
    normalized_target = normalize_ext(target_ext)
    if src_file.suffix.lower() == ".pdf" and normalized_target == "docx":
        return analyze_pdf_to_docx_route(src_file).get("engine_label", PDF2DOCX_ENGINE_LABEL)
    return LIBREOFFICE_ENGINE_LABEL


def filter_files_by_category(files: List[Path], category: str) -> List[Path]:
    valid_extensions = CATEGORY_TO_SOURCE_EXTENSIONS[category]
    return [file_path for file_path in files if file_path.suffix.lower() in valid_extensions]


def build_output_path(
        src_file: Path,
        dst_root: Path,
        target_ext: str,
) -> Path:
    output_dir = dst_root
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir / f"{src_file.stem}.{target_ext}"


def suggest_output_dir_from_files(files: List[Path]) -> Optional[Path]:
    valid_files = [Path(item).expanduser().resolve() for item in files if Path(item).exists()]
    if not valid_files:
        return None
    return valid_files[0].parent


def get_default_download_dir() -> Path:
    candidate = Path.home() / "Downloads"
    if candidate.exists() and candidate.is_dir():
        return candidate
    return Path.home()


def safe_strip_text(value: Any) -> str:
    if value is None:
        return ""
    return str(value).strip()


def build_distinct_output_path(src_file: Path, dst_root: Path, target_ext: str, mode: str = "convert") -> Path:
    candidate = build_output_path(src_file=src_file, dst_root=dst_root, target_ext=target_ext)
    try:
        same_path = candidate.resolve() == src_file.resolve()
    except Exception:
        same_path = candidate == src_file

    if not same_path:
        return candidate

    suffix_map = {
        "compress": "compressed",
        "watermark": "no_watermark",
        "scan": "ocr",
        "convert": "converted",
    }
    label = suffix_map.get(mode, "output")
    return dst_root / f"{src_file.stem}_{label}.{target_ext}"


def _strip_trailing_url_punctuation(url: str) -> str:
    while url and url[-1] in TRAILING_URL_PUNCTUATION:
        url = url[:-1]
    return url


def normalize_video_url(url: str) -> str:
    cleaned_url = _strip_trailing_url_punctuation(url.strip().strip("()[]{}<>（）【】「」『』"))
    parsed = urlparse(cleaned_url)
    if parsed.scheme not in {"http", "https"} or not parsed.netloc:
        raise ValueError(f"链接格式无效：{url}")

    host = parsed.netloc.lower()
    query_data = parse_qs(parsed.query, keep_blank_values=False)

    if host.endswith("bilibili.com") or host.endswith("b23.tv"):
        canonical_query: Dict[str, List[str]] = {}
        if "p" in query_data:
            canonical_query["p"] = query_data["p"]
        cleaned_url = urlunparse((
            parsed.scheme,
            parsed.netloc,
            parsed.path.rstrip("/"),
            "",
            urlencode(canonical_query, doseq=True),
            "",
        ))
    else:
        cleaned_url = urlunparse((
            parsed.scheme,
            parsed.netloc,
            parsed.path,
            "",
            parsed.query,
            "",
        ))

    return cleaned_url


def extract_video_urls_from_text(text: str) -> List[str]:
    urls: List[str] = []
    seen: set = set()

    for line_number, raw_line in enumerate(text.splitlines(), start=1):
        line = raw_line.strip()
        if not line:
            continue

        matches = URL_EXTRACT_PATTERN.findall(line)
        if not matches:
            raise ValueError(f"第 {line_number} 行未识别到有效链接：{raw_line}")

        for match in matches:
            normalized_url = normalize_video_url(match)
            if normalized_url not in seen:
                seen.add(normalized_url)
                urls.append(normalized_url)

    return urls


def create_image_preview(
        file_path: Path,
        max_size: Optional[Tuple[int, int]] = None,
) -> Optional[QPixmap]:
    try:
        pixmap = QPixmap(str(file_path))
        if pixmap.isNull():
            return None

        if not max_size:
            return pixmap

        max_width = max(1, int(max_size[0]))
        max_height = max(1, int(max_size[1]))

        return pixmap.scaled(
            max_width,
            max_height,
            Qt.KeepAspectRatio,
            Qt.SmoothTransformation
        )
    except Exception:
        return None


def build_preview_text(
        src_file: Path,
        dst_root: Path,
        target_ext: str,
) -> str:
    try:
        size_text = format_file_size(src_file.stat().st_size)
    except Exception:
        size_text = "未知"

    preview_output = build_output_path(
        src_file=src_file,
        dst_root=dst_root,
        target_ext=target_ext,
    )

    detected_category = detect_category_by_file(src_file)
    category_text = CATEGORY_NAME_MAP.get(detected_category, "未知类型")

    return (
        f"文件名：{src_file.name}\n"
        f"源路径：{src_file}\n"
        f"文件大小：{size_text}\n"
        f"自动识别类别：{category_text}\n"
        f"当前格式：{src_file.suffix.lower().replace('.', '') or '无'}\n"
        f"目标格式：{target_ext}\n"
        f"预计输出：{preview_output}"
    )


def convert_image(src_file: Path, dst_file: Path) -> Tuple[bool, str]:
    target_format = dst_file.suffix.replace(".", "").upper()

    if target_format == "JPG":
        target_format = "JPEG"
    elif target_format == "TIF":
        target_format = "TIFF"

    try:
        with Image.open(src_file) as img:
            if target_format in {"JPEG", "JPG", "BMP"}:
                if img.mode in ("RGBA", "LA", "P"):
                    if img.mode == "P":
                        img = img.convert("RGBA")
                    background = Image.new("RGB", img.size, (255, 255, 255))
                    alpha_channel = img.split()[-1] if "A" in img.mode else None
                    background.paste(img, mask=alpha_channel)
                    img = background
                else:
                    img = img.convert("RGB")
            img.save(dst_file, format=target_format)
        return True, f"成功：{src_file.name} -> {dst_file.name}"
    except Exception as exc:
        return False, f"失败：{src_file.name} -> {dst_file.name} | {exc}"


def _prepare_pil_image_for_export(img: Image.Image, force_rgb: bool = False) -> Image.Image:
    img = ImageOps.exif_transpose(img)
    if force_rgb:
        if img.mode in ("RGBA", "LA", "P"):
            if img.mode == "P":
                img = img.convert("RGBA")
            background = Image.new("RGB", img.size, (255, 255, 255))
            alpha_channel = img.split()[-1] if "A" in img.mode else None
            background.paste(img, mask=alpha_channel)
            return background
        if img.mode != "RGB":
            return img.convert("RGB")
    return img


def convert_image_to_pdf(src_file: Path, dst_file: Path) -> Tuple[bool, str]:
    try:
        with Image.open(src_file) as img:
            pdf_image = _prepare_pil_image_for_export(img, force_rgb=True)
            pdf_image.save(dst_file, format="PDF", resolution=300.0)
        return True, f"成功：{src_file.name} -> {dst_file.name}"
    except Exception as exc:
        return False, f"失败：{src_file.name} -> {dst_file.name} | {exc}"


def convert_image_to_docx(src_file: Path, dst_file: Path, libreoffice_bin: Optional[str]) -> Tuple[bool, str]:
    if not libreoffice_bin:
        return False, f"失败：{src_file.name} -> {dst_file.name} | 图片转 DOCX 需要 LibreOffice"

    try:
        with tempfile.TemporaryDirectory(prefix="image_to_docx_") as temp_dir:
            html_path = Path(temp_dir) / f"{src_file.stem}.html"
            image_uri = src_file.resolve().as_uri()
            html_content = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="utf-8">
    <style>
        @page {{ size: A4; margin: 16mm; }}
        body {{ margin: 0; padding: 0; font-family: "Microsoft YaHei", sans-serif; }}
        .page {{ width: 100%; text-align: center; }}
        img {{ max-width: 100%; height: auto; display: inline-block; }}
    </style>
</head>
<body>
    <div class="page">
        <img src="{image_uri}" alt="{escape(src_file.name)}">
    </div>
</body>
</html>
"""
            html_path.write_text(html_content, encoding="utf-8")
            return convert_document(html_path, dst_file, libreoffice_bin)
    except Exception as exc:
        return False, f"失败：{src_file.name} -> {dst_file.name} | {exc}"


def convert_media(src_file: Path, dst_file: Path, ffmpeg_bin: str) -> Tuple[bool, str]:
    cmd = [
        ffmpeg_bin,
        "-y",
        "-i", str(src_file),
        str(dst_file)
    ]
    try:
        result = subprocess.run(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True
        )
        if result.returncode == 0:
            return True, f"成功：{src_file.name} -> {dst_file.name}"
        error_text = safe_strip_text(result.stderr) or safe_strip_text(result.stdout) or "外部程序未返回可读错误信息"
        return False, f"失败：{src_file.name} -> {dst_file.name} | {error_text}"
    except Exception as exc:
        return False, f"失败：{src_file.name} -> {dst_file.name} | {exc}"


def libreoffice_filter_name(target_ext: str) -> str:
    mapping = {
        "pdf": "pdf",
        "docx": "docx",
        "xlsx": "xlsx",
        "pptx": "pptx",
        "txt": "txt",
        "html": "html",
        "odt": "odt",
        "ods": "ods",
        "odp": "odp",
    }
    return mapping[target_ext]


def extract_first_image_from_docx(docx_path: Path) -> Optional[bytes]:
    try:
        with ZipFile(docx_path) as archive:
            media_names = [name for name in archive.namelist() if name.startswith("word/media/")]
            if not media_names:
                return None
            media_names.sort()
            return archive.read(media_names[0])
    except Exception:
        return None


def repair_pdf2docx_layout(src_file: Path, docx_path: Path) -> Tuple[bool, str]:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt, Cm, Emu
        from docx.text.paragraph import Paragraph
    except ImportError:
        return False, "版式修复跳过：缺少 python-docx"

    header_extra_space_before_pt = 0.0

    def _paragraph_text(paragraph: Any) -> str:
        return "".join(run.text for run in paragraph.runs).replace("\r", "").strip()

    def _delete_paragraph(paragraph: Any) -> None:
        element = paragraph._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)

    def _move_paragraph_after(paragraph: Any, reference_paragraph: Any) -> None:
        if paragraph is reference_paragraph:
            return
        paragraph_element = paragraph._element
        reference_element = reference_paragraph._element
        parent = reference_element.getparent()
        if parent is None or paragraph_element.getparent() is not parent:
            return
        paragraph_index = parent.index(paragraph_element)
        reference_index = parent.index(reference_element)
        if paragraph_index == reference_index + 1:
            return
        parent.remove(paragraph_element)
        if paragraph_index < reference_index:
            reference_index -= 1
        parent.insert(reference_index + 1, paragraph_element)

    def _insert_paragraph_after(reference_paragraph: Any) -> Any:
        new_paragraph_element = OxmlElement("w:p")
        reference_paragraph._p.addnext(new_paragraph_element)
        new_paragraph = Paragraph(new_paragraph_element, reference_paragraph._parent)
        if getattr(reference_paragraph, "style", None) is not None:
            try:
                new_paragraph.style = reference_paragraph.style
            except Exception:
                pass
        return new_paragraph

    def _replace_paragraph_text(paragraph: Any, text: str) -> None:
        for child in list(paragraph._p):
            if child.tag != qn("w:pPr"):
                paragraph._p.remove(child)
        if text:
            paragraph.add_run(text)

    def _insert_paragraph_after_table(table: Any) -> Any:
        new_paragraph_element = OxmlElement("w:p")
        table._tbl.addnext(new_paragraph_element)
        return Paragraph(new_paragraph_element, table._parent)

    def _delete_table(table: Any) -> None:
        table_element = table._tbl
        parent = table_element.getparent()
        if parent is not None:
            parent.remove(table_element)

    def _set_paragraph_top_border(
            paragraph: Any,
            color: str = "1F4E79",
            size: str = "12",
            space: str = "1",
    ) -> None:
        paragraph_properties = paragraph._p.get_or_add_pPr()
        paragraph_borders = paragraph_properties.find(qn("w:pBdr"))
        if paragraph_borders is None:
            paragraph_borders = OxmlElement("w:pBdr")
            paragraph_properties.append(paragraph_borders)

        for child in list(paragraph_borders):
            if child.tag == qn("w:top"):
                paragraph_borders.remove(child)

        top_border = OxmlElement("w:top")
        top_border.set(qn("w:val"), "single")
        top_border.set(qn("w:sz"), size)
        top_border.set(qn("w:space"), space)
        top_border.set(qn("w:color"), color)
        paragraph_borders.append(top_border)

    def _is_heading(paragraph: Any) -> bool:
        text = _paragraph_text(paragraph)
        if not text or text not in RESUME_SECTION_TITLES:
            return False
        return any(run.bold for run in paragraph.runs if run.text.strip())

    def _is_contact_paragraph(paragraph: Any) -> bool:
        text = _paragraph_text(paragraph)
        return "手机/微信" in text and "邮箱" in text

    def _is_bullet_paragraph(paragraph: Any) -> bool:
        text = _paragraph_text(paragraph).lstrip()
        if not text:
            return False
        if re.match(r"^[•·●○▪■◆◦◉◌]\s*", text):
            return True
        style_name = paragraph.style.name if getattr(paragraph, "style", None) is not None else ""
        if "List" in style_name:
            return True
        if text.startswith("?") and _left_indent_points(paragraph) >= 18.0:
            return True
        return False

    def _has_drawing(paragraph: Any) -> bool:
        return bool(paragraph._p.xpath(".//w:drawing"))

    def _left_indent_points(paragraph: Any) -> float:
        left_indent = paragraph.paragraph_format.left_indent
        return float(left_indent.pt) if left_indent else 0.0

    def _looks_like_new_resume_block(text: str) -> bool:
        stripped = safe_strip_text(text)
        if not stripped:
            return False
        if stripped in RESUME_SECTION_TITLES:
            return True
        if stripped.startswith(("•", "技术栈", "项目经历", "教育背景", "核心技能", "工作/实习经历", "自我评价")):
            return True
        if re.search(r"20\d{2}[./]\d{2}\s*[-~]\s*(?:20\d{2}[./]\d{2}|至今)", stripped):
            return True
        if ("|" in stripped or "｜" in stripped) and len(stripped) <= 80:
            return True
        return False

    def _normalize_sections(document: Any) -> None:
        for cols in document.element.xpath(".//w:sectPr/w:cols"):
            for child in list(cols):
                cols.remove(child)
            for attr_name in list(cols.attrib):
                if attr_name != qn("w:space"):
                    del cols.attrib[attr_name]
            cols.set(qn("w:num"), "1")

    def _find_resume_key_paragraphs(document: Any) -> Tuple[Optional[Any], Optional[Any], Optional[Any], Optional[Any]]:
        paragraphs = list(document.paragraphs)
        name_para = next(
            (
                paragraph for paragraph in paragraphs
                if _paragraph_text(paragraph)
                and any(
                    run.font.size and run.font.size.pt and run.font.size.pt >= 20
                    for run in paragraph.runs
                )
            ),
            None,
        )
        image_para = next(
            (
                paragraph for paragraph in paragraphs
                if paragraph._p.xpath(".//w:drawing")
            ),
            None,
        )
        contact_para = next((paragraph for paragraph in paragraphs if _is_contact_paragraph(paragraph)), None)
        heading_para = next(
            (
                paragraph for paragraph in paragraphs
                if paragraph is not name_para and _is_heading(paragraph)
            ),
            None,
        )
        return name_para, image_para, contact_para, heading_para

    def _shift_anchor_paragraph_down(paragraph: Any, offset_emu: int = 220000) -> None:
        anchors = paragraph._p.xpath(".//wp:anchor")
        for anchor in anchors:
            vertical_nodes = anchor.xpath("./wp:positionV/wp:posOffset")
            if not vertical_nodes:
                continue
            try:
                current_value = int(vertical_nodes[0].text)
            except Exception:
                current_value = 0
            vertical_nodes[0].text = str(current_value + offset_emu)

    def _has_anchor(paragraph: Any) -> bool:
        return bool(paragraph._p.xpath(".//wp:anchor"))

    def _anchor_height_points(paragraph: Any) -> float:
        extents = paragraph._p.xpath(".//wp:anchor/wp:extent")
        if not extents:
            return 0.0
        height_emu = extents[0].get("cy")
        if not height_emu:
            return 0.0
        try:
            return int(height_emu) / 12700.0
        except Exception:
            return 0.0

    def _normalize_bullet_prefix(paragraph: Any) -> bool:
        bullet_chars = "•·●○▪■◆◦◉◌?"
        runs = list(paragraph.runs)
        for index, run in enumerate(runs):
            text = run.text or ""
            stripped = text.lstrip()
            if not stripped:
                continue
            bullet_char = stripped[0]
            if bullet_char not in bullet_chars:
                return False

            leading = text[: len(text) - len(stripped)]
            normalized_bullet = "•"
            if len(stripped) > 1:
                remainder = stripped[1:].lstrip(" \t")
                run.text = f"{leading}{normalized_bullet}\t{remainder}"
                return True

            next_index = index + 1
            while next_index < len(runs):
                next_text = runs[next_index].text or ""
                if next_text.startswith("\t"):
                    return False
                if next_text == "\t":
                    return False
                if not next_text:
                    next_index += 1
                    continue
                if not next_text.strip():
                    runs[next_index].text = "\t"
                    run.text = f"{leading}{normalized_bullet}"
                    return True
                runs[next_index].text = f"\t{next_text.lstrip()}"
                run.text = f"{leading}{normalized_bullet}"
                return True
            return False
        return False

    def _split_mixed_bullet_paragraphs(document: Any) -> int:
        bullet_pattern = re.compile(r"(?:(?<=^)|(?<=[\s]))([•·●○▪■◆◦◉◌?])")

        def _normalize_prefix_text(text: str) -> str:
            normalized_lines = []
            for line in text.splitlines():
                stripped_line = re.sub(r"[ \t]+", " ", line).strip()
                if stripped_line:
                    normalized_lines.append(stripped_line)
            return "\n".join(normalized_lines)

        def _normalize_bullet_text(text: str) -> str:
            return re.sub(r"[\s\t\n]+", " ", text).strip()

        split_count = 0
        paragraphs = list(document.paragraphs)
        index = 0

        while index < len(paragraphs):
            paragraph = paragraphs[index]
            raw_text = "".join(run.text for run in paragraph.runs).replace("\r", "")
            if not raw_text or not any(ch in raw_text for ch in "•·●○▪■◆◦◉◌?"):
                index += 1
                continue

            bullet_matches = list(bullet_pattern.finditer(raw_text))
            if not bullet_matches:
                index += 1
                continue

            prefix_text = _normalize_prefix_text(raw_text[:bullet_matches[0].start()])
            bullet_items = []
            for match_index, match in enumerate(bullet_matches):
                start = match.start()
                end = bullet_matches[match_index + 1].start() if match_index + 1 < len(bullet_matches) else len(raw_text)
                segment_text = raw_text[start:end].lstrip()
                if not segment_text:
                    continue
                segment_body = _normalize_bullet_text(segment_text[1:])
                if segment_body:
                    bullet_items.append(segment_body)

            has_mixed_structure = bool(prefix_text) or len(bullet_items) >= 2 or "\n" in raw_text
            if not has_mixed_structure or not bullet_items:
                index += 1
                continue

            if prefix_text:
                _replace_paragraph_text(paragraph, prefix_text)
                last_paragraph = paragraph
            else:
                _replace_paragraph_text(paragraph, f"•\t{bullet_items[0]}")
                last_paragraph = paragraph
                bullet_items = bullet_items[1:]

            inserted_count = 0
            for item_text in bullet_items:
                new_paragraph = _insert_paragraph_after(last_paragraph)
                _replace_paragraph_text(new_paragraph, f"•\t{item_text}")
                last_paragraph = new_paragraph
                paragraphs.insert(index + 1 + inserted_count, new_paragraph)
                inserted_count += 1

            split_count += inserted_count + (1 if prefix_text else 0)
            index += inserted_count + 1

        return split_count

    def _split_mixed_numbered_paragraphs(document: Any) -> int:
        item_pattern = re.compile(r"(?:(?<=^)|(?<=\n)|(?<=\s))(\d+\.)\s+")

        def _normalize_prefix_text(text: str) -> str:
            normalized_lines = []
            for line in text.splitlines():
                stripped_line = re.sub(r"[ \t]+", " ", line).strip()
                if stripped_line:
                    normalized_lines.append(stripped_line)
            return "\n".join(normalized_lines)

        def _normalize_item_text(marker: str, body: str) -> str:
            normalized_body = re.sub(r"[\s\t\n]+", " ", body).strip()
            return f"{marker} {normalized_body}".strip()

        split_count = 0
        paragraphs = list(document.paragraphs)
        index = 0

        while index < len(paragraphs):
            paragraph = paragraphs[index]
            raw_text = "".join(run.text for run in paragraph.runs).replace("\r", "")
            if not raw_text or not re.search(r"\d+\.\s", raw_text):
                index += 1
                continue
            if re.match(r"^\s*(Fig\.?|Figure|Table)\s+\d+", raw_text.strip(), re.IGNORECASE):
                index += 1
                continue

            matches = list(item_pattern.finditer(raw_text))
            if len(matches) < 2:
                index += 1
                continue

            prefix_text = _normalize_prefix_text(raw_text[:matches[0].start()])
            item_texts = []
            for match_index, match in enumerate(matches):
                start = match.start()
                end = matches[match_index + 1].start() if match_index + 1 < len(matches) else len(raw_text)
                marker = match.group(1)
                body = raw_text[match.end():end]
                normalized_item = _normalize_item_text(marker, body)
                if normalized_item:
                    item_texts.append(normalized_item)

            if not item_texts:
                index += 1
                continue

            if prefix_text:
                _replace_paragraph_text(paragraph, prefix_text)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.left_indent = Pt(0)
                paragraph.paragraph_format.first_line_indent = Pt(0)
                last_paragraph = paragraph
            else:
                _replace_paragraph_text(paragraph, item_texts[0])
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.left_indent = Pt(0)
                paragraph.paragraph_format.first_line_indent = Pt(0)
                last_paragraph = paragraph
                item_texts = item_texts[1:]

            inserted_count = 0
            for item_text in item_texts:
                new_paragraph = _insert_paragraph_after(last_paragraph)
                _replace_paragraph_text(new_paragraph, item_text)
                new_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                new_paragraph.paragraph_format.left_indent = Pt(0)
                new_paragraph.paragraph_format.first_line_indent = Pt(0)
                new_paragraph.paragraph_format.space_before = Pt(0)
                new_paragraph.paragraph_format.space_after = Pt(2)
                last_paragraph = new_paragraph
                paragraphs.insert(index + 1 + inserted_count, new_paragraph)
                inserted_count += 1

            split_count += inserted_count + (1 if prefix_text else 0)
            index += inserted_count + 1

        return split_count

    def _split_heading_body_paragraphs(document: Any) -> int:
        heading_pattern = re.compile(
            r"^(Abstract|Keywords|Appendix|References|\d+(?:\.\d+)*\s+.+)$",
            re.IGNORECASE,
        )

        split_count = 0
        for paragraph in list(document.paragraphs):
            raw_text = "".join(run.text for run in paragraph.runs).replace("\r", "")
            if "\n" not in raw_text:
                continue

            first_line, remainder = raw_text.split("\n", 1)
            heading_text = re.sub(r"[ \t]+", " ", first_line).strip()
            body_text = re.sub(r"[\s\t\n]+", " ", remainder).strip()
            if not heading_text or not body_text:
                continue
            if not heading_pattern.match(heading_text):
                continue

            _replace_paragraph_text(paragraph, heading_text)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = Pt(0)
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(2)
            heading_run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(heading_text)
            heading_run.bold = True
            heading_run.font.size = Pt(12)

            body_paragraph = _insert_paragraph_after(paragraph)
            _replace_paragraph_text(body_paragraph, body_text)
            body_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            body_paragraph.paragraph_format.left_indent = Pt(0)
            body_paragraph.paragraph_format.first_line_indent = Pt(0)
            body_paragraph.paragraph_format.space_before = Pt(0)
            body_paragraph.paragraph_format.space_after = Pt(2)
            split_count += 1

        return split_count

    def _split_caption_body_paragraphs(document: Any) -> int:
        split_count = 0
        paragraphs = list(document.paragraphs)
        for index, paragraph in enumerate(paragraphs):
            raw_text = "".join(run.text for run in paragraph.runs).replace("\r", "")
            if "\n" not in raw_text:
                continue
            first_line, remainder = raw_text.split("\n", 1)
            if not re.match(r"^(Fig\.?\s*\d+|Figure\s+\d+|Table\s+\d+)", first_line.strip(), re.IGNORECASE):
                continue
            caption_text = re.sub(r"[ \t]+", " ", first_line).strip()
            body_text = re.sub(r"[\s\t\n]+", " ", remainder).strip()
            _replace_paragraph_text(paragraph, caption_text)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.left_indent = Pt(0)
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(2)
            if body_text:
                body_paragraph = _insert_paragraph_after(paragraph)
                _replace_paragraph_text(body_paragraph, body_text)
                body_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                body_paragraph.paragraph_format.left_indent = Pt(0)
                body_paragraph.paragraph_format.first_line_indent = Pt(0)
                body_paragraph.paragraph_format.space_before = Pt(0)
                body_paragraph.paragraph_format.space_after = Pt(2)
            split_count += 1
        return split_count

    def _repair_paper_front_matter(src_file: Path, document: Any) -> int:
        if src_file.suffix.lower() != ".pdf":
            return 0

        try:
            import fitz
        except ImportError:
            return 0

        paragraphs = list(document.paragraphs)
        abstract_index = next(
            (index for index, paragraph in enumerate(paragraphs) if _paragraph_text(paragraph).strip().startswith("Abstract")),
            None,
        )
        if abstract_index is None:
            return 0

        adjusted_count = 0
        for paragraph in paragraphs[:abstract_index]:
            text = _paragraph_text(paragraph).strip()
            if not text:
                continue
            if any(token in text for token in ("@", "Corresponding", "Master of Data Science", "University of Science Malaysia")):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.left_indent = Pt(0)
                paragraph.paragraph_format.first_line_indent = Pt(0)
                adjusted_count += 1

            if "\n" in "".join(run.text for run in paragraph.runs):
                lines = [re.sub(r"[ \t]+", " ", line).strip() for line in "".join(run.text for run in paragraph.runs).splitlines()]
                lines = [line for line in lines if line]
                if len(lines) >= 2 and "@" not in text:
                    _replace_paragraph_text(paragraph, lines[0])
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.left_indent = Pt(0)
                    paragraph.paragraph_format.first_line_indent = Pt(0)
                    last_paragraph = paragraph
                    for line in lines[1:]:
                        new_paragraph = _insert_paragraph_after(last_paragraph)
                        _replace_paragraph_text(new_paragraph, line)
                        new_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        new_paragraph.paragraph_format.left_indent = Pt(0)
                        new_paragraph.paragraph_format.first_line_indent = Pt(0)
                        new_paragraph.paragraph_format.space_before = Pt(0)
                        new_paragraph.paragraph_format.space_after = Pt(2)
                        last_paragraph = new_paragraph
                    adjusted_count += len(lines) - 1

        try:
            pdf_doc = fitz.open(str(src_file))
            front_text = "".join(pdf_doc[page_index].get_text() for page_index in range(min(3, len(pdf_doc))))
        except Exception:
            front_text = ""

        email_matches = re.findall(r"[A-Za-z0-9._%+-]+@student\.usm\.my", front_text)
        if email_matches:
            unique_emails = []
            for email in email_matches:
                if email not in unique_emails:
                    unique_emails.append(email)
            corresponding_line = "Corresponding author: " + " / ".join(unique_emails)
            corresponding_paragraph = next(
                (paragraph for paragraph in document.paragraphs if "Corresponding author" in _paragraph_text(paragraph)),
                None,
            )
            if corresponding_paragraph is not None:
                _replace_paragraph_text(corresponding_paragraph, corresponding_line)
                corresponding_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                corresponding_paragraph.paragraph_format.left_indent = Pt(0)
                corresponding_paragraph.paragraph_format.first_line_indent = Pt(0)
                adjusted_count += 1

        return adjusted_count

    def _normalize_pdf_font_name(font_name: str) -> str:
        normalized = (font_name or "").strip().split("+")[-1]
        known_mappings = {
            "TimesNewRomanPSMT": "Times New Roman",
            "TimesNewRomanPS-BoldMT": "Times New Roman",
            "TimesNewRomanPS-ItalicMT": "Times New Roman",
            "TimesNewRomanPS-BoldItalicMT": "Times New Roman",
            "ArialMT": "Arial",
            "Arial-BoldMT": "Arial",
            "Calibri": "Calibri",
            "Calibri-Bold": "Calibri",
            "Helvetica": "Arial",
            "Helvetica-Bold": "Arial",
            "CourierNewPSMT": "Courier New",
        }
        if normalized in known_mappings:
            return known_mappings[normalized]
        if "TimesNewRoman" in normalized or normalized.startswith("Times"):
            return "Times New Roman"
        if "Arial" in normalized:
            return "Arial"
        if "Calibri" in normalized:
            return "Calibri"
        return normalized or "Times New Roman"

    def _set_run_font_name(run: Any, font_name: str) -> None:
        resolved_name = _normalize_pdf_font_name(font_name)
        run.font.name = resolved_name
        run_properties = run._element.get_or_add_rPr()
        run_fonts = run_properties.rFonts
        if run_fonts is None:
            run_fonts = OxmlElement("w:rFonts")
            run_properties.append(run_fonts)
        for attr_name in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            run_fonts.set(qn(attr_name), resolved_name)

    def _strip_toc_line_title(line: str) -> str:
        normalized_line = re.sub(r"\s+", " ", line or "").strip()
        return re.sub(r"\s*\.{2,}\s*\d+\s*$", "", normalized_line).strip()

    def _parse_toc_line(line: str) -> Tuple[str, str]:
        normalized_line = re.sub(r"\s+", " ", line or "").strip()
        match = re.match(r"^(?P<title>.+?)\s*\.{2,}\s*(?P<page>\d+)\s*$", normalized_line)
        if not match:
            return normalized_line, ""
        return match.group("title").strip(), match.group("page").strip()

    def _clear_paragraph_tab_stops(paragraph: Any) -> None:
        paragraph_properties = paragraph._p.get_or_add_pPr()
        tabs = paragraph_properties.find(qn("w:tabs"))
        if tabs is not None:
            paragraph_properties.remove(tabs)

    def _get_toc_indent_level(title: str) -> int:
        match = re.match(r"^(\d+(?:\.\d+)*)\.?\s+", title.strip())
        if not match:
            return 0
        number_token = match.group(1).strip(".")
        parts = [part for part in number_token.split(".") if part]
        return max(0, len(parts) - 1)

    def _build_toc_paragraph_text(title: str, page_number: str) -> str:
        if not page_number:
            return title
        return f"{title}\t{page_number}"

    def _extract_pdf_paper_profile(src_file: Path) -> dict[str, Any]:
        profile = {
            "body_font_name": "Times New Roman",
            "body_font_size": 11.0,
            "body_line_ratio": 1.35,
            "toc_lines": [],
            "section_markers": [],
        }
        if src_file.suffix.lower() != ".pdf":
            return profile

        try:
            import fitz
        except ImportError:
            return profile

        try:
            pdf_doc = fitz.open(str(src_file))
        except Exception:
            return profile

        body_font_counts: dict[str, int] = {}
        body_size_counts: dict[float, int] = {}
        line_ratios = []
        toc_lines: list[str] = []

        for page_index, page in enumerate(pdf_doc):
            page_dict = page.get_text("dict")
            page_text_lines = [line.strip() for line in page.get_text().splitlines() if line.strip()]
            if "Table of Contents" in page.get_text() and not toc_lines:
                try:
                    toc_start = page_text_lines.index("Table of Contents") + 1
                except ValueError:
                    toc_start = 0
                toc_lines = [
                    line for line in page_text_lines[toc_start:]
                    if line and not line.isdigit()
                ]

            for block in page_dict.get("blocks", []):
                if block.get("type") != 0:
                    continue
                block_body_lines = []
                for line in block.get("lines", []):
                    spans = [span for span in line.get("spans", []) if span.get("text", "").strip()]
                    if not spans:
                        continue
                    line_text = "".join(span.get("text", "") for span in spans).strip()
                    if not line_text:
                        continue
                    size_values = [round(float(span.get("size", 0.0)), 1) for span in spans if span.get("size")]
                    avg_size = round(sum(size_values) / len(size_values), 1) if size_values else 0.0
                    x0 = min(span["bbox"][0] for span in spans)
                    y0 = min(span["bbox"][1] for span in spans)
                    x1 = max(span["bbox"][2] for span in spans)
                    y1 = max(span["bbox"][3] for span in spans)
                    line_info = {
                        "text": line_text,
                        "size": avg_size,
                        "x0": x0,
                        "y0": y0,
                        "x1": x1,
                        "y1": y1,
                    }

                    if (
                            page_index >= 2
                            and len(line_text) >= 50
                            and 9.0 <= avg_size <= 12.5
                            and "@" not in line_text
                            and "..." not in line_text
                            and not re.match(r"^(Abstract|Keywords|References|Appendix|\d+(?:\.\d+)*\s)", line_text)
                    ):
                        block_body_lines.append(line_info)
                        body_size_counts[avg_size] = body_size_counts.get(avg_size, 0) + 1
                        for span in spans:
                            span_text = span.get("text", "").strip()
                            if not span_text:
                                continue
                            normalized_font = _normalize_pdf_font_name(span.get("font", ""))
                            body_font_counts[normalized_font] = body_font_counts.get(normalized_font, 0) + len(span_text)

                for current_line, next_line in zip(block_body_lines, block_body_lines[1:]):
                    if abs(current_line["x0"] - next_line["x0"]) > 6:
                        continue
                    gap = next_line["y0"] - current_line["y0"]
                    if current_line["size"] <= 0:
                        continue
                    ratio = gap / current_line["size"]
                    if 1.0 <= ratio <= 1.6:
                        line_ratios.append(ratio)

        if body_font_counts:
            profile["body_font_name"] = max(body_font_counts.items(), key=lambda item: item[1])[0]
        if body_size_counts:
            profile["body_font_size"] = max(body_size_counts.items(), key=lambda item: item[1])[0]
        if line_ratios:
            line_ratios.sort()
            median_ratio = round(line_ratios[len(line_ratios) // 2], 2)
            profile["body_line_ratio"] = max(1.15, min(median_ratio, 1.45))
        if toc_lines:
            profile["toc_lines"] = toc_lines
            profile["section_markers"] = [
                title for title in (_strip_toc_line_title(line) for line in toc_lines)
                if title and title not in {"Abstract", "Keywords", "References", "Appendix"}
            ]

        return profile

    def _repair_paper_toc(document: Any, toc_lines: list[str]) -> int:
        if not toc_lines:
            return 0

        paragraphs = list(document.paragraphs)
        toc_heading_index = next(
            (index for index, paragraph in enumerate(paragraphs) if _paragraph_text(paragraph).strip() == "Table of Contents"),
            None,
        )
        if toc_heading_index is None:
            return 0

        content_start_index = next(
            (
                index for index, paragraph in enumerate(paragraphs[toc_heading_index + 1:], start=toc_heading_index + 1)
                if "Real-World Cloud Security Breach Analysis" in _paragraph_text(paragraph)
            ),
            None,
        )
        if content_start_index is None:
            return 0

        toc_heading = paragraphs[toc_heading_index]
        content_start_paragraph = paragraphs[content_start_index]
        section = document.sections[0] if document.sections else None
        usable_width = None
        if section is not None:
            try:
                usable_width = section.page_width - section.left_margin - section.right_margin
            except Exception:
                usable_width = None
        last_paragraph = toc_heading
        removed_count = 0
        for paragraph in paragraphs[toc_heading_index + 1:content_start_index]:
            _delete_paragraph(paragraph)
            removed_count += 1

        for line in toc_lines:
            title, page_number = _parse_toc_line(line)
            new_paragraph = _insert_paragraph_after(last_paragraph)
            _replace_paragraph_text(new_paragraph, _build_toc_paragraph_text(title, page_number))
            try:
                new_paragraph.style = document.styles["Normal"]
            except Exception:
                pass
            new_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            indent_level = _get_toc_indent_level(title)
            new_paragraph.paragraph_format.left_indent = Pt(18 * indent_level)
            new_paragraph.paragraph_format.first_line_indent = Pt(0)
            new_paragraph.paragraph_format.space_before = Pt(0)
            new_paragraph.paragraph_format.space_after = Pt(0)
            new_paragraph.paragraph_format.line_spacing = 1.15
            _clear_paragraph_tab_stops(new_paragraph)
            if usable_width is not None and page_number:
                try:
                    new_paragraph.paragraph_format.tab_stops.add_tab_stop(
                        usable_width,
                        WD_TAB_ALIGNMENT.RIGHT,
                        WD_TAB_LEADER.DOTS,
                    )
                except Exception:
                    pass
            last_paragraph = new_paragraph

        page_break_paragraph = _insert_paragraph_after(last_paragraph)
        page_break_paragraph.paragraph_format.space_before = Pt(0)
        page_break_paragraph.paragraph_format.space_after = Pt(0)
        page_break_paragraph.add_run().add_break(WD_BREAK.PAGE)

        return len(toc_lines)

    def _split_paper_inline_sections(document: Any, profile: dict[str, Any]) -> int:
        section_markers = [
            marker for marker in profile.get("section_markers", [])
            if marker and marker not in {"Abstract", "Keywords", "References", "Appendix"}
        ]
        if not section_markers:
            return 0

        marker_pattern = re.compile(
            "|".join(re.escape(marker) for marker in sorted(section_markers, key=len, reverse=True))
        )

        paragraphs = list(document.paragraphs)
        abstract_index = next(
            (index for index, paragraph in enumerate(paragraphs) if _paragraph_text(paragraph).strip() == "Abstract"),
            None,
        )
        if abstract_index is None:
            return 0

        def _apply_kind(paragraph: Any, kind: str) -> None:
            paragraph.paragraph_format.left_indent = Pt(0)
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(2)
            if kind == "heading":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_before = Pt(6)
                heading_run = paragraph.runs[0] if paragraph.runs else None
                if heading_run is not None:
                    heading_run.bold = True
                    heading_run.font.size = Pt(12)
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        split_count = 0
        for paragraph_index in range(abstract_index + 1, len(paragraphs)):
            paragraph = paragraphs[paragraph_index]
            raw_text = _paragraph_text(paragraph)
            if not raw_text or _has_drawing(paragraph):
                continue
            if "..." in raw_text:
                continue

            matches = list(marker_pattern.finditer(raw_text))
            if not matches:
                continue
            if not (matches[0].start() > 0 or len(matches) > 1):
                continue

            segments: list[tuple[str, str]] = []
            prefix_text = re.sub(r"\s+", " ", raw_text[:matches[0].start()]).strip()
            if prefix_text:
                segments.append(("body", prefix_text))

            for match_index, match in enumerate(matches):
                marker_text = match.group(0)
                segment_end = matches[match_index + 1].start() if match_index + 1 < len(matches) else len(raw_text)
                segment_text = re.sub(r"\s+", " ", raw_text[match.start():segment_end]).strip()
                body_text = segment_text[len(marker_text):].strip(" :-\u2014")
                segments.append(("heading", marker_text))
                if body_text:
                    segments.append(("body", body_text))

            if len(segments) <= 1:
                continue

            first_kind, first_text = segments[0]
            _replace_paragraph_text(paragraph, first_text)
            _apply_kind(paragraph, first_kind)
            last_paragraph = paragraph

            for kind, text in segments[1:]:
                new_paragraph = _insert_paragraph_after(last_paragraph)
                _replace_paragraph_text(new_paragraph, text)
                _apply_kind(new_paragraph, kind)
                last_paragraph = new_paragraph

            split_count += len(segments) - 1

        return split_count

    def _apply_paper_typography_profile(document: Any, profile: dict[str, Any]) -> int:
        body_font_name = _normalize_pdf_font_name(profile.get("body_font_name", "Times New Roman"))
        body_size = Pt(profile.get("body_font_size", 11.0))
        body_line_ratio = float(profile.get("body_line_ratio", 1.35))
        adjusted_count = 0

        paragraphs = list(document.paragraphs)
        abstract_index = next(
            (index for index, paragraph in enumerate(paragraphs) if _paragraph_text(paragraph).strip() == "Abstract"),
            None,
        )

        for paragraph_index, paragraph in enumerate(paragraphs):
            text = _paragraph_text(paragraph).strip()
            if not text or _has_drawing(paragraph):
                continue
            if _is_bullet_paragraph(paragraph):
                continue
            if text.isdigit():
                continue
            if paragraph_index < (abstract_index or 0):
                continue
            if any(text.startswith(prefix) for prefix in ("Abstract", "Keywords", "References", "Appendix", "Fig ", "Figure ", "Table ")):
                continue
            if re.match(r"^\d+(?:\.\d+)*\s+.+", text):
                continue

            paragraph.paragraph_format.line_spacing = body_line_ratio
            for run in paragraph.runs:
                if run.text.strip() and not run.bold:
                    _set_run_font_name(run, body_font_name)
                    run.font.size = body_size
            adjusted_count += 1

        return adjusted_count

    def _normalize_bullet_paragraphs(document: Any) -> int:
        def _resolve_neighbor_font_size(paragraphs: list[Any], paragraph_index: int) -> Optional[Pt]:
            for distance in range(1, 5):
                for neighbor_index in (paragraph_index - distance, paragraph_index + distance):
                    if neighbor_index < 0 or neighbor_index >= len(paragraphs):
                        continue
                    neighbor = paragraphs[neighbor_index]
                    if _is_bullet_paragraph(neighbor):
                        continue
                    for run in neighbor.runs:
                        if run.font.size and 8.0 <= run.font.size.pt <= 14.0:
                            return run.font.size
            return None

        adjusted_count = 0
        paragraphs = list(document.paragraphs)
        for paragraph_index, paragraph in enumerate(paragraphs):
            if not _is_bullet_paragraph(paragraph):
                continue

            _normalize_bullet_prefix(paragraph)
            try:
                paragraph.style = document.styles["Normal"]
            except Exception:
                pass
            paragraph_format = paragraph.paragraph_format
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph_format.left_indent = Pt(18)
            paragraph_format.first_line_indent = Pt(-18)
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(4)
            paragraph_format.line_spacing = 1.12
            _clear_paragraph_tab_stops(paragraph)
            try:
                paragraph_format.tab_stops.add_tab_stop(Pt(18))
            except Exception:
                pass

            explicit_sizes = [
                run.font.size for run in paragraph.runs
                if run.text.strip() and run.font.size and 8.0 <= run.font.size.pt <= 14.0
            ]
            target_size = explicit_sizes[0] if explicit_sizes else _resolve_neighbor_font_size(paragraphs, paragraph_index)
            if target_size is not None:
                for run in paragraph.runs:
                    if run.text.strip():
                        run.font.size = target_size
            adjusted_count += 1

        return adjusted_count

    def _flatten_text_tables(document: Any) -> int:
        flattened_count = 0
        separator_cells = {"/", "|", "｜"}

        def _cell_text(cell: Any) -> str:
            parts = []
            for paragraph in cell.paragraphs:
                text = _paragraph_text(paragraph)
                if text:
                    parts.append(text)
            return "\n".join(parts).strip()

        for table in list(document.tables):
            if len(table.rows) != 1 or len(table.columns) < 3:
                continue

            cells = table.rows[0].cells
            texts = [_cell_text(cell) for cell in cells]
            non_empty = [text for text in texts if text]
            if len(non_empty) < 3:
                continue
            if any(len(text) > 120 for text in non_empty):
                continue

            joined_parts = []
            for text in non_empty:
                clean_text = re.sub(r"\s+", " ", text).strip()
                if not clean_text:
                    continue
                if clean_text in separator_cells:
                    if joined_parts:
                        joined_parts[-1] = f"{joined_parts[-1]} {clean_text}"
                    else:
                        joined_parts.append(clean_text)
                    continue
                joined_parts.append(clean_text)

            if len(joined_parts) < 2:
                continue

            flattened_text = " ".join(joined_parts)
            if not any(token in flattened_text for token in ("/", "@", "doi", "pp.", "vol.", "author")):
                continue

            new_paragraph = _insert_paragraph_after_table(table)
            _replace_paragraph_text(new_paragraph, flattened_text)
            new_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            new_paragraph.paragraph_format.left_indent = Pt(0)
            new_paragraph.paragraph_format.first_line_indent = Pt(0)
            new_paragraph.paragraph_format.space_before = Pt(0)
            new_paragraph.paragraph_format.space_after = Pt(2)
            _delete_table(table)
            flattened_count += 1

        return flattened_count

    def _convert_anchor_images_to_inline(document: Any) -> int:
        converted_count = 0
        for paragraph in document.paragraphs:
            if not _has_anchor(paragraph):
                continue

            relationship_ids = paragraph._p.xpath(".//a:blip/@r:embed")
            if not relationship_ids:
                continue

            extent_nodes = paragraph._p.xpath(".//wp:anchor/wp:extent")
            width_emu = None
            if extent_nodes:
                try:
                    width_emu = int(extent_nodes[0].get("cx"))
                except Exception:
                    width_emu = None

            image_bytes = None
            for rel_id in relationship_ids:
                related_part = paragraph.part.related_parts.get(rel_id)
                if related_part is None:
                    continue
                image_bytes = getattr(related_part, "blob", None)
                if image_bytes:
                    break
            if not image_bytes:
                continue

            for child in list(paragraph._p):
                if child.tag != qn("w:pPr"):
                    paragraph._p.remove(child)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = None
            paragraph.paragraph_format.first_line_indent = None
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
            picture_run = paragraph.add_run()
            if width_emu and width_emu > 0:
                picture_run.add_picture(io.BytesIO(image_bytes), width=Emu(width_emu))
            else:
                picture_run.add_picture(io.BytesIO(image_bytes))
            converted_count += 1

        return converted_count

    def _normalize_regular_paragraph_indents(document: Any) -> int:
        adjusted_count = 0
        paragraphs = list(document.paragraphs)
        for paragraph_index, paragraph in enumerate(paragraphs):
            text = _paragraph_text(paragraph)
            if not text or _is_bullet_paragraph(paragraph) or _has_drawing(paragraph):
                continue

            paragraph_format = paragraph.paragraph_format
            changed = False
            left_indent = paragraph_format.left_indent.pt if paragraph_format.left_indent else 0.0
            first_indent = paragraph_format.first_line_indent.pt if paragraph_format.first_line_indent else 0.0

            if 0.0 < left_indent <= 36.0:
                paragraph_format.left_indent = Pt(0)
                changed = True

            previous_has_drawing = paragraph_index > 0 and _has_drawing(paragraphs[paragraph_index - 1])
            if first_indent:
                if abs(first_indent) <= 24.0 or first_indent > 60.0 or previous_has_drawing:
                    paragraph_format.first_line_indent = Pt(0)
                    changed = True

            if changed:
                adjusted_count += 1

        return adjusted_count

    def _normalize_regular_paragraph_alignment(document: Any) -> int:
        adjusted_count = 0
        paragraphs = list(document.paragraphs)
        abstract_index = next(
            (index for index, paragraph in enumerate(paragraphs) if _paragraph_text(paragraph).strip().startswith("Abstract")),
            None,
        )
        for paragraph_index, paragraph in enumerate(paragraphs):
            text = _paragraph_text(paragraph)
            if not text or _is_bullet_paragraph(paragraph) or _has_drawing(paragraph):
                continue
            if len(text) < 90:
                continue
            if text.isdigit():
                continue
            if re.match(r"^\d+\.\s", text):
                continue
            if any(text.startswith(prefix) for prefix in ("Fig ", "Figure ", "Table ", "Appendix", "Abstract", "Keywords")):
                continue
            if any(run.bold for run in paragraph.runs if run.text.strip()):
                continue
            if abstract_index is not None and paragraph_index < abstract_index:
                continue
            if "@" in text or "Corresponding author" in text:
                continue

            font_sizes = [run.font.size.pt for run in paragraph.runs if run.text.strip() and run.font.size]
            if font_sizes and max(font_sizes) > 14.0:
                continue

            if paragraph.alignment in (None, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                adjusted_count += 1

        return adjusted_count

    def _replace_with_inline_image(paragraph: Any, image_bytes: bytes) -> None:
        for child in list(paragraph._p):
            if child.tag != qn("w:pPr"):
                paragraph._p.remove(child)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run().add_picture(io.BytesIO(image_bytes), width=Cm(3.0))

    def _repair_resume_header(document: Any) -> bool:
        nonlocal header_extra_space_before_pt
        name_para, image_para, contact_para, first_heading_para = _find_resume_key_paragraphs(document)
        if not name_para or not image_para or not contact_para or not first_heading_para:
            return False

        image_bytes = extract_first_image_from_docx(docx_path)
        if _has_anchor(image_para):
            _shift_anchor_paragraph_down(image_para)
            anchor_height_pt = _anchor_height_points(image_para)
            if anchor_height_pt > 0:
                header_extra_space_before_pt = max(10.0, min(anchor_height_pt - 82.0, 52.0))
        elif image_bytes:
            _replace_with_inline_image(image_para, image_bytes)
        name_para.paragraph_format.space_after = Pt(2)
        _move_paragraph_after(contact_para, name_para)
        contact_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        contact_para.paragraph_format.space_before = Pt(0)
        contact_para.paragraph_format.space_after = Pt(8)
        contact_para.paragraph_format.line_spacing = 1.20

        duplicate_contact_found = False
        for paragraph in list(document.paragraphs):
            if not _is_contact_paragraph(paragraph):
                continue
            if duplicate_contact_found:
                _delete_paragraph(paragraph)
                continue
            duplicate_contact_found = True

        return True

    def _merge_wrapped_resume_paragraphs(document: Any) -> int:
        merged_count = 0
        paragraphs = list(document.paragraphs)
        index = 1

        while index < len(paragraphs):
            previous_paragraph = paragraphs[index - 1]
            current_paragraph = paragraphs[index]
            previous_text = _paragraph_text(previous_paragraph)
            current_text = _paragraph_text(current_paragraph)

            if not current_text:
                if _has_drawing(current_paragraph):
                    index += 1
                    continue
                _delete_paragraph(current_paragraph)
                paragraphs.pop(index)
                continue

            current_is_hard_start = (
                _is_heading(current_paragraph)
                or _is_contact_paragraph(current_paragraph)
                or _is_bullet_paragraph(current_paragraph)
                or _looks_like_new_resume_block(current_text)
            )
            if current_is_hard_start or not previous_text:
                index += 1
                continue

            if _is_heading(previous_paragraph) or _is_contact_paragraph(previous_paragraph):
                index += 1
                continue

            current_indent = _left_indent_points(current_paragraph)
            previous_indent = _left_indent_points(previous_paragraph)
            previous_is_bold = any(run.bold for run in previous_paragraph.runs if run.text.strip())

            should_merge = (
                current_indent > previous_indent + 4.0
                or current_text.startswith(("（", "("))
                or (
                    abs(current_indent - previous_indent) <= 1.0
                    and previous_is_bold
                    and len(current_text) <= 120
                    and not _looks_like_new_resume_block(current_text)
                )
            )

            if not should_merge:
                index += 1
                continue

            separator = ""
            if previous_text and not previous_text.endswith(("（", "(", "/", "-", "+", "|", "｜")):
                separator = " "
            previous_paragraph.add_run(f"{separator}{current_text.strip()}")
            _delete_paragraph(current_paragraph)
            paragraphs.pop(index)
            merged_count += 1

        return merged_count

    def _normalize_resume_spacing(document: Any) -> None:
        applied_header_clearance = False
        for paragraph in document.paragraphs:
            text = _paragraph_text(paragraph)
            if not text:
                continue

            paragraph_format = paragraph.paragraph_format
            if _is_heading(paragraph):
                heading_space_before = 10.0
                if not applied_header_clearance and header_extra_space_before_pt > heading_space_before:
                    heading_space_before = header_extra_space_before_pt
                    _set_paragraph_top_border(paragraph)
                    applied_header_clearance = True
                paragraph_format.space_before = Pt(heading_space_before)
                paragraph_format.space_after = Pt(4)
                paragraph_format.line_spacing = 1.10
                paragraph_format.keep_with_next = True
            elif _is_contact_paragraph(paragraph):
                paragraph_format.space_before = Pt(0)
                paragraph_format.space_after = Pt(8)
                paragraph_format.line_spacing = 1.20
            elif _is_bullet_paragraph(paragraph):
                paragraph_format.left_indent = Pt(36)
                paragraph_format.first_line_indent = Pt(-18)
                paragraph_format.space_before = Pt(0)
                paragraph_format.space_after = Pt(5)
                paragraph_format.line_spacing = 1.28
            else:
                paragraph_format.space_before = Pt(0)
                paragraph_format.space_after = Pt(2)
                paragraph_format.line_spacing = 1.15

    document = Document(str(docx_path))
    paper_profile = _extract_pdf_paper_profile(src_file)
    _normalize_sections(document)
    flattened_table_count = _flatten_text_tables(document)
    front_matter_fix_count = _repair_paper_front_matter(src_file, document)
    toc_fix_count = _repair_paper_toc(document, paper_profile.get("toc_lines", []))
    inline_section_split_count = _split_paper_inline_sections(document, paper_profile)
    split_heading_count = _split_heading_body_paragraphs(document)
    split_bullet_count = _split_mixed_bullet_paragraphs(document)
    split_numbered_count = _split_mixed_numbered_paragraphs(document)
    split_caption_count = _split_caption_body_paragraphs(document)
    bullet_count = _normalize_bullet_paragraphs(document)

    header_repaired = _repair_resume_header(document)
    inline_image_count = 0 if header_repaired else _convert_anchor_images_to_inline(document)
    indent_fix_count = _normalize_regular_paragraph_indents(document)
    alignment_fix_count = 0 if header_repaired else _normalize_regular_paragraph_alignment(document)
    typography_fix_count = 0 if header_repaired else _apply_paper_typography_profile(document, paper_profile)
    if not header_repaired:
        late_inline_section_split_count = _split_paper_inline_sections(document, paper_profile)
        if late_inline_section_split_count > 0:
            inline_section_split_count += late_inline_section_split_count
            indent_fix_count += _normalize_regular_paragraph_indents(document)
            alignment_fix_count += _normalize_regular_paragraph_alignment(document)
            typography_fix_count += _apply_paper_typography_profile(document, paper_profile)
    if (
            not header_repaired
            and bullet_count <= 0
            and inline_image_count <= 0
            and indent_fix_count <= 0
            and flattened_table_count <= 0
            and alignment_fix_count <= 0
            and toc_fix_count <= 0
            and inline_section_split_count <= 0
            and typography_fix_count <= 0
    ):
        return False, "版式修复跳过：未命中简历版式特征"

    merged_count = _merge_wrapped_resume_paragraphs(document) if header_repaired else 0
    if header_repaired:
        _normalize_resume_spacing(document)
    document.save(str(docx_path))

    note_parts = []
    if header_repaired:
        note_parts.append("已修复简历版式（保留非表格头部）")
    if flattened_table_count > 0:
        note_parts.append(f"已扁平化 {flattened_table_count} 处误判表格")
    if front_matter_fix_count > 0:
        note_parts.append(f"已修正 {front_matter_fix_count} 处论文前言格式")
    if toc_fix_count > 0:
        note_parts.append(f"已重建 {toc_fix_count} 处目录段落")
    if inline_section_split_count > 0:
        note_parts.append(f"已拆分 {inline_section_split_count} 处论文内联节标题")
    if split_heading_count > 0:
        note_parts.append(f"已拆分 {split_heading_count} 处标题正文混合段落")
    if split_bullet_count > 0:
        note_parts.append(f"已拆分 {split_bullet_count} 处混合项目符号段落")
    if split_numbered_count > 0:
        note_parts.append(f"已拆分 {split_numbered_count} 处混合编号段落")
    if split_caption_count > 0:
        note_parts.append(f"已拆分 {split_caption_count} 处图表标题段落")
    if bullet_count > 0:
        note_parts.append(f"已优化 {bullet_count} 处项目符号间距")
    if inline_image_count > 0:
        note_parts.append(f"已转换 {inline_image_count} 张浮动图片为内联图片")
    if indent_fix_count > 0:
        note_parts.append(f"已修正 {indent_fix_count} 处正文异常缩进")
    if alignment_fix_count > 0:
        note_parts.append(f"已修正 {alignment_fix_count} 处正文对齐方式")
    if typography_fix_count > 0:
        note_parts.append(f"已统一 {typography_fix_count} 处论文正文排版")
    if merged_count > 0:
        note_parts.append(f"合并 {merged_count} 处续行段落")
    repair_note = "，".join(note_parts) if note_parts else "已完成段落后处理"
    return True, repair_note


def run_ocrmypdf(src_file: Path, dst_file: Path) -> Tuple[bool, str]:
    commands: List[List[str]] = []
    cli_path = shutil.which("ocrmypdf")
    if cli_path:
        commands.append([cli_path, "--skip-text", "--optimize", "0", str(src_file), str(dst_file)])
    if importlib.util.find_spec(OCRMYPDF_MODULE_NAME) is not None:
        commands.append(
            [sys.executable, "-m", OCRMYPDF_MODULE_NAME, "--skip-text", "--optimize", "0", str(src_file), str(dst_file)]
        )

    if not commands:
        return False, f"未安装 {OCRMYPDF_ENGINE_LABEL}"

    last_error = ""
    for command in commands:
        try:
            completed = subprocess.run(
                command,
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="ignore",
                check=False,
            )
        except Exception as exc:
            last_error = str(exc)
            continue

        if completed.returncode == 0 and dst_file.exists():
            return True, f"{OCRMYPDF_ENGINE_LABEL} 已完成 OCR"

        stderr_text = safe_strip_text(completed.stderr) or safe_strip_text(completed.stdout)
        if stderr_text:
            last_error = stderr_text

    return False, last_error or f"{OCRMYPDF_ENGINE_LABEL} 运行失败"


def convert_pdf_to_docx_with_pdf2docx(
        conversion_src_file: Path,
        profile_src_file: Path,
        dst_file: Path,
        engine_label: str,
) -> Tuple[bool, str]:
    try:
        from pdf2docx import Converter
    except ImportError:
        return (
            False,
            f"失败：{profile_src_file.name} -> {dst_file.name} | 未安装 {PDF2DOCX_ENGINE_LABEL}，"
            f"请先执行 pip install pdf2docx",
        )

    converter = None
    try:
        with tempfile.TemporaryDirectory(prefix="pdf_to_docx_") as temp_dir:
            temp_output = Path(temp_dir) / dst_file.name
            converter = Converter(str(conversion_src_file))
            converter.convert(
                str(temp_output),
                multi_processing=False,
            )
            converter.close()
            converter = None

            if not temp_output.exists():
                return False, (
                    f"失败：{profile_src_file.name} -> {dst_file.name} | "
                    f"{PDF2DOCX_ENGINE_LABEL} 未生成输出文件"
                )

            if dst_file.exists():
                dst_file.unlink()
            shutil.move(str(temp_output), str(dst_file))

        repaired, repair_note = repair_pdf2docx_layout(profile_src_file, dst_file)
        success_message = f"成功：{profile_src_file.name} -> {dst_file.name} | 引擎：{engine_label}"
        if repaired:
            success_message += f" | {repair_note}"
        elif repair_note:
            success_message += f" | {repair_note}"

        return True, success_message
    except Exception as exc:
        if dst_file.exists() and dst_file.stat().st_size == 0:
            dst_file.unlink(missing_ok=True)
        return False, f"失败：{profile_src_file.name} -> {dst_file.name} | {engine_label}: {exc}"
    finally:
        if converter is not None:
            try:
                converter.close()
            except Exception:
                pass


def convert_pdf_to_docx(src_file: Path, dst_file: Path) -> Tuple[bool, str]:
    route = analyze_pdf_to_docx_route(src_file)

    if route.get("engine") == OCRMYPDF_ENGINE_LABEL:
        if not has_ocrmypdf_engine():
            return (
                False,
                f"失败：{src_file.name} -> {dst_file.name} | 检测为扫描件，但未安装 {OCRMYPDF_ENGINE_LABEL}",
            )
        if not has_pdf2docx_engine():
            return (
                False,
                f"失败：{src_file.name} -> {dst_file.name} | OCR 后仍需 {PDF2DOCX_ENGINE_LABEL}，请先执行 pip install pdf2docx",
            )

        with tempfile.TemporaryDirectory(prefix="pdf_to_docx_ocr_") as temp_dir:
            searchable_pdf = Path(temp_dir) / f"{src_file.stem}_ocr.pdf"
            ocr_ok, ocr_message = run_ocrmypdf(src_file, searchable_pdf)
            if not ocr_ok:
                return False, f"失败：{src_file.name} -> {dst_file.name} | {ocr_message}"
            converted_ok, converted_message = convert_pdf_to_docx_with_pdf2docx(
                searchable_pdf,
                searchable_pdf,
                dst_file,
                OCRMYPDF_ENGINE_LABEL,
            )
            if converted_ok:
                converted_message += f" | 路由：{route.get('reason', '')}"
            return converted_ok, converted_message

    if route.get("engine") == FREEP2W_ENGINE_LABEL:
        converted_ok, converted_message = convert_pdf_to_docx_with_pdf2docx(
            src_file,
            src_file,
            dst_file,
            PDF2DOCX_ENGINE_LABEL,
        )
        if converted_ok:
            converted_message += " | 复杂版面已命中 FreeP2W 预留路由，当前版本暂回退到 pdf2docx"
        return converted_ok, converted_message

    return convert_pdf_to_docx_with_pdf2docx(
        src_file,
        src_file,
        dst_file,
        route.get("engine_label", PDF2DOCX_ENGINE_LABEL),
    )


def convert_document_with_libreoffice(
        src_file: Path,
        dst_file: Path,
        libreoffice_bin: str,
) -> Tuple[bool, str]:
    output_dir = dst_file.parent
    target_ext = dst_file.suffix.replace(".", "").lower()

    cmd = [
        libreoffice_bin,
        "--headless",
        "--convert-to", libreoffice_filter_name(target_ext),
        "--outdir", str(output_dir),
        str(src_file)
    ]

    try:
        result = subprocess.run(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True
        )
        if result.returncode != 0:
            error_text = (
                safe_strip_text(result.stderr)
                or safe_strip_text(result.stdout)
                or f"{LIBREOFFICE_ENGINE_LABEL} 未返回可读错误信息"
            )
            return False, f"失败：{src_file.name} -> {dst_file.name} | {error_text}"

        generated_file = output_dir / f"{src_file.stem}.{target_ext}"
        if not generated_file.exists():
            return False, f"失败：{src_file.name} -> {dst_file.name} | 未找到输出文件"

        if generated_file.resolve() != dst_file.resolve():
            if dst_file.exists():
                dst_file.unlink()
            generated_file.replace(dst_file)

        return True, f"成功：{src_file.name} -> {dst_file.name} | 引擎：{LIBREOFFICE_ENGINE_LABEL}"
    except Exception as exc:
        return False, f"失败：{src_file.name} -> {dst_file.name} | {exc}"


def convert_document(
        src_file: Path,
        dst_file: Path,
        libreoffice_bin: Optional[str],
) -> Tuple[bool, str]:
    target_ext = dst_file.suffix.replace(".", "").lower()
    engine_label = get_document_conversion_engine_label(src_file, target_ext)

    if src_file.suffix.lower() == ".pdf" and target_ext == "docx":
        return convert_pdf_to_docx(src_file, dst_file)

    if not libreoffice_bin:
        return (
            False,
            f"失败：{src_file.name} -> {dst_file.name} | 未检测到 {engine_label}，"
            f"当前转换需要该开源引擎",
        )

    return convert_document_with_libreoffice(src_file, dst_file, libreoffice_bin)


def compress_image_lossless(src_file: Path, dst_file: Path, quality: int = 85) -> Tuple[bool, str]:
    try:
        src_size = src_file.stat().st_size
        with Image.open(src_file) as img:
            save_kwargs = {}
            target_ext = dst_file.suffix.lower()

            if target_ext in [".jpg", ".jpeg"]:
                save_kwargs["quality"] = quality
                save_kwargs["optimize"] = False
                if img.mode in ("RGBA", "LA", "P"):
                    if img.mode == "P":
                        img = img.convert("RGBA")
                    background = Image.new("RGB", img.size, (255, 255, 255))
                    alpha_channel = img.split()[-1] if "A" in img.mode else None
                    background.paste(img, mask=alpha_channel)
                    img = background
                else:
                    img = img.convert("RGB")
            elif target_ext == ".png":
                save_kwargs["optimize"] = False
                save_kwargs["compress_level"] = 6
            elif target_ext == ".gif":
                save_kwargs["optimize"] = False
            elif target_ext == ".webp":
                save_kwargs["quality"] = quality
                save_kwargs["method"] = 4

            img.save(dst_file, **save_kwargs)

        dst_size = dst_file.stat().st_size
        if dst_size >= src_size:
            if dst_file.exists():
                dst_file.unlink()
            return False, f"失败：{src_file.name} | 压缩后文件未变小 ({format_file_size(src_size)} -> {format_file_size(dst_size)})"
        ratio = (1 - dst_size / src_size) * 100 if src_size > 0 else 0

        return True, f"成功：{src_file.name} | 压缩率: {ratio:.1f}% ({format_file_size(src_size)} -> {format_file_size(dst_size)})"
    except Exception as exc:
        return False, f"失败：{src_file.name} | {exc}"


def _probe_media_duration_seconds(ffmpeg_bin: str, src_file: Path) -> float:
    try:
        result = subprocess.run(
            [ffmpeg_bin, "-i", str(src_file)],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            encoding="utf-8",
            errors="replace",
        )
        output_text = f"{safe_strip_text(result.stdout)}\n{safe_strip_text(result.stderr)}"
        match = re.search(r"Duration:\s*(\d+:\d+:\d+(?:\.\d+)?)", output_text)
        if not match:
            return 0.0
        return _parse_ffmpeg_time_to_seconds(match.group(1))
    except Exception:
        return 0.0


def probe_media_info(ffmpeg_bin: str, src_file: Path) -> Dict[str, Any]:
    ffprobe_bin = str(Path(ffmpeg_bin).with_name("ffprobe.exe"))
    if not Path(ffprobe_bin).exists():
        ffprobe_bin = "ffprobe"

    try:
        result = subprocess.run(
            [
                ffprobe_bin,
                "-v", "error",
                "-show_entries",
                "format=duration,size,bit_rate:stream=index,codec_name,codec_type,width,height,avg_frame_rate,bit_rate",
                "-of", "json",
                str(src_file),
            ],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            encoding="utf-8",
            errors="replace",
        )
        if result.returncode != 0:
            raise RuntimeError(safe_strip_text(result.stderr) or "ffprobe 执行失败")
        data = json.loads(result.stdout or "{}")
        return data if isinstance(data, dict) else {}
    except Exception:
        return {}


def ffmpeg_supports_encoder(ffmpeg_bin: str, encoder_name: str) -> bool:
    cache_key = str(Path(ffmpeg_bin).expanduser().resolve())
    if cache_key not in FFMPEG_ENCODER_CACHE:
        try:
            result = subprocess.run(
                [ffmpeg_bin, "-hide_banner", "-encoders"],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                encoding="utf-8",
                errors="replace",
            )
            output = f"{safe_strip_text(result.stdout)}\n{safe_strip_text(result.stderr)}"
            encoders = set()
            for line in output.splitlines():
                parts = line.split()
                if len(parts) >= 2 and parts[0].startswith("V"):
                    encoders.add(parts[1])
            FFMPEG_ENCODER_CACHE[cache_key] = encoders
        except Exception:
            FFMPEG_ENCODER_CACHE[cache_key] = set()
    return encoder_name in FFMPEG_ENCODER_CACHE[cache_key]


def map_video_compress_settings(quality: int) -> Tuple[int, str]:
    quality = max(50, min(100, int(quality)))
    crf = int(round(34 - (quality - 50) * 16 / 50))

    if quality >= 92:
        preset = "faster"
    elif quality >= 75:
        preset = "veryfast"
    else:
        preset = "superfast"

    return crf, preset


def build_video_compress_plan(ffmpeg_bin: str, src_file: Path, quality: int) -> Dict[str, Any]:
    quality = max(50, min(100, int(quality)))
    crf, preset = map_video_compress_settings(quality)
    info = probe_media_info(ffmpeg_bin, src_file)
    streams = info.get("streams", []) if isinstance(info, dict) else []
    video_stream = next((item for item in streams if item.get("codec_type") == "video"), {})
    codec_name = str(video_stream.get("codec_name") or "").lower()

    raw_video_bitrate = video_stream.get("bit_rate") or info.get("format", {}).get("bit_rate")
    try:
        source_bitrate = int(raw_video_bitrate)
    except Exception:
        source_bitrate = 0

    bitrate_ratio = 0.58 + (quality - 50) * 0.30 / 50
    target_bitrate = int(source_bitrate * bitrate_ratio) if source_bitrate > 0 else 0
    target_bitrate = max(900_000, target_bitrate) if target_bitrate > 0 else 0

    use_bitrate_mode = codec_name in {"h264", "hevc", "h265", "av1", "vp9"} and target_bitrate > 0
    encoder = "hevc_nvenc" if ffmpeg_supports_encoder(ffmpeg_bin, "hevc_nvenc") else "libx264"
    encoder_preset = "p5" if encoder == "hevc_nvenc" else preset

    plan: Dict[str, Any] = {
        "codec_name": codec_name or "unknown",
        "encoder": encoder,
        "preset": encoder_preset,
        "crf": crf,
        "target_bitrate": target_bitrate,
        "use_bitrate_mode": use_bitrate_mode,
    }
    return plan


def compress_video_lossless(
        src_file: Path,
        dst_file: Path,
        ffmpeg_bin: str,
        crf: int = 23,
        preset: str = "veryfast",
        target_bitrate: int = 0,
        use_bitrate_mode: bool = False,
        source_codec: str = "",
        encoder: str = "libx264",
        progress_callback: Optional[callable] = None,
        log_callback: Optional[callable] = None,
) -> Tuple[bool, str]:
    cmd = [
        ffmpeg_bin,
        "-y",
        "-i", str(src_file),
        "-c:a", "copy",
    ]

    if encoder == "hevc_nvenc":
        cmd.extend([
            "-c:v", "hevc_nvenc",
            "-preset", preset,
            "-rc", "vbr",
            "-cq", str(max(20, min(34, crf + 1))),
        ])
    else:
        cmd.extend([
            "-c:v", "libx264",
            "-preset", preset,
        ])

    if use_bitrate_mode and target_bitrate > 0:
        target_k = max(1, target_bitrate // 1000)
        maxrate_k = max(target_k, int(target_k * 1.15))
        bufsize_k = max(maxrate_k * 2, target_k * 2)
        cmd.extend([
            "-b:v", f"{target_k}k",
            "-maxrate", f"{maxrate_k}k",
            "-bufsize", f"{bufsize_k}k",
        ])
    else:
        cmd.extend(["-crf", str(crf)])

    cmd.append(str(dst_file))
    try:
        src_size = src_file.stat().st_size
        duration_seconds = _probe_media_duration_seconds(ffmpeg_bin, src_file)
        if log_callback is not None:
            try:
                if use_bitrate_mode and target_bitrate > 0:
                    log_callback(
                        f"正在压缩视频：{src_file.name} | 源编码={source_codec or 'unknown'} | "
                        f"encoder={encoder} | preset={preset} | target_bitrate={target_bitrate // 1000}k"
                    )
                else:
                    log_callback(f"正在压缩视频：{src_file.name} | encoder={encoder} | preset={preset} | crf={crf}")
            except Exception:
                pass

        process = subprocess.Popen(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            encoding="utf-8",
            errors="replace",
            bufsize=1,
        )

        stderr_lines: List[str] = []
        time_pattern = re.compile(r'time=(\d+:\d+:\d+(?:\.\d+)?)')

        try:
            if process.stderr is not None:
                for raw_line in process.stderr:
                    line = raw_line.rstrip('\r\n')
                    stderr_lines.append(line)
                    match = time_pattern.search(line)
                    if match and duration_seconds > 0 and progress_callback is not None:
                        current_seconds = _parse_ffmpeg_time_to_seconds(match.group(1))
                        percent = int(max(0.0, min(100.0, current_seconds * 100.0 / max(duration_seconds, 0.1))))
                        try:
                            progress_callback(percent)
                        except Exception:
                            pass
            return_code = process.wait()
        finally:
            if process.stdout:
                process.stdout.close()
            if process.stderr:
                process.stderr.close()

        if return_code == 0:
            dst_size = dst_file.stat().st_size
            if dst_size >= src_size:
                if dst_file.exists():
                    dst_file.unlink()
                return False, f"失败：{src_file.name} | 压缩后文件未变小 ({format_file_size(src_size)} -> {format_file_size(dst_size)})"
            if progress_callback is not None:
                try:
                    progress_callback(100)
                except Exception:
                    pass
            ratio = (1 - dst_size / src_size) * 100 if src_size > 0 else 0
            return True, f"成功：{src_file.name} | 压缩率: {ratio:.1f}% ({format_file_size(src_size)} -> {format_file_size(dst_size)})"

        error_text = safe_strip_text("\n".join(stderr_lines)) or "ffmpeg 未返回可读错误信息"
        return False, f"失败：{src_file.name} | {error_text}"
    except Exception as exc:
        return False, f"失败：{src_file.name} | {exc}"


def remove_pdf_watermark(src_file: Path, dst_file: Path) -> Tuple[bool, str]:
    try:
        try:
            import fitz
        except ImportError:
            return False, f"失败：{src_file.name} | 需要安装PyMuPDF库 (pip install PyMuPDF)"

        doc = fitz.open(str(src_file))
        removed_annotations = 0

        for page in doc:
            annot = page.first_annot
            while annot is not None:
                next_annot = annot.next
                page.delete_annot(annot)
                removed_annotations += 1
                annot = next_annot

        if removed_annotations == 0:
            doc.close()
            return False, f"失败：{src_file.name} | 未检测到可安全移除的注释型水印"

        doc.save(str(dst_file), garbage=4, deflate=True)
        doc.close()

        return True, f"成功：{src_file.name} -> {dst_file.name} | 已移除 {removed_annotations} 个注释对象"
    except Exception as exc:
        return False, f"失败：{src_file.name} | {exc}"


def scan_image_to_document(
        src_file: Path,
        dst_file: Path,
        tesseract_bin: Optional[str],
        libreoffice_bin: Optional[str]
) -> Tuple[bool, str]:
    try:
        try:
            import pytesseract
            if tesseract_bin:
                pytesseract.pytesseract.tesseract_cmd = tesseract_bin
        except ImportError:
            return False, f"失败：{src_file.name} | 需要安装pytesseract库 (pip install pytesseract)"

        target_ext = dst_file.suffix.lower().replace(".", "")

        if target_ext == "pdf":
            pdf_bytes = pytesseract.image_to_pdf_or_hocr(
                str(src_file),
                extension="pdf",
                lang="chi_sim+eng",
            )
            if not pdf_bytes:
                return False, f"失败：{src_file.name} | OCR 未生成 PDF 内容"
            with open(dst_file, "wb") as f:
                f.write(pdf_bytes)
            return True, f"成功：{src_file.name} -> {dst_file.name}"

        with Image.open(src_file) as img:
            img = _prepare_pil_image_for_export(img, force_rgb=True)

            text = pytesseract.image_to_string(img, lang='chi_sim+eng')
            if not text.strip():
                return False, f"失败：{src_file.name} | OCR 未识别到可写入的文本"

        if target_ext == "txt":
            with open(dst_file, "w", encoding="utf-8") as f:
                f.write(text)
            return True, f"成功：{src_file.name} -> {dst_file.name}"

        temp_txt = dst_file.parent / f"{dst_file.stem}_temp.txt"
        with open(temp_txt, "w", encoding="utf-8") as f:
            f.write(text)

        if target_ext in ["docx", "doc"]:
            if not libreoffice_bin:
                temp_txt.unlink()
                return False, f"失败：需要LibreOffice来生成DOCX格式"

            cmd = [
                libreoffice_bin,
                "--headless",
                "--convert-to", "docx",
                "--outdir", str(dst_file.parent),
                str(temp_txt)
            ]
            result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)

            generated = dst_file.parent / f"{temp_txt.stem}.docx"
            if generated.exists():
                if dst_file.exists():
                    dst_file.unlink()
                generated.replace(dst_file)
                temp_txt.unlink()
                return True, f"成功：{src_file.name} -> {dst_file.name}"
            else:
                temp_txt.unlink()
                return False, f"失败：{src_file.name} | 无法生成DOCX文件"

        temp_txt.unlink()
        return False, f"失败：不支持的目标格式 {target_ext}"

    except Exception as exc:
        return False, f"失败：{src_file.name} | {exc}"


def load_cv2_image(src_file: Path) -> Any:
    try:
        import cv2
    except ImportError as exc:
        raise RuntimeError("需要安装 opencv-python 库 (pip install opencv-python)") from exc

    raw_bytes = np.fromfile(str(src_file), dtype=np.uint8)
    if raw_bytes.size == 0:
        raise ValueError(f"{src_file.name} 读取失败")

    image = cv2.imdecode(raw_bytes, cv2.IMREAD_COLOR)
    if image is None:
        raise ValueError(f"{src_file.name} 不是有效图片或解码失败")
    return image


def order_quad_points(points: np.ndarray) -> np.ndarray:
    pts = np.asarray(points, dtype=np.float32)
    if pts.shape != (4, 2):
        raise ValueError("证件边界点数量不正确")

    ordered = np.zeros((4, 2), dtype=np.float32)
    point_sum = pts.sum(axis=1)
    point_diff = np.diff(pts, axis=1).reshape(-1)

    ordered[0] = pts[np.argmin(point_sum)]
    ordered[2] = pts[np.argmax(point_sum)]
    ordered[1] = pts[np.argmin(point_diff)]
    ordered[3] = pts[np.argmax(point_diff)]
    return ordered


def four_point_transform(image: Any, points: np.ndarray) -> Any:
    try:
        import cv2
    except ImportError as exc:
        raise RuntimeError("需要安装 opencv-python 库 (pip install opencv-python)") from exc

    rect = order_quad_points(points)
    top_left, top_right, bottom_right, bottom_left = rect

    width_top = np.linalg.norm(top_right - top_left)
    width_bottom = np.linalg.norm(bottom_right - bottom_left)
    target_width = int(max(width_top, width_bottom))

    height_right = np.linalg.norm(top_right - bottom_right)
    height_left = np.linalg.norm(top_left - bottom_left)
    target_height = int(max(height_left, height_right))

    if target_width < 40 or target_height < 40:
        raise ValueError("证件区域过小，无法矫正")

    destination = np.array([
        [0, 0],
        [target_width - 1, 0],
        [target_width - 1, target_height - 1],
        [0, target_height - 1],
    ], dtype=np.float32)

    matrix = cv2.getPerspectiveTransform(rect, destination)
    return cv2.warpPerspective(image, matrix, (target_width, target_height))


def detect_certificate_quad(image: Any) -> np.ndarray:
    try:
        import cv2
    except ImportError as exc:
        raise RuntimeError("需要安装 opencv-python 库 (pip install opencv-python)") from exc

    original_height, original_width = image.shape[:2]
    max_side = max(original_width, original_height)
    scale = max_side / 1400.0 if max_side > 1400 else 1.0

    if scale > 1.0:
        resized = cv2.resize(
            image,
            (int(original_width / scale), int(original_height / scale)),
            interpolation=cv2.INTER_AREA,
        )
    else:
        resized = image.copy()

    gray = cv2.cvtColor(resized, cv2.COLOR_BGR2GRAY)
    blur = cv2.GaussianBlur(gray, (5, 5), 0)
    edges = cv2.Canny(blur, 60, 160)
    edges = cv2.morphologyEx(edges, cv2.MORPH_CLOSE, np.ones((5, 5), np.uint8), iterations=2)

    binary = cv2.adaptiveThreshold(
        blur,
        255,
        cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY,
        31,
        7,
    )
    binary = cv2.bitwise_not(binary)
    candidate_mask = cv2.bitwise_or(edges, binary)
    candidate_mask = cv2.morphologyEx(
        candidate_mask,
        cv2.MORPH_CLOSE,
        np.ones((7, 7), np.uint8),
        iterations=2,
    )

    contours, _ = cv2.findContours(candidate_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    if not contours:
        raise ValueError("未检测到证件边界")

    image_area = float(resized.shape[0] * resized.shape[1])
    selected_quad: Optional[np.ndarray] = None

    for contour in sorted(contours, key=cv2.contourArea, reverse=True)[:15]:
        contour_area = float(cv2.contourArea(contour))
        if contour_area < image_area * 0.10:
            continue

        perimeter = cv2.arcLength(contour, True)
        approx = cv2.approxPolyDP(contour, 0.02 * perimeter, True)
        if len(approx) == 4 and cv2.isContourConvex(approx):
            selected_quad = approx.reshape(4, 2).astype(np.float32)
            break

    if selected_quad is None:
        largest = max(contours, key=cv2.contourArea)
        rect = cv2.minAreaRect(largest)
        box = cv2.boxPoints(rect).astype(np.float32)
        if cv2.contourArea(box) < image_area * 0.10:
            raise ValueError("未检测到足够大的证件区域")
        selected_quad = box

    return selected_quad * scale


def auto_crop_certificate_image(image: Any) -> Any:
    try:
        import cv2
    except ImportError as exc:
        raise RuntimeError("需要安装 opencv-python 库 (pip install opencv-python)") from exc

    height, width = image.shape[:2]
    if height < 40 or width < 40:
        raise ValueError("证件图像过小，无法自动裁剪")

    patch_h = max(6, int(height * 0.06))
    patch_w = max(6, int(width * 0.06))
    corner_samples = np.vstack([
        image[:patch_h, :patch_w].reshape(-1, 3),
        image[:patch_h, width - patch_w:].reshape(-1, 3),
        image[height - patch_h:, :patch_w].reshape(-1, 3),
        image[height - patch_h:, width - patch_w:].reshape(-1, 3),
    ]).astype(np.float32)
    background_color = np.median(corner_samples, axis=0)

    color_distance = np.linalg.norm(image.astype(np.float32) - background_color, axis=2)
    diff_mask = np.where(color_distance > 18.0, 255, 0).astype(np.uint8)

    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    edges = cv2.Canny(gray, 50, 150)
    edges = cv2.dilate(edges, np.ones((5, 5), np.uint8), iterations=2)

    candidate_mask = cv2.bitwise_or(diff_mask, edges)
    candidate_mask = cv2.morphologyEx(
        candidate_mask,
        cv2.MORPH_CLOSE,
        np.ones((7, 7), np.uint8),
        iterations=2,
    )

    contours, _ = cv2.findContours(candidate_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    if not contours:
        raise ValueError("自动裁剪失败，未检测到有效证件区域")

    image_area = float(height * width)
    main_contour = max(contours, key=cv2.contourArea)
    contour_area = float(cv2.contourArea(main_contour))
    if contour_area < image_area * 0.35:
        raise ValueError("自动裁剪失败，证件区域过小")

    x, y, w, h = cv2.boundingRect(main_contour)
    pad_x = max(2, int(width * 0.008))
    pad_y = max(2, int(height * 0.008))
    x1 = max(0, x - pad_x)
    y1 = max(0, y - pad_y)
    x2 = min(width, x + w + pad_x)
    y2 = min(height, y + h + pad_y)

    if x2 - x1 < 40 or y2 - y1 < 40:
        raise ValueError("自动裁剪失败，裁剪结果尺寸异常")

    return image[y1:y2, x1:x2].copy()


def remove_certificate_background(image: Any) -> Image.Image:
    height, width = image.shape[:2]
    if height < 40 or width < 40:
        raise ValueError("证件图像过小，无法抠背景")

    rgb_image = image[:, :, ::-1]
    rgba_image = Image.fromarray(rgb_image).convert("RGBA")

    scale = 4
    large_mask = Image.new("L", (width * scale, height * scale), 0)
    draw = ImageDraw.Draw(large_mask)

    inset_x = max(8, int(width * 0.02))
    inset_y = max(8, int(height * 0.02))
    radius = max(16, int(min(width, height) * 0.075))
    draw.rounded_rectangle(
        (
            inset_x * scale,
            inset_y * scale,
            (width - inset_x - 1) * scale,
            (height - inset_y - 1) * scale,
        ),
        radius=radius * scale,
        fill=255,
    )
    alpha_mask = large_mask.resize((width, height), Image.Resampling.LANCZOS)
    rgba_image.putalpha(alpha_mask)
    return rgba_image


def rotate_cv2_image(image: Any, angle: int) -> Any:
    normalized_angle = int(angle) % 360
    if normalized_angle == 0:
        return image.copy()

    try:
        import cv2
    except ImportError as exc:
        raise RuntimeError("需要安装 opencv-python 库 (pip install opencv-python)") from exc

    if normalized_angle == 90:
        return cv2.rotate(image, cv2.ROTATE_90_COUNTERCLOCKWISE)
    if normalized_angle == 180:
        return cv2.rotate(image, cv2.ROTATE_180)
    if normalized_angle == 270:
        return cv2.rotate(image, cv2.ROTATE_90_CLOCKWISE)
    raise ValueError(f"不支持的旋转角度：{angle}")


def run_tesseract_on_pil_image(
        pil_image: Image.Image,
        tesseract_bin: str,
        psm: int,
) -> str:
    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_file:
        temp_path = Path(temp_file.name)

    try:
        prepared_image = _prepare_pil_image_for_export(pil_image, force_rgb=True)
        prepared_image.save(temp_path, format="PNG")
        cmd = [
            tesseract_bin,
            str(temp_path),
            "stdout",
            "-l",
            "chi_sim+eng",
            "--psm",
            str(psm),
        ]
        result = subprocess.run(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            encoding="utf-8",
            errors="ignore",
        )
        if result.returncode != 0:
            stderr_text = safe_strip_text(result.stderr)
            raise RuntimeError(f"Tesseract 执行失败：{stderr_text or '未知错误'}")
        return normalize_ocr_text(result.stdout or "")
    finally:
        temp_path.unlink(missing_ok=True)


def score_certificate_ocr_text(text: str) -> float:
    normalized_text = normalize_ocr_text(text)
    if not normalized_text:
        return 0.0

    compact_text = normalized_text.replace(" ", "")
    score = 0.0

    keyword_weights = {
        "姓名": 12.0,
        "性别": 8.0,
        "民族": 8.0,
        "出生": 8.0,
        "住址": 8.0,
        "公民身份号码": 16.0,
        "身份号码": 12.0,
        "中华人民共和国": 16.0,
        "中华人民": 10.0,
        "居民身份证": 16.0,
        "居民身份": 12.0,
        "签发机关": 12.0,
        "有效期限": 12.0,
        "有效期": 10.0,
        "机关": 6.0,
    }
    for keyword, weight in keyword_weights.items():
        if keyword in compact_text:
            score += weight

    if re.search(r"[1-9]\d{5}(?:19|20)\d{2}(?:0\d|1[0-2])(?:0\d|[12]\d|3[01])\d{3}[0-9Xx]", compact_text):
        score += 28.0

    if re.search(r"\d{4}[.\-/年]\d{1,2}[.\-/月]\d{1,2}", compact_text):
        score += 8.0
    if re.search(r"\d{4}[.\-/年]\d{1,2}[.\-/月]\d{1,2}.*\d{4}[.\-/年]\d{1,2}[.\-/月]\d{1,2}", compact_text):
        score += 6.0

    chinese_count = len(re.findall(r"[\u4e00-\u9fff]", normalized_text))
    score += min(15.0, chinese_count * 0.4)

    digit_groups = re.findall(r"\d{6,18}", compact_text)
    if digit_groups:
        score += min(12.0, float(max(len(item) for item in digit_groups)))

    return score


def find_best_certificate_ocr_result(
        image: Any,
        tesseract_bin: str,
) -> Tuple[Any, str, int, int, float]:
    try:
        import cv2
    except ImportError as exc:
        raise RuntimeError("需要安装 opencv-python 库 (pip install opencv-python)") from exc

    best_result: Optional[Tuple[Any, str, int, int, float]] = None

    for angle in CERTIFICATE_ROTATION_ANGLES:
        rotated_image = rotate_cv2_image(image, angle)
        pil_candidate = Image.fromarray(cv2.cvtColor(rotated_image, cv2.COLOR_BGR2RGB))
        for psm in CERTIFICATE_TESSERACT_PSMS:
            recognized_text = run_tesseract_on_pil_image(pil_candidate, tesseract_bin, psm)
            score = score_certificate_ocr_text(recognized_text)
            if best_result is None or score > best_result[4]:
                best_result = (rotated_image, recognized_text, angle, psm, score)

    if best_result is None:
        raise ValueError("OCR 未生成任何识别结果")

    if best_result[4] < 12.0:
        raise ValueError(f"OCR 未识别到足够明确的证件信息，最佳得分 {best_result[4]:.1f}")

    return best_result


def normalize_ocr_text(text: str) -> str:
    lines = []
    for raw_line in text.splitlines():
        line = re.sub(r"\s+", " ", raw_line).strip()
        if line:
            lines.append(line)
    return "\n".join(lines)


def extract_certificate_info_lines(text: str) -> List[str]:
    normalized_text = normalize_ocr_text(text)
    if not normalized_text:
        raise ValueError("OCR 未识别到证件信息")

    summary_lines: List[str] = []
    for line in normalized_text.splitlines():
        compact_line = line.replace(" ", "")
        lower_line = compact_line.lower()
        if any(keyword.lower() in lower_line for keyword in CERTIFICATE_INFO_KEYWORDS):
            summary_lines.append(line)

    id_matches = re.findall(r"[1-9]\d{5}(?:19|20)\d{2}(?:0\d|1[0-2])(?:0\d|[12]\d|3[01])\d{3}[0-9Xx]", normalized_text)
    for value in id_matches:
        summary_lines.append(f"证件号码：{value}")

    date_matches = re.findall(r"\d{4}[.\-/年]\d{1,2}[.\-/月]\d{1,2}", normalized_text)
    for value in date_matches[:2]:
        summary_lines.append(f"证件日期：{value}")

    if not summary_lines:
        raise ValueError("OCR 未提取到证件关键字段")

    unique_lines: List[str] = []
    seen_lines = set()
    for line in summary_lines:
        if line not in seen_lines:
            seen_lines.add(line)
            unique_lines.append(line)
    return unique_lines


def process_certificate_image(
        src_file: Path,
        tesseract_bin: Optional[str],
) -> Tuple[Image.Image, str, List[str]]:
    try:
        import cv2
    except ImportError as exc:
        raise RuntimeError("需要安装 opencv-python 库 (pip install opencv-python)") from exc

    if not tesseract_bin:
        raise RuntimeError("未检测到 tesseract.exe")

    image = load_cv2_image(src_file)
    certificate_quad = detect_certificate_quad(image)
    corrected_image = four_point_transform(image, certificate_quad)
    cropped_image = auto_crop_certificate_image(corrected_image)
    oriented_image, recognized_text, best_angle, best_psm, best_score = find_best_certificate_ocr_result(
        cropped_image,
        tesseract_bin,
    )
    refined_quad = detect_certificate_quad(oriented_image)
    refined_image = four_point_transform(oriented_image, refined_quad)
    final_image = auto_crop_certificate_image(refined_image)
    cutout_image = remove_certificate_background(final_image)

    final_pil_image = Image.fromarray(final_image[:, :, ::-1])
    recognized_text = run_tesseract_on_pil_image(final_pil_image, tesseract_bin, best_psm)
    best_score = score_certificate_ocr_text(recognized_text)

    cleaned_text = normalize_ocr_text(recognized_text)
    if not cleaned_text:
        raise ValueError("OCR 未识别到证件信息")

    info_lines = extract_certificate_info_lines(cleaned_text)
    info_lines.insert(0, f"OCR方向：{best_angle}° | psm={best_psm} | score={best_score:.1f}")
    return cutout_image, cleaned_text, info_lines


def save_images_as_pdf(images: List[Image.Image], dst_file: Path) -> None:
    if not images:
        raise ValueError("没有可写入 PDF 的图片")

    normalized_pages = [_prepare_pil_image_for_export(item, force_rgb=True) for item in images]
    first_page = normalized_pages[0]
    other_pages = normalized_pages[1:]
    first_page.save(str(dst_file), "PDF", resolution=300.0, save_all=True, append_images=other_pages)


def download_video_from_url(
        url: str,
        output_dir: Path,
        yt_dlp_bin: str,
        video_format: str = "mp4",
        progress_callback: Optional[callable] = None,
        log_callback: Optional[callable] = None,
) -> Tuple[bool, str, Optional[Path]]:
    def _emit_progress(percent: float) -> None:
        if progress_callback is None:
            return
        try:
            progress_callback(max(0.0, min(100.0, float(percent))))
        except Exception:
            pass

    def _emit_log(message: str) -> None:
        if log_callback is None or not message:
            return
        try:
            log_callback(message)
        except Exception:
            pass

    def _extract_downloaded_file(stdout_text: str, stderr_text: str) -> Optional[Path]:
        downloaded_file: Optional[Path] = None
        output_lines = [line.strip() for line in stdout_text.splitlines() if line.strip()]

        for line in reversed(output_lines):
            candidate = Path(line)
            if not candidate.is_absolute():
                candidate = output_dir / candidate
            if candidate.exists() and candidate.is_file():
                downloaded_file = candidate
                break

        if not downloaded_file:
            combined_lines = output_lines + [line.strip() for line in stderr_text.splitlines() if line.strip()]
            for line in combined_lines:
                if "Destination:" in line:
                    raw_path = line.split("Destination:", 1)[-1].strip()
                    candidate = Path(raw_path)
                    if not candidate.is_absolute():
                        candidate = output_dir / candidate
                    if candidate.exists() and candidate.is_file():
                        downloaded_file = candidate
                        break

        if not downloaded_file:
            for potential_file in sorted(
                    output_dir.glob(f"*.{video_format}"),
                    key=lambda x: x.stat().st_mtime,
                    reverse=True,
            ):
                if potential_file.is_file() and potential_file.stat().st_size > 0:
                    downloaded_file = potential_file
                    break

        return downloaded_file

    def _run_attempt(cmd: List[str], env: Dict[str, str]) -> Tuple[int, str, str]:
        stdout_lines: List[str] = []
        stderr_lines: List[str] = []
        process = subprocess.Popen(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            encoding="utf-8",
            errors="replace",
            cwd=str(output_dir),
            env=env,
            bufsize=1,
        )

        percentage_pattern = re.compile(r"\[download\]\s+(\d+(?:\.\d+)?)%")
        destination_pattern = re.compile(r"Destination:\s+(.+)$")

        try:
            if process.stdout is not None:
                for raw_line in process.stdout:
                    line = raw_line.rstrip('\r\n')
                    stdout_lines.append(line)
                    stripped = line.strip()
                    if not stripped:
                        continue

                    match = percentage_pattern.search(stripped)
                    if match:
                        _emit_progress(float(match.group(1)))
                        continue

                    if stripped.startswith("[Merger]") or stripped.startswith("[Fixup"):
                        _emit_progress(99.0)

                    destination_match = destination_pattern.search(stripped)
                    if destination_match:
                        candidate_name = Path(destination_match.group(1).strip()).name
                        _emit_log(f"检测到输出文件：{candidate_name}")

            stderr_text, _ = process.communicate()
            if stderr_text:
                stderr_lines.append(stderr_text)
        finally:
            if process.stdout:
                process.stdout.close()
            if process.stderr:
                process.stderr.close()

        return process.returncode, "\n".join(stdout_lines), "\n".join(stderr_lines)

    try:
        output_dir.mkdir(parents=True, exist_ok=True)
        output_template = str(output_dir / "%(title)s.%(ext)s")
        _emit_progress(0.0)

        common_args = [
            "-f", f"bv*+ba/bv*+b/best",
            "--merge-output-format", video_format,
            "-o", output_template,
            "--print", "after_move:filepath",
            "--newline",
            "--no-playlist",
            "--no-warnings",
            "--socket-timeout", "20",
            "--retries", "3",
            "--fragment-retries", "3",
            "--extractor-retries", "3",
            "--user-agent",
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/146.0.0.0 Safari/537.36",
            "--add-header", "Referer: https://www.bilibili.com/",
        ]
        ffmpeg_for_merge = find_executable(FFMPEG_BINARY_CANDIDATES)
        if ffmpeg_for_merge:
            common_args.extend(["--ffmpeg-location", ffmpeg_for_merge])

        is_bilibili = "bilibili.com" in url.lower() or "b23.tv" in url.lower()
        attempt_profiles: List[Dict[str, Any]] = [
            {
                "name": "常规模式",
                "extra_args": [],
                "clean_proxy": False,
            },
            {
                "name": "直连重试",
                "extra_args": ["-4"],
                "clean_proxy": True,
            },
            {
                "name": "证书宽松模式",
                "extra_args": ["-4", "--no-check-certificates"],
                "clean_proxy": True,
            },
        ]

        if is_bilibili:
            attempt_profiles.insert(
                1,
                {
                    "name": "B站兼容模式",
                    "extra_args": ["--extractor-args", "bilibili:try_look=1"],
                    "clean_proxy": False,
                },
            )
            attempt_profiles.append(
                {
                    "name": "B站兼容+证书宽松模式",
                    "extra_args": ["-4", "--extractor-args", "bilibili:try_look=1", "--no-check-certificates"],
                    "clean_proxy": True,
                },
            )

        error_messages: List[str] = []

        for profile in attempt_profiles:
            _emit_progress(0.0)
            _emit_log(f"正在尝试：{profile['name']}")
            cmd = [yt_dlp_bin, *common_args, *profile["extra_args"], url]
            env = os.environ.copy()
            if profile["clean_proxy"]:
                for key in [
                    "http_proxy", "https_proxy", "all_proxy",
                    "HTTP_PROXY", "HTTPS_PROXY", "ALL_PROXY",
                ]:
                    env.pop(key, None)

            return_code, stdout_text, stderr_text = _run_attempt(cmd, env)
            stdout_text = (stdout_text or "").strip()
            stderr_text = (stderr_text or "").strip()
            downloaded_file = _extract_downloaded_file(stdout_text, stderr_text)

            if return_code == 0 and downloaded_file and downloaded_file.exists():
                _emit_progress(100.0)
                return True, f"下载成功: {downloaded_file.name}", downloaded_file

            error_text = stderr_text or stdout_text or "yt-dlp 未返回可读错误信息"
            error_messages.append(f"{profile['name']}: {error_text}")

            if downloaded_file and downloaded_file.exists():
                _emit_progress(100.0)
                return True, f"下载成功: {downloaded_file.name}", downloaded_file

        return False, "下载失败: " + " | ".join(error_messages[-3:]), None

    except Exception as exc:
        return False, f"下载异常: {exc}", None


def _clamp_int(value: int, min_value: int, max_value: int) -> int:
    return max(min_value, min(max_value, int(value)))


def _normalize_box(x: int, y: int, w: int, h: int, frame_width: int, frame_height: int) -> Optional[Dict[str, int]]:
    x = _clamp_int(x, 0, frame_width - 1)
    y = _clamp_int(y, 0, frame_height - 1)
    w = max(1, int(w))
    h = max(1, int(h))
    if x + w > frame_width:
        w = frame_width - x
    if y + h > frame_height:
        h = frame_height - y
    if w <= 0 or h <= 0:
        return None
    return {"x": x, "y": y, "width": w, "height": h}


def _expand_box(box: Dict[str, int], pad_x: int, pad_y: int, frame_width: int, frame_height: int) -> Dict[str, int]:
    return _normalize_box(
        box["x"] - pad_x,
        box["y"] - pad_y,
        box["width"] + pad_x * 2,
        box["height"] + pad_y * 2,
        frame_width,
        frame_height,
    )


def _boxes_close(box1: Dict[str, int], box2: Dict[str, int], gap: int = 24) -> bool:
    left1, top1 = box1["x"], box1["y"]
    right1, bottom1 = left1 + box1["width"], top1 + box1["height"]
    left2, top2 = box2["x"], box2["y"]
    right2, bottom2 = left2 + box2["width"], top2 + box2["height"]
    return not (
            right1 < left2 - gap or
            right2 < left1 - gap or
            bottom1 < top2 - gap or
            bottom2 < top1 - gap
    )


def _merge_boxes(boxes: List[Dict[str, Any]], frame_width: int, frame_height: int, gap: int = 24) -> List[
    Dict[str, Any]]:
    merged: List[Dict[str, Any]] = []
    for src in boxes:
        box = {
            "x": int(src["x"]),
            "y": int(src["y"]),
            "width": int(src["width"]),
            "height": int(src["height"]),
            "source": src.get("source", "unknown"),
            "texts": list(src.get("texts", [])),
            "score": float(src.get("score", 0.0)),
            "hits": int(src.get("hits", 0)),
            "bins": set(src.get("bins", set())),
        }
        if box["width"] <= 0 or box["height"] <= 0:
            continue

        merged_into_existing = False
        for existing in merged:
            if _boxes_close(existing, box, gap=gap):
                min_x = min(existing["x"], box["x"])
                min_y = min(existing["y"], box["y"])
                max_x = max(existing["x"] + existing["width"], box["x"] + box["width"])
                max_y = max(existing["y"] + existing["height"], box["y"] + box["height"])
                existing.update({
                    "x": min_x,
                    "y": min_y,
                    "width": max_x - min_x,
                    "height": max_y - min_y,
                    "score": max(existing.get("score", 0.0), box.get("score", 0.0)),
                    "hits": max(int(existing.get("hits", 0)), int(box.get("hits", 0))),
                })
                existing_bins = existing.setdefault("bins", set())
                existing_bins.update(box.get("bins", set()))
                existing_texts = existing.setdefault("texts", [])
                for item in box.get("texts", []):
                    if item and item not in existing_texts:
                        existing_texts.append(item)
                existing_sources = set(str(existing.get("source", "unknown")).split("+"))
                existing_sources.update(str(box.get("source", "unknown")).split("+"))
                existing["source"] = "+".join(sorted(s for s in existing_sources if s))
                merged_into_existing = True
                break
        if not merged_into_existing:
            merged.append(box)

    normalized: List[Dict[str, Any]] = []
    for item in merged:
        fixed = _normalize_box(item["x"], item["y"], item["width"], item["height"], frame_width, frame_height)
        if fixed:
            fixed.update({
                "source": item.get("source", "unknown"),
                "texts": item.get("texts", []),
                "score": item.get("score", 0.0),
                "hits": int(item.get("hits", 0)),
                "bins": sorted(list(item.get("bins", set()))),
            })
            normalized.append(fixed)
    return normalized


def _build_preferred_area_mask(
        frame_width: int,
        frame_height: int,
        detection_mode: str = "fast",
):
    import numpy as np

    mask = np.zeros((frame_height, frame_width), dtype=np.uint8)
    if detection_mode == "extreme":
        corner_w = max(78, int(frame_width * 0.22))
        corner_h = max(56, int(frame_height * 0.14))
    else:
        corner_w = max(90, int(frame_width * 0.26))
        corner_h = max(70, int(frame_height * 0.20))

    mask[0:corner_h, 0:corner_w] = 255
    mask[0:corner_h, frame_width - corner_w:frame_width] = 255
    mask[frame_height - corner_h:frame_height, 0:corner_w] = 255
    mask[frame_height - corner_h:frame_height, frame_width - corner_w:frame_width] = 255

    # 极速模式只保留顶部和两侧，明确排除底部中间，避免误删字幕/底部字母。
    if detection_mode == "extreme":
        edge_h = max(28, int(frame_height * 0.05))
        edge_w = max(24, int(frame_width * 0.05))
    else:
        # 快速模式也补充上下边带，兼顾 B 站这类靠边但不完全贴角的固定水印。
        edge_h = max(36, int(frame_height * (0.08 if detection_mode == "fast" else 0.10)))
        edge_w = max(30, int(frame_width * (0.06 if detection_mode == "fast" else 0.10)))

    mask[0:edge_h, :] = 255
    mask[:, 0:edge_w] = 255
    mask[:, frame_width - edge_w:frame_width] = 255

    if detection_mode != "extreme":
        mask[frame_height - edge_h:frame_height, :] = 255
    else:
        bottom_safe_h = max(corner_h, int(frame_height * 0.10))
        center_left = int(frame_width * 0.18)
        center_right = int(frame_width * 0.82)
        mask[frame_height - bottom_safe_h:frame_height, center_left:center_right] = 0

    if detection_mode == "comprehensive":
        center_band_h = max(64, int(frame_height * 0.16))
        center_top = max(0, frame_height // 2 - center_band_h // 2)
        center_bottom = min(frame_height, center_top + center_band_h)
        mask[center_top:center_bottom, :] = 255
    return mask


def _bin_index(frame_ratio: float, total_bins: int = 5) -> int:
    frame_ratio = max(0.0, min(0.999999, float(frame_ratio)))
    return min(total_bins - 1, int(frame_ratio * total_bins))


def _boxes_overlap_ratio(box1: Dict[str, int], box2: Dict[str, int]) -> float:
    left = max(box1['x'], box2['x'])
    top = max(box1['y'], box2['y'])
    right = min(box1['x'] + box1['width'], box2['x'] + box2['width'])
    bottom = min(box1['y'] + box1['height'], box2['y'] + box2['height'])
    if right <= left or bottom <= top:
        return 0.0
    inter = (right - left) * (bottom - top)
    area1 = max(1, box1['width'] * box1['height'])
    area2 = max(1, box2['width'] * box2['height'])
    return inter / float(max(1, min(area1, area2)))


def _sample_video_frames(
        video_path: Path,
        detection_mode: str = "fast",
        progress_callback: Optional[callable] = None,
        log_callback: Optional[callable] = None,
):
    try:
        import cv2
    except ImportError:
        return None, None

    cap = cv2.VideoCapture(str(video_path))
    if not cap.isOpened():
        return None, None

    total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
    width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
    height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
    fps = float(cap.get(cv2.CAP_PROP_FPS) or 0.0)
    duration = (total_frames / fps) if fps > 0 else 0.0
    if total_frames <= 0 or width <= 0 or height <= 0:
        cap.release()
        return None, None

    if detection_mode == "extreme":
        sample_frames = 6
    elif detection_mode == "fast":
        sample_frames = 14
    else:
        sample_frames = 24

    # 极速模式只看前几秒，换速度；快速/全面模式仍兼顾全片分布，避免把片头标题当成水印。
    distributed_ratios: List[float] = []
    start_ratio = 0.02
    end_ratio = 0.98
    if duration > 0 and duration < 15:
        start_ratio = 0.0
        end_ratio = 1.0

    if detection_mode == "extreme":
        focus_seconds = min(4.0, max(1.5, duration * 0.35 if duration > 0 else 4.0))
        focus_ratio_end = min(1.0, focus_seconds / max(duration, 0.1)) if duration > 0 else 0.25
        for i in range(sample_frames):
            ratio = focus_ratio_end * (i / max(1, sample_frames - 1))
            distributed_ratios.append(ratio)
    else:
        early_focus_count = min(6 if detection_mode == "fast" else 8, sample_frames)
        for i in range(early_focus_count):
            ratio = min(0.22, 0.02 + i * (0.20 / max(1, early_focus_count - 1)))
            distributed_ratios.append(ratio)

        remaining = max(0, sample_frames - len(distributed_ratios))
        for i in range(remaining):
            ratio = start_ratio + (end_ratio - start_ratio) * (i / max(1, remaining - 1))
            distributed_ratios.append(ratio)

    indices = []
    for ratio in distributed_ratios:
        idx = max(0, min(total_frames - 1, int((total_frames - 1) * ratio)))
        indices.append(idx)
    indices = sorted(set(indices))

    frames = []
    total_indices = max(1, len(indices))
    if log_callback:
        mode_name = "极速" if detection_mode == "extreme" else ("快速" if detection_mode == "fast" else "全面")
        if detection_mode == "extreme":
            log_callback(f"抽样检测：前几秒极速采样 {len(indices)} 帧（模式：{mode_name}）")
        else:
            log_callback(f"抽样检测：全视频分布采样 {len(indices)} 帧（模式：{mode_name}）")

    for pos, idx in enumerate(indices, start=1):
        cap.set(cv2.CAP_PROP_POS_FRAMES, idx)
        ret, frame = cap.read()
        if ret and frame is not None:
            ratio = idx / max(1, total_frames - 1)
            frames.append({
                'frame': frame,
                'index': idx,
                'ratio': ratio,
                'bin': _bin_index(ratio),
            })
        if progress_callback:
            try:
                progress_callback(int(pos * 100 / total_indices))
            except Exception:
                pass

    cap.release()
    if not frames:
        return None, None
    return frames, {"width": width, "height": height, "fps": fps, "duration": duration, "total_frames": total_frames}


def _collect_roi_boxes(frame_width: int, frame_height: int, detection_mode: str) -> List[Tuple[int, int, int, int]]:
    try:
        import cv2
    except ImportError:
        return []

    roi_mask = _build_preferred_area_mask(frame_width, frame_height, detection_mode)
    roi_boxes = []
    contours, _ = cv2.findContours(roi_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    for contour in contours:
        x, y, w, h = cv2.boundingRect(contour)
        roi_boxes.append((x, y, w, h))
    return roi_boxes


def _finalize_region_candidates(
        candidates: List[Dict[str, Any]],
        sample_count: int,
        frame_width: int,
        frame_height: int,
) -> List[Dict[str, Any]]:
    finalized: List[Dict[str, Any]] = []

    for entry in candidates:
        hits = int(entry.get('hits', 0))
        bins = set(entry.get('bins', set()))
        if hits <= 0:
            continue
        if not entry.get('x_values') or not entry.get('y_values') or not entry.get('w_values') or not entry.get('h_values'):
            continue
        x = int(sum(entry['x_values']) / len(entry['x_values']))
        y = int(sum(entry['y_values']) / len(entry['y_values']))
        w = int(sum(entry['w_values']) / len(entry['w_values']))
        h = int(sum(entry['h_values']) / len(entry['h_values']))
        fixed = _normalize_box(x, y, w, h, frame_width, frame_height)
        if not fixed:
            continue
        fixed.update({
            'texts': entry.get('texts', [])[:10],
            'source': entry.get('source', 'unknown'),
            'score': float(entry.get('score', 0.0)) + hits * 8 + len(bins) * 6,
            'hits': hits,
            'bins': bins,
        })
        finalized.append(_expand_box(fixed, 8, 6, frame_width, frame_height))
    return finalized


def _detect_static_overlay_regions(
        samples: List[Dict[str, Any]],
        frame_width: int,
        frame_height: int,
        detection_mode: str = "fast",
) -> List[Dict[str, Any]]:
    try:
        import cv2
        import numpy as np
    except ImportError:
        return []

    gray_stack = np.stack([cv2.cvtColor(item['frame'], cv2.COLOR_BGR2GRAY) for item in samples]).astype(np.float32)
    mean_gray = np.mean(gray_stack, axis=0).astype(np.uint8)
    std_gray = np.std(gray_stack, axis=0).astype(np.uint8)

    mask_static = cv2.threshold(std_gray, 16 if detection_mode == 'fast' else 18, 255, cv2.THRESH_BINARY_INV)[1]
    mask_bright = cv2.threshold(mean_gray, 160, 255, cv2.THRESH_BINARY)[1]
    mask_dark = cv2.threshold(mean_gray, 82, 255, cv2.THRESH_BINARY_INV)[1]
    edges = cv2.Canny(mean_gray, 80, 180)
    edge_regions = cv2.dilate(edges, np.ones((3, 3), np.uint8), iterations=1)

    candidate = cv2.bitwise_or(cv2.bitwise_and(mask_static, mask_bright), cv2.bitwise_and(mask_static, mask_dark))
    candidate = cv2.bitwise_or(candidate, cv2.bitwise_and(mask_static, edge_regions))
    candidate = cv2.bitwise_and(candidate, _build_preferred_area_mask(frame_width, frame_height, detection_mode))
    candidate = cv2.morphologyEx(candidate, cv2.MORPH_CLOSE, np.ones((5, 5), np.uint8), iterations=2)
    candidate = cv2.morphologyEx(candidate, cv2.MORPH_OPEN, np.ones((3, 3), np.uint8), iterations=1)

    contours, _ = cv2.findContours(candidate, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    min_area = max(120, int(frame_width * frame_height * 0.00005))
    max_area = int(frame_width * frame_height * 0.12)

    raw_candidates: List[Dict[str, Any]] = []
    for contour in contours:
        x, y, w, h = cv2.boundingRect(contour)
        area = w * h
        if area < min_area or area > max_area:
            continue
        if w < 18 or h < 12:
            continue
        if h > frame_height * 0.35 or w > frame_width * 0.58:
            continue
        aspect = w / max(h, 1)
        if aspect < 0.3 or aspect > 22:
            continue

        hits = 0
        bins = set()
        for sample in samples:
            gray = cv2.cvtColor(sample['frame'], cv2.COLOR_BGR2GRAY)
            roi = gray[y:y + h, x:x + w]
            if roi.size == 0:
                continue
            roi_std = float(np.std(roi))
            roi_mean = float(np.mean(roi))
            edge_density = float(np.mean(cv2.Canny(roi, 80, 180) > 0))
            if roi_std > 6 and (roi_mean > 145 or roi_mean < 110 or edge_density > 0.03):
                hits += 1
                bins.add(sample['bin'])

        raw_candidates.append({
            'x_values': [x],
            'y_values': [y],
            'w_values': [w],
            'h_values': [h],
            'texts': [],
            'hits': hits,
            'bins': bins,
            'score': float(area),
            'source': 'static',
        })

    raw_candidates.sort(key=lambda item: (len(item.get('bins', set())), item.get('hits', 0), item.get('score', 0.0)), reverse=True)
    return _finalize_region_candidates(raw_candidates[:16], len(samples), frame_width, frame_height)


def _detect_stable_edge_regions(
        samples: List[Dict[str, Any]],
        frame_width: int,
        frame_height: int,
        detection_mode: str = "fast",
) -> List[Dict[str, Any]]:
    try:
        import cv2
        import numpy as np
    except ImportError:
        return []

    if not samples:
        return []

    edge_maps = []
    for item in samples:
        gray = cv2.cvtColor(item['frame'], cv2.COLOR_BGR2GRAY)
        blur = cv2.GaussianBlur(gray, (3, 3), 0)
        edges = cv2.Canny(blur, 60, 160)
        edge_maps.append((edges > 0).astype(np.uint8))

    stable_ratio = np.mean(np.stack(edge_maps, axis=0), axis=0)
    stable_mask = (stable_ratio >= (0.58 if detection_mode == 'fast' else 0.52)).astype(np.uint8) * 255
    stable_mask = cv2.bitwise_and(stable_mask, _build_preferred_area_mask(frame_width, frame_height, detection_mode))
    stable_mask = cv2.morphologyEx(stable_mask, cv2.MORPH_CLOSE, np.ones((5, 5), np.uint8), iterations=2)
    stable_mask = cv2.dilate(stable_mask, np.ones((3, 3), np.uint8), iterations=1)

    contours, _ = cv2.findContours(stable_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    min_area = max(120, int(frame_width * frame_height * 0.00005))
    max_area = int(frame_width * frame_height * 0.14)

    raw_candidates: List[Dict[str, Any]] = []
    for contour in contours:
        x, y, w, h = cv2.boundingRect(contour)
        area = w * h
        if area < min_area or area > max_area:
            continue
        if w < 22 or h < 12:
            continue
        if h > frame_height * 0.30 or w > frame_width * 0.62:
            continue

        hits = 0
        bins = set()
        for sample in samples:
            gray = cv2.cvtColor(sample['frame'], cv2.COLOR_BGR2GRAY)
            roi = gray[y:y + h, x:x + w]
            if roi.size == 0:
                continue
            local_edges = cv2.Canny(roi, 60, 160)
            edge_density = float(np.mean(local_edges > 0))
            if edge_density >= 0.035:
                hits += 1
                bins.add(sample['bin'])

        raw_candidates.append({
            'x_values': [x],
            'y_values': [y],
            'w_values': [w],
            'h_values': [h],
            'texts': [],
            'hits': hits,
            'bins': bins,
            'score': float(area) + len(bins) * 30,
            'source': 'stable_edge',
        })

    raw_candidates.sort(key=lambda item: (len(item.get('bins', set())), item.get('hits', 0), item.get('score', 0.0)), reverse=True)
    return _finalize_region_candidates(raw_candidates[:16], len(samples), frame_width, frame_height)


def _normalize_detected_text(text: str) -> str:
    text = re.sub(r'\s+', '', str(text or ''))
    text = re.sub(r'[^\w一-鿿@.-]+', '', text)
    return text.strip().lower()


def _detect_persistent_ocr_regions(
        samples: List[Dict[str, Any]],
        frame_width: int,
        frame_height: int,
        detection_mode: str = "fast",
        static_boxes: Optional[List[Dict[str, Any]]] = None,
) -> List[Dict[str, Any]]:
    try:
        import cv2
        import pytesseract
    except ImportError:
        return []

    candidate_texts: Dict[Tuple[int, int, int, int, str], Dict[str, Any]] = {}
    roi_boxes: List[Tuple[int, int, int, int]] = []

    if static_boxes:
        for item in static_boxes[: (3 if detection_mode == 'fast' else 6)]:
            roi_boxes.append((item['x'], item['y'], item['width'], item['height']))

    if not roi_boxes and detection_mode == 'comprehensive':
        roi_boxes = _collect_roi_boxes(frame_width, frame_height, detection_mode)

    for sample in samples:
        gray = cv2.cvtColor(sample['frame'], cv2.COLOR_BGR2GRAY)
        for rx, ry, rw, rh in roi_boxes:
            roi = gray[ry:ry + rh, rx:rx + rw]
            if roi.size == 0:
                continue
            enlarged = cv2.resize(roi, None, fx=1.8, fy=1.8, interpolation=cv2.INTER_CUBIC)
            data = pytesseract.image_to_data(
                enlarged,
                output_type=pytesseract.Output.DICT,
                lang='chi_sim+eng',
                config='--oem 3 --psm 6'
            )
            count = len(data.get('text', []))
            for i in range(count):
                raw_text = str(data['text'][i] or '').strip()
                text = _normalize_detected_text(raw_text)
                conf_raw = str(data.get('conf', ['-1'])[i])
                try:
                    conf = float(conf_raw)
                except Exception:
                    conf = -1.0
                if len(text) < 2 or conf < 28:
                    continue
                x = int(data['left'][i] / 1.8) + rx
                y = int(data['top'][i] / 1.8) + ry
                w = max(1, int(data['width'][i] / 1.8))
                h = max(1, int(data['height'][i] / 1.8))
                key = (round(x / 24), round(y / 24), round(w / 24), round(h / 24), text)
                entry = candidate_texts.setdefault(key, {
                    'x_values': [], 'y_values': [], 'w_values': [], 'h_values': [],
                    'texts': [], 'hits': 0, 'score': 0.0, 'bins': set(), 'source': 'ocr',
                })
                entry['x_values'].append(x)
                entry['y_values'].append(y)
                entry['w_values'].append(w)
                entry['h_values'].append(h)
                if raw_text and raw_text not in entry['texts']:
                    entry['texts'].append(raw_text)
                entry['hits'] += 1
                entry['bins'].add(sample['bin'])
                entry['score'] = max(entry['score'], conf)

    return _finalize_region_candidates(list(candidate_texts.values()), len(samples), frame_width, frame_height)


def detect_text_watermark_region(
        video_path: Path,
        ffmpeg_bin: str,
        detection_mode: str = "fast",
        progress_callback: Optional[callable] = None,
        log_callback: Optional[callable] = None,
) -> Optional[Dict[str, Any]]:
    try:
        def _emit_progress(value: int) -> None:
            if progress_callback is None:
                return
            try:
                progress_callback(max(0, min(100, int(value))))
            except Exception:
                pass

        _emit_progress(0)
        samples, info = _sample_video_frames(
            video_path,
            detection_mode=detection_mode,
            progress_callback=lambda p: _emit_progress(int(p * 0.35)),
            log_callback=log_callback,
        )
        if not samples or not info:
            return None

        frame_width = int(info['width'])
        frame_height = int(info['height'])
        duration = float(info.get('duration', 0.0))

        if log_callback:
            log_callback("正在识别固定水印区域…")
        static_boxes = _detect_static_overlay_regions(samples, frame_width, frame_height, detection_mode=detection_mode)
        stable_edge_boxes = _detect_stable_edge_regions(samples, frame_width, frame_height, detection_mode=detection_mode)
        _emit_progress(65)

        ocr_boxes: List[Dict[str, Any]] = []
        preferred_seed_boxes = list(static_boxes) + list(stable_edge_boxes)
        if detection_mode == 'extreme':
            ocr_boxes = []
        elif detection_mode == 'comprehensive' or not preferred_seed_boxes:
            if log_callback:
                log_callback("正在进行 OCR 复核…")
            ocr_boxes = _detect_persistent_ocr_regions(
                samples,
                frame_width,
                frame_height,
                detection_mode=detection_mode,
                static_boxes=preferred_seed_boxes,
            )
        elif preferred_seed_boxes:
            ocr_boxes = _detect_persistent_ocr_regions(
                samples,
                frame_width,
                frame_height,
                detection_mode='fast',
                static_boxes=preferred_seed_boxes[:4],
            )
        _emit_progress(88)

        candidate_boxes: List[Dict[str, Any]] = []
        candidate_boxes.extend(static_boxes)
        candidate_boxes.extend(stable_edge_boxes)
        candidate_boxes.extend(ocr_boxes)

        merged_boxes = _merge_boxes(candidate_boxes, frame_width, frame_height, gap=26)
        if not merged_boxes:
            return None

        merged_boxes = sorted(
            merged_boxes,
            key=lambda item: (
                int(item.get('hits', 0)),
                float(item.get('score', 0.0)),
                item['width'] * item['height'],
            ),
            reverse=True,
        )[:6]

        primary = merged_boxes[0].copy()
        primary['regions'] = merged_boxes
        primary['sample_count'] = len(samples)
        primary['duration'] = duration
        primary['mode'] = detection_mode
        primary['frame_width'] = frame_width
        primary['frame_height'] = frame_height
        text_values: List[str] = []
        for item in merged_boxes:
            for txt in item.get('texts', []):
                if txt and txt not in text_values:
                    text_values.append(txt)
        primary['texts'] = text_values[:12]
        _emit_progress(100)
        return primary

    except Exception:
        return None


def _parse_ffmpeg_time_to_seconds(time_text: str) -> float:
    try:
        hours, minutes, seconds = time_text.split(':')
        return int(hours) * 3600 + int(minutes) * 60 + float(seconds)
    except Exception:
        return 0.0


def _sanitize_delogo_region(
        x: int,
        y: int,
        w: int,
        h: int,
        frame_width: int,
        frame_height: int,
        band: int = 4,
) -> Optional[Dict[str, int]]:
    if frame_width <= 2 or frame_height <= 2:
        return None

    band = max(0, int(band))
    usable_left = band
    usable_top = band
    usable_right = frame_width - band - 1
    usable_bottom = frame_height - band - 1

    if usable_right <= usable_left or usable_bottom <= usable_top:
        return None

    x = _clamp_int(x, usable_left, usable_right)
    y = _clamp_int(y, usable_top, usable_bottom)
    max_w = max(1, usable_right - x)
    max_h = max(1, usable_bottom - y)
    w = max(1, min(int(w), max_w))
    h = max(1, min(int(h), max_h))

    if w < 4 or h < 4:
        return None

    return {"x": x, "y": y, "width": w, "height": h, "band": band}


def _run_ffmpeg_with_progress(
        ffmpeg_bin: str,
        src_file: Path,
        dst_file: Path,
        filter_complex: str,
        duration_seconds: float,
        progress_callback: Optional[callable] = None,
        log_callback: Optional[callable] = None,
) -> Tuple[bool, str]:
    cmd = [
        ffmpeg_bin,
        '-y',
        '-i', str(src_file),
        '-map', '0:v:0',
        '-map', '0:a?',
        '-vf', filter_complex,
        '-c:a', 'aac',
        '-b:a', '192k',
        '-movflags', '+faststart',
    ]

    encoder_name = 'libx264'
    if ffmpeg_supports_encoder(ffmpeg_bin, 'h264_nvenc'):
        encoder_name = 'h264_nvenc'
        cmd.extend([
            '-c:v', 'h264_nvenc',
            '-preset', 'p4',
            '-rc', 'vbr',
            '-cq', '21',
            '-b:v', '0',
        ])
    elif ffmpeg_supports_encoder(ffmpeg_bin, 'hevc_nvenc'):
        encoder_name = 'hevc_nvenc'
        cmd.extend([
            '-c:v', 'hevc_nvenc',
            '-preset', 'p4',
            '-rc', 'vbr',
            '-cq', '22',
            '-b:v', '0',
        ])
    else:
        cmd.extend([
            '-c:v', 'libx264',
            '-preset', 'veryfast',
            '-crf', '18',
        ])

    cmd.append(str(dst_file))

    if log_callback is not None:
        try:
            log_callback(f"去水印编码器：{encoder_name}")
        except Exception:
            pass

    process = subprocess.Popen(
        cmd,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        encoding='utf-8',
        errors='replace',
        bufsize=1,
    )

    stderr_lines: List[str] = []
    time_pattern = re.compile(r'time=(\d+:\d+:\d+(?:\.\d+)?)')

    try:
        if process.stderr is not None:
            for raw_line in process.stderr:
                line = raw_line.rstrip('\r\n')
                stderr_lines.append(line)
                match = time_pattern.search(line)
                if match and duration_seconds > 0 and progress_callback is not None:
                    current_seconds = _parse_ffmpeg_time_to_seconds(match.group(1))
                    percent = int(max(0.0, min(100.0, current_seconds * 100.0 / max(duration_seconds, 0.1))))
                    try:
                        progress_callback(percent)
                    except Exception:
                        pass
        return_code = process.wait()
    finally:
        if process.stdout:
            process.stdout.close()
        if process.stderr:
            process.stderr.close()

    if return_code == 0 and dst_file.exists() and dst_file.stat().st_size > 0:
        if progress_callback is not None:
            try:
                progress_callback(100)
            except Exception:
                pass
        return True, ''

    stderr_text = '\n'.join(stderr_lines).strip()
    return False, stderr_text or 'ffmpeg 未返回可读错误信息'


def probe_video_frame_size(ffmpeg_bin: str, src_file: Path) -> Tuple[int, int]:
    info = probe_media_info(ffmpeg_bin, src_file)
    streams = info.get("streams", []) if isinstance(info, dict) else []
    video_stream = next((item for item in streams if item.get("codec_type") == "video"), {})
    try:
        width = int(video_stream.get("width") or 0)
    except Exception:
        width = 0
    try:
        height = int(video_stream.get("height") or 0)
    except Exception:
        height = 0
    return width, height


def remove_video_watermark(
        src_file: Path,
        dst_file: Path,
        ffmpeg_bin: str,
        watermark_region: Optional[Dict[str, Any]] = None,
        detection_mode: str = "fast",
        detect_progress_callback: Optional[callable] = None,
        encode_progress_callback: Optional[callable] = None,
        log_callback: Optional[callable] = None,
) -> Tuple[bool, str]:
    try:
        if watermark_region is None:
            watermark_region = detect_text_watermark_region(
                src_file,
                ffmpeg_bin,
                detection_mode=detection_mode,
                progress_callback=detect_progress_callback,
                log_callback=log_callback,
            )

        if watermark_region is None:
            shutil.copy2(src_file, dst_file)
            return True, f"成功：{src_file.name} -> {dst_file.name}（未检测到可用水印区域，直接复制）"

        regions = list(watermark_region.get('regions', [])) or [watermark_region]
        if not regions:
            shutil.copy2(src_file, dst_file)
            return True, f"成功：{src_file.name} -> {dst_file.name}（未检测到可用水印区域，直接复制）"

        frame_width = int(watermark_region.get('frame_width', 0) or 0)
        frame_height = int(watermark_region.get('frame_height', 0) or 0)
        probed_width, probed_height = probe_video_frame_size(ffmpeg_bin, src_file)
        if probed_width > 0 and probed_height > 0:
            if (frame_width, frame_height) != (probed_width, probed_height) and log_callback:
                log_callback(
                    f"检测尺寸与实际视频尺寸不一致，使用实际尺寸："
                    f"{frame_width}x{frame_height} -> {probed_width}x{probed_height}"
                )
            frame_width, frame_height = probed_width, probed_height
        if frame_width <= 0 or frame_height <= 0:
            try:
                import cv2
                cap = cv2.VideoCapture(str(src_file))
                if cap.isOpened():
                    frame_width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH) or 0)
                    frame_height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT) or 0)
                cap.release()
            except Exception:
                frame_width = frame_width or 0
                frame_height = frame_height or 0

        filter_parts = []
        region_desc = []
        delogo_band = 4
        for item in regions[:6]:
            raw_x = int(item.get('x', 0))
            raw_y = int(item.get('y', 0))
            raw_w = int(item.get('width', 1))
            raw_h = int(item.get('height', 1))
            fixed = _sanitize_delogo_region(raw_x, raw_y, raw_w, raw_h, frame_width, frame_height, band=delogo_band)
            if not fixed:
                continue
            x = fixed['x']
            y = fixed['y']
            w = fixed['width']
            h = fixed['height']
            band = int(fixed.get('band', delogo_band))
            filter_parts.append(f"delogo=x={x}:y={y}:w={w}:h={h}:show=0")
            src_name = item.get('source', 'unknown')
            hits = int(item.get('hits', 0))
            bins = len(set(item.get('bins', [])))
            region_desc.append(f"{src_name}@({x},{y},{w},{h})/band={band}/hits={hits}/bins={bins}")

        if not filter_parts:
            shutil.copy2(src_file, dst_file)
            return True, f"成功：{src_file.name} -> {dst_file.name}（未生成有效 delogo 区域，直接复制）"

        filter_complex = ','.join(filter_parts)
        duration_seconds = float(watermark_region.get('duration', 0.0))
        if log_callback:
            log_callback("正在编码输出去水印视频…")

        ok, ffmpeg_message = _run_ffmpeg_with_progress(
            ffmpeg_bin=ffmpeg_bin,
            src_file=src_file,
            dst_file=dst_file,
            filter_complex=filter_complex,
            duration_seconds=duration_seconds,
            progress_callback=encode_progress_callback,
            log_callback=log_callback,
        )

        if ok:
            texts = watermark_region.get('texts', [])
            watermark_text = '、'.join(texts[:6]) if texts else '未识别到稳定文本，按固定图像区域处理'
            return True, (
                f"成功：{src_file.name} -> {dst_file.name}（已处理 {len(filter_parts)} 个区域；"
                f"检测结果：{watermark_text}；区域：{' | '.join(region_desc[:4])}）"
            )

        return False, f"失败：{src_file.name} | {ffmpeg_message}"

    except Exception as exc:
        return False, f"失败：{src_file.name} | {exc}"


class VideoDownloadWorker(QThread):
    log_signal = Signal(str)
    progress_signal = Signal(int, int)
    stage_signal = Signal(str)
    finished_signal = Signal(int, int, int)
    video_downloaded = Signal(str)

    def __init__(
            self,
            video_urls: List[str],
            local_video_files: List[str],
            output_dir: str,
            yt_dlp_bin: Optional[str],
            ffmpeg_bin: Optional[str],
            remove_watermark: bool,
            video_format: str,
            detection_mode: str,
    ):
        super().__init__()
        self.video_urls = video_urls
        self.local_video_files = [str(Path(item).expanduser().resolve()) for item in local_video_files]
        self.output_dir = Path(output_dir).expanduser().resolve()
        self.yt_dlp_bin = yt_dlp_bin
        self.ffmpeg_bin = ffmpeg_bin
        self.remove_watermark = remove_watermark
        self.video_format = video_format
        self.detection_mode = detection_mode
        self._is_running = True

    def stop(self) -> None:
        self._is_running = False

    def _emit_overall_progress(self, item_index: int, total_items: int, phase_start: int, phase_end: int, phase_percent: float) -> None:
        phase_percent = max(0.0, min(100.0, float(phase_percent)))
        per_item = 1000
        start_unit = int(per_item * phase_start / 100)
        end_unit = int(per_item * phase_end / 100)
        current_item_units = int(start_unit + (end_unit - start_unit) * phase_percent / 100.0)
        current_units = item_index * per_item + current_item_units
        total_units = max(1, total_items * per_item)
        self.progress_signal.emit(current_units, total_units)

    def _build_non_conflict_output_path(self, src_path: Path) -> Path:
        if self.remove_watermark:
            candidate = self.output_dir / f"{src_path.stem}_no_watermark{src_path.suffix}"
        else:
            candidate = self.output_dir / src_path.name
        if candidate.resolve() != src_path.resolve() and not candidate.exists():
            return candidate

        index = 1
        while True:
            suffix_name = f"{src_path.stem}_no_watermark_{index}{src_path.suffix}" if self.remove_watermark else f"{src_path.stem}_{index}{src_path.suffix}"
            candidate = self.output_dir / suffix_name
            if not candidate.exists():
                return candidate
            index += 1

    def _process_final_video(self, src_video: Path, final_output: Path, item_index: int, total_items: int) -> Tuple[bool, str]:
        if not self.remove_watermark or not self.ffmpeg_bin:
            shutil.move(str(src_video), str(final_output))
            return True, f"成功：{final_output.name}"

        self.stage_signal.emit("抽样检测中")
        ok, message = remove_video_watermark(
            src_file=src_video,
            dst_file=final_output,
            ffmpeg_bin=self.ffmpeg_bin,
            detection_mode=self.detection_mode,
            detect_progress_callback=lambda p: self._emit_overall_progress(item_index, total_items, 55, 75, p),
            encode_progress_callback=lambda p: self._emit_overall_progress(item_index, total_items, 75, 100, p),
            log_callback=self.log_signal.emit,
        )
        return ok, message

    def run(self) -> None:
        success_count = 0
        fail_count = 0
        skip_count = 0

        try:
            self.output_dir.mkdir(parents=True, exist_ok=True)
            temp_dir = self.output_dir / "temp_downloads"
            temp_dir.mkdir(parents=True, exist_ok=True)

            tasks: List[Dict[str, Any]] = []
            for url in self.video_urls:
                tasks.append({"type": "url", "value": url})
            for file_path in self.local_video_files:
                tasks.append({"type": "local", "value": file_path})

            total = len(tasks)
            if total == 0:
                self.log_signal.emit("没有可处理的视频任务。")
                self.finished_signal.emit(0, 0, 0)
                return

            for current, task in enumerate(tasks):
                if not self._is_running:
                    self.log_signal.emit("任务已停止。")
                    break

                task_type = task['type']
                task_value = task['value']
                self._emit_overall_progress(current, total, 0, 100, 0)

                if task_type == 'url':
                    url = str(task_value)
                    self.stage_signal.emit("下载中")
                    self.log_signal.emit(f"正在下载: {url}")
                    ok, message, downloaded_file = download_video_from_url(
                        url=url,
                        output_dir=temp_dir,
                        yt_dlp_bin=self.yt_dlp_bin,
                        video_format=self.video_format,
                        progress_callback=lambda p, item_index=current, total_count=total: self._emit_overall_progress(item_index, total_count, 0, 55, p),
                        log_callback=self.log_signal.emit,
                    )

                    if not ok or not downloaded_file:
                        fail_count += 1
                        self.log_signal.emit(f"下载失败: {message}")
                        self._emit_overall_progress(current, total, 0, 100, 100)
                        continue

                    self.log_signal.emit(f"下载完成: {downloaded_file.name}")
                    final_file = self.output_dir / downloaded_file.name
                    ok, message = self._process_final_video(downloaded_file, final_file, current, total)
                    if ok:
                        success_count += 1
                        self.log_signal.emit(message)
                        self.video_downloaded.emit(str(final_file))
                    else:
                        self.log_signal.emit(f"去水印失败，已自动回退保留原文件: {message}")
                        shutil.move(str(downloaded_file), str(final_file))
                        skip_count += 1
                        self.video_downloaded.emit(str(final_file))
                    try:
                        if downloaded_file.exists():
                            downloaded_file.unlink()
                    except Exception:
                        pass

                else:
                    src_file = Path(task_value)
                    if not src_file.exists() or not src_file.is_file():
                        fail_count += 1
                        self.log_signal.emit(f"本地文件不存在，已跳过：{src_file}")
                        self._emit_overall_progress(current, total, 0, 100, 100)
                        continue

                    self.stage_signal.emit("准备本地视频")
                    self.log_signal.emit(f"正在处理本地视频: {src_file.name}")
                    final_file = self._build_non_conflict_output_path(src_file)
                    if self.remove_watermark and self.ffmpeg_bin:
                        self._emit_overall_progress(current, total, 0, 15, 100)
                        ok, message = remove_video_watermark(
                            src_file=src_file,
                            dst_file=final_file,
                            ffmpeg_bin=self.ffmpeg_bin,
                            detection_mode=self.detection_mode,
                            detect_progress_callback=lambda p, item_index=current, total_count=total: self._emit_overall_progress(item_index, total_count, 15, 45, p),
                            encode_progress_callback=lambda p, item_index=current, total_count=total: self._emit_overall_progress(item_index, total_count, 45, 100, p),
                            log_callback=self.log_signal.emit,
                        )
                        if ok:
                            success_count += 1
                            self.log_signal.emit(message)
                            self.video_downloaded.emit(str(final_file))
                        else:
                            shutil.copy2(src_file, final_file)
                            skip_count += 1
                            self.log_signal.emit(f"去水印失败，已自动回退保留原文件: {message}")
                            self.video_downloaded.emit(str(final_file))
                    else:
                        self.stage_signal.emit("复制本地视频")
                        shutil.copy2(src_file, final_file)
                        success_count += 1
                        self.log_signal.emit(f"成功：{src_file.name} -> {final_file.name}")
                        self.video_downloaded.emit(str(final_file))

                self._emit_overall_progress(current, total, 0, 100, 100)

            self.stage_signal.emit("已完成")
            try:
                temp_dir.rmdir()
            except Exception:
                pass

            self.finished_signal.emit(success_count, fail_count, skip_count)
        except Exception as exc:
            self.log_signal.emit(f"程序异常：{exc}")
            self.finished_signal.emit(success_count, fail_count, skip_count)


class ConvertWorker(QThread):
    log_signal = Signal(str)
    progress_signal = Signal(int, int)
    finished_signal = Signal(int, int, int)

    def __init__(
            self,
            selected_files: List[str],
            output_dir: str,
            target_ext: str,
            overwrite: bool,
            ffmpeg_bin: Optional[str],
            libreoffice_bin: Optional[str],
            tesseract_bin: Optional[str],
            mode: str = "convert",
            compress_quality: int = 85,
    ):
        super().__init__()
        self.selected_files = [Path(item).expanduser().resolve() for item in selected_files]
        self.output_dir = Path(output_dir).expanduser().resolve()
        self.target_ext = target_ext
        self.overwrite = overwrite
        self.ffmpeg_bin = ffmpeg_bin
        self.libreoffice_bin = libreoffice_bin
        self.tesseract_bin = tesseract_bin
        self.mode = mode
        self.compress_quality = compress_quality
        self._is_running = True

    def stop(self) -> None:
        self._is_running = False

    def run(self) -> None:
        success_count = 0
        fail_count = 0
        skip_count = 0

        try:
            files = [file_path for file_path in self.selected_files if file_path.is_file()]

            if not files:
                self.log_signal.emit("没有找到需要处理的文件。")
                self.finished_signal.emit(0, 0, 0)
                return

            self.output_dir.mkdir(parents=True, exist_ok=True)

            total = len(files)
            current = 0

            for src_file in files:
                if not self._is_running:
                    self.log_signal.emit("任务已停止。")
                    break

                self.log_signal.emit(f"开始处理：{src_file.name}")

                if self.mode == "compress":
                    dst_file = build_distinct_output_path(
                        src_file=src_file,
                        dst_root=self.output_dir,
                        target_ext=src_file.suffix.lower().replace(".", ""),
                        mode="compress",
                    )
                elif self.mode == "scan":
                    dst_file = build_distinct_output_path(
                        src_file=src_file,
                        dst_root=self.output_dir,
                        target_ext=self.target_ext,
                        mode="scan",
                    )
                elif self.mode == "watermark":
                    dst_file = build_distinct_output_path(
                        src_file=src_file,
                        dst_root=self.output_dir,
                        target_ext="pdf",
                        mode="watermark",
                    )
                else:
                    dst_file = build_distinct_output_path(
                        src_file=src_file,
                        dst_root=self.output_dir,
                        target_ext=self.target_ext,
                        mode="convert",
                    )

                if dst_file.exists() and not self.overwrite:
                    message = f"跳过：{dst_file.name} | 目标文件已存在"
                    skip_count += 1
                    self.log_signal.emit(message)
                else:
                    ok = False
                    message = ""

                    if self.mode == "compress":
                        category = detect_category_by_file(src_file)
                        if category == "image":
                            self.progress_signal.emit(current * 100, total * 100)
                            ok, message = compress_image_lossless(src_file, dst_file, self.compress_quality)
                            self.progress_signal.emit((current + 1) * 100, total * 100)
                        elif category == "video":
                            if not self.ffmpeg_bin:
                                ok, message = False, f"失败：{src_file.name} | 未检测到 ffmpeg"
                            else:
                                video_plan = build_video_compress_plan(self.ffmpeg_bin, src_file, self.compress_quality)
                                ok, message = compress_video_lossless(
                                    src_file,
                                    dst_file,
                                    self.ffmpeg_bin,
                                    video_plan["crf"],
                                    video_plan["preset"],
                                    target_bitrate=video_plan["target_bitrate"],
                                    use_bitrate_mode=video_plan["use_bitrate_mode"],
                                    source_codec=video_plan["codec_name"],
                                    encoder=video_plan["encoder"],
                                    progress_callback=lambda p, item_index=current, total_count=total: self.progress_signal.emit(
                                        item_index * 100 + int(max(0, min(100, p))),
                                        total_count * 100,
                                    ),
                                    log_callback=self.log_signal.emit,
                                )
                        else:
                            ok, message = False, f"失败：{src_file.name} | 仅支持图片和视频压缩"

                    elif self.mode == "scan":
                        ok, message = scan_image_to_document(
                            src_file, dst_file, self.tesseract_bin, self.libreoffice_bin
                        )

                    elif self.mode == "watermark":
                        ok, message = remove_pdf_watermark(src_file, dst_file)

                    else:
                        category = detect_category_by_file(src_file)

                        if not category:
                            message = f"跳过：{src_file.name} | 无法识别文件类别"
                            skip_count += 1
                            self.log_signal.emit(message)
                            current += 1
                            self.progress_signal.emit(current, total)
                            continue

                        if src_file.suffix.lower() == f".{self.target_ext}":
                            message = f"跳过：{src_file.name} | 源文件已是目标格式"
                            skip_count += 1
                            self.log_signal.emit(message)
                            current += 1
                            self.progress_signal.emit(current, total)
                            continue

                        supported_targets = get_supported_targets_for_file(src_file)
                        if self.target_ext not in supported_targets:
                            ok = False
                            message = (
                                f"失败：{src_file.name} | 自动识别类别为"
                                f"{CATEGORY_NAME_MAP.get(category, category)}，不支持转换为 {self.target_ext}"
                            )
                        else:
                            if category == "image":
                                if self.target_ext == "pdf":
                                    ok, message = convert_image_to_pdf(src_file, dst_file)
                                elif self.target_ext == "docx":
                                    ok, message = convert_image_to_docx(src_file, dst_file, self.libreoffice_bin)
                                else:
                                    ok, message = convert_image(src_file, dst_file)
                            elif category in {"audio", "video"}:
                                if not self.ffmpeg_bin:
                                    ok, message = False, f"失败：{src_file.name} | 未检测到 ffmpeg"
                                else:
                                    ok, message = convert_media(src_file, dst_file, self.ffmpeg_bin)
                            elif category == "document":
                                ok, message = convert_document(src_file, dst_file, self.libreoffice_bin)
                            else:
                                ok, message = False, f"失败：{src_file.name} | 不支持的类别"

                    if ok:
                        success_count += 1
                    else:
                        fail_count += 1

                    self.log_signal.emit(message)

                current += 1
                self.progress_signal.emit(current, total)

            self.finished_signal.emit(success_count, fail_count, skip_count)
        except Exception as exc:
            self.log_signal.emit(f"程序异常：{exc}")
            self.finished_signal.emit(success_count, fail_count, skip_count)


class CertificatePdfWorker(QThread):
    log_signal = Signal(str)
    progress_signal = Signal(int, int)
    summary_signal = Signal(str)
    finished_signal = Signal(int, int, int, str, str)

    def __init__(
            self,
            selected_files: List[str],
            output_dir: str,
            output_name: str,
            overwrite: bool,
            tesseract_bin: Optional[str],
    ):
        super().__init__()
        self.selected_files = [Path(item).expanduser().resolve() for item in selected_files]
        self.output_dir = Path(output_dir).expanduser().resolve()
        self.output_name = output_name.strip()
        self.overwrite = overwrite
        self.tesseract_bin = tesseract_bin
        self._is_running = True

    def stop(self) -> None:
        self._is_running = False

    def run(self) -> None:
        success_count = 0
        fail_count = 0
        skip_count = 0
        pdf_result_path = ""
        text_result_path = ""

        try:
            files = [file_path for file_path in self.selected_files if file_path.is_file()]
            if not files:
                self.log_signal.emit("没有找到需要处理的证件图片。")
                self.finished_signal.emit(0, 0, 0, "", "")
                return

            if not self.output_name:
                self.log_signal.emit("未填写输出 PDF 文件名。")
                self.finished_signal.emit(0, 0, 0, "", "")
                return

            output_name = self.output_name if self.output_name.lower().endswith(".pdf") else f"{self.output_name}.pdf"
            pdf_path = self.output_dir / output_name
            text_path = self.output_dir / f"{Path(output_name).stem}_识别结果.txt"

            if pdf_path.exists() and not self.overwrite:
                self.log_signal.emit(f"失败：{pdf_path.name} | 目标 PDF 已存在")
                self.finished_signal.emit(0, 1, 0, "", "")
                return

            if text_path.exists() and not self.overwrite:
                self.log_signal.emit(f"失败：{text_path.name} | 识别结果文件已存在")
                self.finished_signal.emit(0, 1, 0, "", "")
                return

            self.output_dir.mkdir(parents=True, exist_ok=True)

            processed_pages: List[Image.Image] = []
            text_blocks: List[str] = []
            total = len(files)

            for index, src_file in enumerate(files, start=1):
                if not self._is_running:
                    self.log_signal.emit("任务已停止。")
                    break

                self.log_signal.emit(f"开始整理证件：{src_file.name}")
                try:
                    processed_image, recognized_text, info_lines = process_certificate_image(
                        src_file,
                        self.tesseract_bin,
                    )
                    processed_pages.append(processed_image)
                    success_count += 1

                    summary_block = (
                        f"[{src_file.name}]\n"
                        + "\n".join(info_lines)
                        + "\n"
                    )
                    text_block = (
                        f"[{src_file.name}]\n"
                        f"{recognized_text}\n"
                    )
                    text_blocks.append(text_block)
                    self.summary_signal.emit(summary_block)
                    self.log_signal.emit(f"成功：{src_file.name} | 已完成边框识别、自动裁剪、抠背景并识别文字")
                except Exception as exc:
                    fail_count += 1
                    self.log_signal.emit(f"失败：{src_file.name} | {exc}")

                self.progress_signal.emit(index, total)

            if not processed_pages:
                self.finished_signal.emit(success_count, fail_count, skip_count, "", "")
                return

            if not self._is_running:
                self.finished_signal.emit(success_count, fail_count, skip_count, "", "")
                return

            save_images_as_pdf(processed_pages, pdf_path)
            pdf_result_path = str(pdf_path)

            with open(text_path, "w", encoding="utf-8") as file_obj:
                file_obj.write("\n".join(text_blocks).strip())
            text_result_path = str(text_path)

            self.log_signal.emit(f"已生成 PDF：{pdf_path.name}")
            self.log_signal.emit(f"已导出识别文本：{text_path.name}")
            self.finished_signal.emit(success_count, fail_count, skip_count, pdf_result_path, text_result_path)
        except Exception as exc:
            self.log_signal.emit(f"程序异常：{exc}")
            self.finished_signal.emit(success_count, fail_count, skip_count, pdf_result_path, text_result_path)


class HomePage(QWidget):
    feature_selected = Signal(str)

    def __init__(self):
        super().__init__()
        self.init_ui()

    def init_ui(self) -> None:
        title_label = QLabel("文件格式转换工具")
        title_label.setAlignment(Qt.AlignCenter)
        title_label.setStyleSheet("""
            font-size: 32px;
            font-weight: bold;
            color: #1f2937;
            margin-bottom: 8px;
        """)

        subtitle_label = QLabel("首页直接进入功能，不再区分场景")
        subtitle_label.setAlignment(Qt.AlignCenter)
        subtitle_label.setStyleSheet("""
            font-size: 16px;
            color: #6b7280;
            margin-bottom: 10px;
        """)

        feature_items = [
            ("format_convert", "🔄", "格式转换", "批量转换图片、音频、视频、文档"),
            ("scan_to_doc", "📷", "扫描图片转文档", "OCR 识别，PDF 保留原版式"),
            ("certificate_pdf", "🪪", "证件识别整理", "识别边框、抠背景并合并为 PDF"),
            ("pdf_watermark", "📄", "PDF去水印", "移除注释型 PDF 水印"),
            ("compress", "🗜️", "图片/视频压缩", "高质量压缩，尽量保持画质"),
            ("video_download", "🎬", "视频下载去水印", "支持链接下载和本地视频处理"),
        ]

        button_grid_widget = QWidget()
        button_grid = QGridLayout(button_grid_widget)
        button_grid.setSpacing(18)
        button_grid.setContentsMargins(60, 20, 60, 20)

        for index, (feature_key, icon, title, desc) in enumerate(feature_items):
            row = index // 2
            col = index % 2
            feature_button = self.create_feature_button(feature_key, icon, title, desc)
            if len(feature_items) % 2 == 1 and index == len(feature_items) - 1:
                button_grid.addWidget(feature_button, row, col, 1, 2)
            else:
                button_grid.addWidget(feature_button, row, col)

        quick_tip_label = QLabel("建议流程：进入页面后先选文件，系统会自动补齐大部分参数。")
        quick_tip_label.setAlignment(Qt.AlignCenter)
        quick_tip_label.setStyleSheet("""
            font-size: 13px;
            color: #6b7280;
            margin-top: 8px;
        """)

        main_layout = QVBoxLayout()
        main_layout.addStretch()
        main_layout.addWidget(title_label)
        main_layout.addWidget(subtitle_label)
        main_layout.addWidget(button_grid_widget)
        main_layout.addWidget(quick_tip_label)
        main_layout.addStretch()

        self.setLayout(main_layout)

    def create_feature_button(self, feature_key: str, icon: str, title: str, desc: str) -> QPushButton:
        button = QPushButton(f"{icon}  {title}\n{desc}")
        button.setMinimumHeight(110)
        button.setCursor(Qt.PointingHandCursor)
        button.setStyleSheet("""
            QPushButton {
                background-color: #ffffff;
                color: #1f2937;
                border: 2px solid #dbe4f0;
                border-radius: 16px;
                padding: 18px 20px;
                text-align: left;
                font-size: 16px;
                font-weight: 700;
            }
            QPushButton:hover {
                border-color: #2563eb;
                background-color: #eff6ff;
            }
            QPushButton:pressed {
                background-color: #dbeafe;
            }
        """)
        button.clicked.connect(lambda _, key=feature_key: self.feature_selected.emit(key))
        return button


class VideoDownloadPage(QWidget):
    back_signal = Signal()

    def __init__(
            self,
            yt_dlp_bin: Optional[str],
            ffmpeg_bin: Optional[str],
    ):
        super().__init__()
        self.yt_dlp_bin = yt_dlp_bin
        self.ffmpeg_bin = ffmpeg_bin
        self.worker: Optional[VideoDownloadWorker] = None
        self.local_video_files: List[str] = []

        self.init_ui()
        self.apply_styles()
        self.update_output_dir_mode()

    def init_ui(self) -> None:
        back_btn = QPushButton("← 返回首页")
        back_btn.setFixedWidth(120)
        back_btn.clicked.connect(lambda: self.back_signal.emit())
        back_btn.setStyleSheet("""
            QPushButton {
                background-color: #6b7280;
                color: white;
                border: none;
                border-radius: 8px;
                padding: 8px 14px;
                font-weight: 600;
            }
            QPushButton:hover {
                background-color: #4b5563;
            }
        """)

        page_title = QLabel("视频下载 / 本地上传去水印")
        page_title.setStyleSheet("""
            font-size: 20px;
            font-weight: bold;
            color: #1f2937;
        """)

        header_layout = QHBoxLayout()
        header_layout.addWidget(back_btn)
        header_layout.addStretch()
        header_layout.addWidget(page_title)
        header_layout.addStretch()

        desc_label = QLabel("常用流程：粘贴链接或选择本地视频后，输出目录会自动补齐，通常可直接开始。")
        desc_label.setStyleSheet("color: #6b7280; font-size: 13px; margin-bottom: 10px;")

        self.url_label = QLabel("视频链接（每行一个，可留空）")
        self.url_edit = QTextEdit()
        self.url_edit.setPlaceholderText(
            "请输入视频链接，支持多个链接（每行一个）\n也支持直接粘贴分享文案，程序会自动提取其中的 URL\n例如：\n【标题】 https://www.bilibili.com/video/xxx\nhttps://www.youtube.com/watch?v=xxx"
        )
        self.url_edit.setMinimumHeight(110)
        self.url_edit.textChanged.connect(self.handle_source_changed)

        self.local_label = QLabel("本地视频（可多选，可留空）")
        self.local_edit = QLineEdit()
        self.local_edit.setReadOnly(True)
        self.local_edit.setPlaceholderText("请选择本地视频文件")
        self.local_clear_btn = QPushButton("清除")
        self.local_clear_btn.clicked.connect(self.clear_local_videos)
        self.local_browse_btn = QPushButton("选择本地视频")
        self.local_browse_btn.clicked.connect(self.choose_local_videos)

        self.output_label = QLabel("输出文件夹")
        self.output_edit = QLineEdit()
        self.output_edit.setPlaceholderText("本地视频默认原目录，链接默认下载目录")
        self.output_edit.setReadOnly(True)
        self.output_browse_btn = QPushButton("浏览")
        self.output_browse_btn.clicked.connect(self.choose_output_dir)
        self.auto_output_check = QCheckBox("自动使用推荐输出目录")
        self.auto_output_check.setChecked(True)
        self.auto_output_check.stateChanged.connect(self.update_output_dir_mode)

        self.format_label = QLabel("下载视频格式")
        self.format_combo = QComboBox()
        self.format_combo.addItems(["mp4", "webm", "mkv", "avi"])
        self.format_combo.setCurrentText("mp4")

        self.mode_label = QLabel("去水印模式")
        self.mode_combo = QComboBox()
        self.mode_combo.addItems(["极速模式", "快速模式", "全面模式"])
        self.mode_combo.setCurrentText("快速模式")

        self.remove_watermark_check = QCheckBox("自动去除水印")
        self.remove_watermark_check.setChecked(True)

        self.mode_tip_label = QLabel(
            "极速模式：仅检测前几秒，极少抽样，不做 OCR，优先速度。\n"
            "快速模式：优先四角，少量全视频抽样，先静态检测后少量 OCR。\n"
            "全面模式：扩大边缘/中部范围，抽样更多，OCR 复核更充分。"
        )
        self.mode_tip_label.setStyleSheet("color: #6b7280; font-size: 12px;")
        self.mode_tip_label.setWordWrap(True)

        yt_dlp_text = self.yt_dlp_bin if self.yt_dlp_bin else "未检测到"
        ffmpeg_text = self.ffmpeg_bin if self.ffmpeg_bin else "未检测到"
        self.status_label = QLabel(f"yt-dlp：{yt_dlp_text}    |    ffmpeg：{ffmpeg_text}（本地/下载去水印都需要）")
        self.status_label.setWordWrap(True)

        self.stage_label = QLabel("当前阶段：待开始")
        self.stage_label.setStyleSheet("font-weight: 600; color: #2563eb;")

        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)
        self.progress_bar.setFormat("%p%")

        self.start_btn = QPushButton("开始处理")
        self.stop_btn = QPushButton("停止")
        self.stop_btn.setEnabled(False)
        self.start_btn.setDefault(True)

        self.start_btn.clicked.connect(self.start_download)
        self.stop_btn.clicked.connect(self.stop_download)

        self.downloaded_list = QListWidget()
        self.downloaded_list.setMinimumHeight(150)

        self.log_edit = QTextEdit()
        self.log_edit.setReadOnly(True)

        form_group = QGroupBox("处理设置")
        form_layout = QGridLayout()
        form_layout.addWidget(self.url_label, 0, 0)
        form_layout.addWidget(self.url_edit, 0, 1, 1, 2)

        form_layout.addWidget(self.local_label, 1, 0)
        local_path_layout = QHBoxLayout()
        local_path_layout.setContentsMargins(0, 0, 0, 0)
        local_path_layout.addWidget(self.local_edit)
        local_path_layout.addWidget(self.local_clear_btn)
        local_path_layout.addWidget(self.local_browse_btn)
        form_layout.addLayout(local_path_layout, 1, 1, 1, 2)

        form_layout.addWidget(self.output_label, 2, 0)
        form_layout.addWidget(self.output_edit, 2, 1)
        form_layout.addWidget(self.output_browse_btn, 2, 2)
        form_layout.addWidget(self.auto_output_check, 3, 1, 1, 2)

        form_layout.addWidget(self.format_label, 4, 0)
        form_layout.addWidget(self.format_combo, 4, 1)

        form_layout.addWidget(self.mode_label, 5, 0)
        form_layout.addWidget(self.mode_combo, 5, 1)

        option_layout = QHBoxLayout()
        option_layout.addWidget(self.remove_watermark_check)
        option_layout.addStretch()
        form_layout.addLayout(option_layout, 6, 0, 1, 3)
        form_layout.addWidget(self.mode_tip_label, 7, 0, 1, 3)

        form_group.setLayout(form_layout)

        button_layout = QHBoxLayout()
        button_layout.addWidget(self.start_btn)
        button_layout.addWidget(self.stop_btn)
        button_layout.addStretch()

        downloaded_group = QGroupBox("输出文件")
        downloaded_layout = QVBoxLayout()
        downloaded_layout.addWidget(self.downloaded_list)
        downloaded_group.setLayout(downloaded_layout)

        log_group = QGroupBox("运行日志")
        log_layout = QVBoxLayout()
        log_layout.addWidget(self.log_edit)
        log_group.setLayout(log_layout)

        main_layout = QVBoxLayout()
        main_layout.addLayout(header_layout)
        main_layout.addWidget(desc_label)
        main_layout.addWidget(form_group)
        main_layout.addWidget(self.status_label)
        main_layout.addWidget(self.stage_label)
        main_layout.addWidget(self.progress_bar)
        main_layout.addWidget(downloaded_group)
        main_layout.addLayout(button_layout)
        main_layout.addWidget(log_group)

        self.setLayout(main_layout)

    def apply_styles(self) -> None:
        self.setStyleSheet("""
            QWidget {
                background-color: #f7f9fc;
                color: #1f2937;
                font-size: 14px;
                font-family: "Microsoft YaHei", "PingFang SC", "Segoe UI";
            }
            QGroupBox {
                border: 1px solid #d7deea;
                border-radius: 10px;
                margin-top: 12px;
                padding-top: 12px;
                background-color: #ffffff;
                font-weight: 600;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 12px;
                padding: 0 6px 0 6px;
            }
            QLineEdit, QComboBox, QTextEdit, QListWidget {
                border: 1px solid #cfd8e3;
                border-radius: 8px;
                padding: 6px 8px;
                background: #ffffff;
            }
            QPushButton {
                background-color: #2563eb;
                color: white;
                border: none;
                border-radius: 8px;
                padding: 8px 14px;
                font-weight: 600;
            }
            QPushButton:hover {
                background-color: #1d4ed8;
            }
            QPushButton:disabled {
                background-color: #94a3b8;
            }
            QProgressBar {
                border: 1px solid #d1d5db;
                border-radius: 8px;
                text-align: center;
                background: #ffffff;
                min-height: 18px;
            }
            QProgressBar::chunk {
                border-radius: 8px;
                background-color: #22c55e;
            }
            QLabel {
                background: transparent;
            }
        """)

    def choose_local_videos(self) -> None:
        file_filter = "视频文件 (*.mp4 *.avi *.mov *.mkv *.wmv *.flv *.webm *.m4v)"
        files, _ = QFileDialog.getOpenFileNames(self, "选择本地视频", "", file_filter)
        if files:
            self.local_video_files = files
            if len(files) == 1:
                self.local_edit.setText(files[0])
            else:
                self.local_edit.setText(f"已选择 {len(files)} 个文件")
            self.apply_recommended_output_dir()

    def clear_local_videos(self) -> None:
        self.local_video_files = []
        self.local_edit.clear()
        self.apply_recommended_output_dir()

    def choose_output_dir(self) -> None:
        directory = QFileDialog.getExistingDirectory(self, "选择输出文件夹")
        if directory:
            self.output_edit.setText(directory)

    def handle_source_changed(self) -> None:
        if self.auto_output_check.isChecked():
            self.apply_recommended_output_dir()

    def update_output_dir_mode(self) -> None:
        use_auto = self.auto_output_check.isChecked()
        self.output_browse_btn.setEnabled(not use_auto)
        if use_auto:
            self.apply_recommended_output_dir()

    def apply_recommended_output_dir(self) -> None:
        if not self.auto_output_check.isChecked():
            return
        if self.local_video_files:
            local_paths = [Path(item).expanduser().resolve() for item in self.local_video_files]
            recommended_dir = suggest_output_dir_from_files(local_paths)
            if recommended_dir:
                self.output_edit.setText(str(recommended_dir))
                return
        if self.url_edit.toPlainText().strip():
            self.output_edit.setText(str(get_default_download_dir()))
            return
        self.output_edit.clear()

    def append_log(self, message: str) -> None:
        self.log_edit.append(message)

    def update_stage(self, stage_text: str) -> None:
        self.stage_label.setText(f"当前阶段：{stage_text}")

    def update_progress(self, current: int, total: int) -> None:
        if total <= 0:
            self.progress_bar.setValue(0)
            return
        percent = int(current * 100 / total)
        self.progress_bar.setValue(percent)

    def on_video_downloaded(self, file_path: str) -> None:
        self.downloaded_list.addItem(file_path)

    def _get_detection_mode_value(self) -> str:
        current_mode = self.mode_combo.currentText()
        if current_mode == "极速模式":
            return "extreme"
        if current_mode == "全面模式":
            return "comprehensive"
        return "fast"

    def start_download(self) -> None:
        output_dir = self.output_edit.text().strip()
        urls_text = self.url_edit.toPlainText().strip()
        try:
            urls = extract_video_urls_from_text(urls_text) if urls_text else []
        except ValueError as exc:
            QMessageBox.warning(self, "提示", str(exc))
            return
        local_files = list(self.local_video_files)

        if not urls and not local_files:
            QMessageBox.warning(self, "提示", "请输入视频链接或选择本地视频。")
            return

        if not output_dir:
            QMessageBox.warning(self, "提示", "请选择输出文件夹。")
            return

        output_path = Path(output_dir)
        if not output_path.exists() or not output_path.is_dir():
            QMessageBox.warning(self, "提示", "输出文件夹不存在或不是有效目录。")
            return

        if urls and not self.yt_dlp_bin:
            QMessageBox.warning(
                self,
                "提示",
                "当前包含链接下载任务，但未检测到 yt-dlp。\n\n请先在终端执行：\npython -m pip install -U yt-dlp\n\n安装完成后重启本程序。"
            )
            return

        remove_watermark = self.remove_watermark_check.isChecked()
        if remove_watermark and not self.ffmpeg_bin:
            reply = QMessageBox.question(
                self,
                "提示",
                "去水印功能需要 ffmpeg，但未检测到。是否继续处理（仅下载/复制，不去水印）？",
                QMessageBox.Yes | QMessageBox.No,
                QMessageBox.Yes
            )
            if reply == QMessageBox.No:
                return
            remove_watermark = False

        self.log_edit.clear()
        self.downloaded_list.clear()
        self.progress_bar.setValue(0)
        self.update_stage("任务初始化")

        self.worker = VideoDownloadWorker(
            video_urls=urls,
            local_video_files=local_files,
            output_dir=output_dir,
            yt_dlp_bin=self.yt_dlp_bin,
            ffmpeg_bin=self.ffmpeg_bin,
            remove_watermark=remove_watermark,
            video_format=self.format_combo.currentText(),
            detection_mode=self._get_detection_mode_value(),
        )
        self.worker.log_signal.connect(self.append_log)
        self.worker.stage_signal.connect(self.update_stage)
        self.worker.progress_signal.connect(self.update_progress)
        self.worker.finished_signal.connect(self.on_finished)
        self.worker.video_downloaded.connect(self.on_video_downloaded)

        self.start_btn.setEnabled(False)
        self.stop_btn.setEnabled(True)
        self.worker.start()

    def stop_download(self) -> None:
        if self.worker and self.worker.isRunning():
            self.worker.stop()
            self.append_log("正在请求停止当前任务……")
            self.update_stage("停止中")

    def on_finished(self, success_count: int, fail_count: int, skip_count: int) -> None:
        self.start_btn.setEnabled(True)
        self.stop_btn.setEnabled(False)
        self.update_stage("已完成")

        summary = f"处理完成。成功：{success_count}，失败：{fail_count}，跳过：{skip_count}"
        self.append_log("")
        self.append_log(summary)
        QMessageBox.information(self, "完成", summary)


class FormatConvertPage(QWidget):
    back_signal = Signal()

    def __init__(
            self,
            ffmpeg_bin: Optional[str],
            libreoffice_bin: Optional[str],
            tesseract_bin: Optional[str],
            scene_key: str = "student",
    ):
        super().__init__()
        self.ffmpeg_bin = ffmpeg_bin
        self.libreoffice_bin = libreoffice_bin
        self.tesseract_bin = tesseract_bin
        self.scene_key = scene_key
        self.worker: Optional[ConvertWorker] = None
        self.selected_input_files: List[Path] = []
        self.temp_dir = Path(tempfile.gettempdir()) / "batch_convert_preview_temp"
        self.temp_dir.mkdir(parents=True, exist_ok=True)

        self.init_ui()
        self.apply_styles()
        self.update_dependency_status()
        self.update_target_formats_by_selected_files()
        self.update_output_dir_mode()

    def init_ui(self) -> None:
        back_btn = QPushButton("← 返回首页")
        back_btn.setFixedWidth(120)
        back_btn.clicked.connect(lambda: self.back_signal.emit())
        back_btn.setStyleSheet("""
            QPushButton {
                background-color: #6b7280;
                color: white;
                border: none;
                border-radius: 8px;
                padding: 8px 14px;
                font-weight: 600;
            }
            QPushButton:hover {
                background-color: #4b5563;
            }
        """)

        page_title = QLabel("格式转换")
        page_title.setStyleSheet("""
            font-size: 20px;
            font-weight: bold;
            color: #1f2937;
        """)

        header_layout = QHBoxLayout()
        header_layout.addWidget(back_btn)
        header_layout.addStretch()
        header_layout.addWidget(page_title)
        header_layout.addStretch()

        desc_label = QLabel("常用流程：选择文件后，目标格式和输出目录会自动补齐，通常可直接开始转换。")
        desc_label.setStyleSheet("color: #6b7280; font-size: 13px; margin-bottom: 10px;")

        self.input_label = QLabel("批量上传文件")
        self.input_edit = QLineEdit()
        self.input_edit.setReadOnly(True)
        self.input_edit.setPlaceholderText("请选择一个或多个文件")
        self.input_browse_btn = QPushButton("选择文件")
        self.input_browse_btn.clicked.connect(self.choose_input_files)

        self.output_label = QLabel("输出文件夹")
        self.output_edit = QLineEdit()
        self.output_edit.setPlaceholderText("默认自动使用首个文件所在目录")
        self.output_edit.setReadOnly(True)
        self.output_browse_btn = QPushButton("浏览")
        self.output_browse_btn.clicked.connect(self.choose_output_dir)
        self.auto_output_check = QCheckBox("自动使用推荐输出目录")
        self.auto_output_check.setChecked(True)
        self.auto_output_check.stateChanged.connect(self.update_output_dir_mode)

        self.target_label = QLabel("目标格式")
        self.target_combo = QComboBox()
        self.target_combo.currentTextChanged.connect(self.update_preview_panel)

        self.overwrite_check = QCheckBox("覆盖同名文件")
        self.overwrite_check.setChecked(False)
        self.overwrite_check.stateChanged.connect(self.update_preview_panel)

        self.status_label = QLabel("")
        self.status_label.setWordWrap(True)

        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)
        self.progress_bar.setFormat("%p%")

        self.select_all_btn = QPushButton("全选")
        self.select_all_btn.clicked.connect(self.select_all_files)

        self.unselect_all_btn = QPushButton("取消全选")
        self.unselect_all_btn.clicked.connect(self.unselect_all_files)

        self.preview_btn = QPushButton("预览所选文件")
        self.preview_btn.clicked.connect(self.update_preview_panel)

        self.start_btn = QPushButton("开始转换")
        self.stop_btn = QPushButton("停止转换")
        self.stop_btn.setEnabled(False)
        self.start_btn.setDefault(True)

        self.start_btn.clicked.connect(self.start_convert)
        self.stop_btn.clicked.connect(self.stop_convert)

        self.file_list = QListWidget()
        self.file_list.itemChanged.connect(self.update_selection_status)
        self.file_list.currentItemChanged.connect(self.update_preview_panel)

        self.selection_info_label = QLabel("当前未上传文件")
        self.selection_info_label.setWordWrap(True)

        self.preview_image_label = QLabel("预览区域")
        self.preview_image_label.setAlignment(Qt.AlignCenter)
        self.preview_image_label.setMinimumHeight(240)
        self.preview_image_label.setFrameShape(QFrame.StyledPanel)
        self.preview_image_label.setStyleSheet("""
            QLabel {
                background-color: #f8fafc;
                border: 1px solid #d7deea;
                border-radius: 8px;
                padding: 8px;
            }
        """)
        self._preview_original_pixmap: Optional[QPixmap] = None

        self.preview_text = QTextEdit()
        self.preview_text.setReadOnly(True)

        self.log_edit = QTextEdit()
        self.log_edit.setReadOnly(True)
        self.log_edit.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)

        form_group = QGroupBox("转换设置")
        form_layout = QGridLayout()
        form_layout.addWidget(self.input_label, 0, 0)
        form_layout.addWidget(self.input_edit, 0, 1)
        form_layout.addWidget(self.input_browse_btn, 0, 2)

        form_layout.addWidget(self.output_label, 1, 0)
        form_layout.addWidget(self.output_edit, 1, 1)
        form_layout.addWidget(self.output_browse_btn, 1, 2)
        form_layout.addWidget(self.auto_output_check, 2, 1, 1, 2)

        form_layout.addWidget(self.target_label, 3, 0)
        form_layout.addWidget(self.target_combo, 3, 1)

        option_layout = QHBoxLayout()
        option_layout.addWidget(self.overwrite_check)
        option_layout.addStretch()
        form_layout.addLayout(option_layout, 4, 0, 1, 3)

        form_group.setLayout(form_layout)

        left_group = QGroupBox("待转换文件")
        left_layout = QVBoxLayout()
        file_btn_layout = QHBoxLayout()
        file_btn_layout.addWidget(self.select_all_btn)
        file_btn_layout.addWidget(self.unselect_all_btn)
        file_btn_layout.addWidget(self.preview_btn)
        file_btn_layout.addStretch()

        left_layout.addLayout(file_btn_layout)
        left_layout.addWidget(self.selection_info_label)
        left_layout.addWidget(self.file_list)
        left_group.setLayout(left_layout)

        right_group = QGroupBox("转换效果预览")
        right_layout = QVBoxLayout()
        right_layout.addWidget(self.preview_image_label)
        right_layout.addWidget(self.preview_text)
        right_group.setLayout(right_layout)

        splitter = QSplitter(Qt.Horizontal)
        splitter.addWidget(left_group)
        splitter.addWidget(right_group)
        splitter.setSizes([560, 520])

        button_layout = QHBoxLayout()
        button_layout.addWidget(self.start_btn)
        button_layout.addWidget(self.stop_btn)
        button_layout.addStretch()

        log_group = QGroupBox("运行日志")
        log_layout = QVBoxLayout()
        log_layout.addWidget(self.log_edit)
        log_group.setLayout(log_layout)

        main_layout = QVBoxLayout()
        main_layout.addLayout(header_layout)
        main_layout.addWidget(desc_label)
        main_layout.addWidget(form_group)
        main_layout.addWidget(self.status_label)
        main_layout.addWidget(self.progress_bar)
        main_layout.addWidget(splitter)
        main_layout.addLayout(button_layout)
        main_layout.addWidget(log_group)

        self.setLayout(main_layout)

    def apply_styles(self) -> None:
        self.setStyleSheet("""
            QWidget {
                background-color: #f7f9fc;
                color: #1f2937;
                font-size: 14px;
                font-family: "Microsoft YaHei", "PingFang SC", "Segoe UI";
            }
            QGroupBox {
                border: 1px solid #d7deea;
                border-radius: 10px;
                margin-top: 12px;
                padding-top: 12px;
                background-color: #ffffff;
                font-weight: 600;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 12px;
                padding: 0 6px 0 6px;
            }
            QLineEdit, QComboBox, QTextEdit, QListWidget {
                border: 1px solid #cfd8e3;
                border-radius: 8px;
                padding: 6px 8px;
                background: #ffffff;
            }
            QPushButton {
                background-color: #2563eb;
                color: white;
                border: none;
                border-radius: 8px;
                padding: 8px 14px;
                font-weight: 600;
            }
            QPushButton:hover {
                background-color: #1d4ed8;
            }
            QPushButton:disabled {
                background-color: #94a3b8;
            }
            QProgressBar {
                border: 1px solid #d1d5db;
                border-radius: 8px;
                text-align: center;
                background: #ffffff;
                min-height: 18px;
            }
            QProgressBar::chunk {
                border-radius: 8px;
                background-color: #22c55e;
            }
            QLabel {
                background: transparent;
            }
        """)

    def get_selected_files(self) -> List[str]:
        selected_files: List[str] = []
        for index in range(self.file_list.count()):
            item = self.file_list.item(index)
            if item.checkState() == Qt.Checked:
                selected_files.append(item.data(Qt.UserRole))
        return selected_files

    def update_dependency_status(self) -> None:
        ffmpeg_text = self.ffmpeg_bin if self.ffmpeg_bin else "未检测到"
        libreoffice_text = self.libreoffice_bin if self.libreoffice_bin else "未检测到"
        tesseract_text = self.tesseract_bin if self.tesseract_bin else "未检测到"
        pdf2docx_text = "已安装" if has_pdf2docx_engine() else "未安装"
        ocrmypdf_text = "已安装" if has_ocrmypdf_engine() else "未安装"
        freep2w_text = "已安装" if has_freep2w_engine() else "未安装"
        self.status_label.setText(
            "环境检测："
            f"ffmpeg：{ffmpeg_text}    |    "
            f"LibreOffice：{libreoffice_text}    |    "
            f"pdf2docx：{pdf2docx_text}    |    "
            f"OCRmyPDF：{ocrmypdf_text}    |    "
            f"FreeP2W：{freep2w_text}    |    "
            f"Tesseract：{tesseract_text}"
        )

    def update_target_formats_by_selected_files(self) -> None:
        available_targets = set()
        has_pdf = False

        for file_path in self.selected_input_files:
            available_targets.update(get_supported_targets_for_file(file_path))
            if file_path.suffix.lower() == ".pdf":
                has_pdf = True

        sorted_targets = sorted(available_targets)

        self.target_combo.blockSignals(True)
        self.target_combo.clear()
        self.target_combo.addItems(sorted_targets)

        if sorted_targets:
            if has_pdf and "docx" in sorted_targets:
                self.target_combo.setCurrentText("docx")
            elif "pdf" in sorted_targets:
                self.target_combo.setCurrentText("pdf")

        self.target_combo.blockSignals(False)

        self.update_preview_panel()

    def choose_input_files(self) -> None:
        file_paths, _ = QFileDialog.getOpenFileNames(
            self,
            "选择一个或多个文件",
            "",
            "所有支持文件 (*.jpg *.jpeg *.png *.bmp *.gif *.webp *.tif *.tiff *.ico "
            "*.mp3 *.wav *.flac *.aac *.ogg *.m4a *.wma "
            "*.mp4 *.avi *.mov *.mkv *.wmv *.flv *.webm *.m4v "
            "*.doc *.docx *.odt *.rtf *.txt *.xls *.xlsx *.ods *.csv *.ppt *.pptx *.odp *.pdf)"
        )

        if not file_paths:
            return

        new_files = [Path(item).expanduser().resolve() for item in file_paths]
        self.selected_input_files = new_files
        self.input_edit.setText(f"已选择 {len(new_files)} 个文件")
        self.load_file_list_from_uploaded_files()
        self.apply_recommended_output_dir()
        self.update_target_formats_by_selected_files()

    def choose_output_dir(self) -> None:
        directory = QFileDialog.getExistingDirectory(self, "选择输出文件夹")
        if directory:
            self.output_edit.setText(directory)
            self.update_preview_panel()

    def update_output_dir_mode(self) -> None:
        use_auto = self.auto_output_check.isChecked()
        self.output_browse_btn.setEnabled(not use_auto)
        if use_auto:
            self.apply_recommended_output_dir()
        self.update_preview_panel()

    def apply_recommended_output_dir(self) -> None:
        recommended_dir = suggest_output_dir_from_files(self.selected_input_files)
        if recommended_dir:
            self.output_edit.setText(str(recommended_dir))
        elif self.auto_output_check.isChecked():
            self.output_edit.clear()

    def append_log(self, message: str) -> None:
        self.log_edit.append(message)

    def update_progress(self, current: int, total: int) -> None:
        if total <= 0:
            self.progress_bar.setValue(0)
            return
        percent = int(current * 100 / total)
        self.progress_bar.setValue(percent)

    def load_file_list_from_uploaded_files(self) -> None:
        self.file_list.blockSignals(True)
        self.file_list.clear()

        valid_files: List[Path] = []
        invalid_files: List[Path] = []

        for file_path in self.selected_input_files:
            if detect_category_by_file(file_path):
                valid_files.append(file_path)
            else:
                invalid_files.append(file_path)

        self.selected_input_files = valid_files

        for file_path in valid_files:
            category = detect_category_by_file(file_path)
            category_text = CATEGORY_NAME_MAP.get(category, "未知类型")
            item = QListWidgetItem(f"[{category_text}] {file_path.name}")
            item.setData(Qt.UserRole, str(file_path.resolve()))
            item.setToolTip(str(file_path))
            item.setFlags(item.flags() | Qt.ItemIsUserCheckable | Qt.ItemIsSelectable | Qt.ItemIsEnabled)
            item.setCheckState(Qt.Checked)
            self.file_list.addItem(item)

        self.file_list.blockSignals(False)
        self.update_selection_status()

        if self.file_list.count() > 0:
            self.file_list.setCurrentRow(0)

        if invalid_files:
            invalid_names = "\n".join(str(item) for item in invalid_files[:10])
            self.append_log(f"以下文件未识别到支持格式，已忽略：\n{invalid_names}")

        self.update_preview_panel()

    def select_all_files(self) -> None:
        self.file_list.blockSignals(True)
        for index in range(self.file_list.count()):
            self.file_list.item(index).setCheckState(Qt.Checked)
        self.file_list.blockSignals(False)
        self.update_selection_status()
        self.update_preview_panel()

    def unselect_all_files(self) -> None:
        self.file_list.blockSignals(True)
        for index in range(self.file_list.count()):
            self.file_list.item(index).setCheckState(Qt.Unchecked)
        self.file_list.blockSignals(False)
        self.update_selection_status()
        self.update_preview_panel()

    def update_selection_status(self) -> None:
        total = self.file_list.count()
        selected = len(self.get_selected_files())
        self.selection_info_label.setText(
            f"已上传 {total} 个可转换文件，当前选中 {selected} 个用于转换。"
        )

    def update_preview_panel(self) -> None:
        current_item = self.file_list.currentItem()
        output_dir = self.output_edit.text().strip()
        target_ext = normalize_ext(self.target_combo.currentText())

        if not current_item:
            self._preview_original_pixmap = None
            self.preview_image_label.setPixmap(QPixmap())
            self.preview_image_label.setText("预览区域")
            self.preview_text.setPlainText("请先上传文件，并在左侧文件列表中选择一个文件查看预览。")
            return

        if not output_dir or not target_ext:
            self._preview_original_pixmap = None
            self.preview_image_label.setPixmap(QPixmap())
            self.preview_image_label.setText("预览区域")
            self.preview_text.setPlainText("请选择输出文件夹和目标格式后查看预览。")
            return

        file_path = Path(current_item.data(Qt.UserRole))
        output_root = Path(output_dir)

        self.preview_text.setPlainText(
            build_preview_text(
                src_file=file_path,
                dst_root=output_root,
                target_ext=target_ext,
            )
        )

        if detect_category_by_file(file_path) == "image":
            original_pixmap = create_image_preview(file_path, None)
            if original_pixmap:
                self._preview_original_pixmap = original_pixmap
                self.refresh_preview_image()
                self.preview_image_label.setText("")
            else:
                self._preview_original_pixmap = None
                self.preview_image_label.setPixmap(QPixmap())
                self.preview_image_label.setText("当前图片无法预览")
        else:
            self._preview_original_pixmap = None
            self.preview_image_label.setPixmap(QPixmap())
            self.preview_image_label.setText("当前类型暂不提供画面预览，可查看下方转换信息预览")

    def refresh_preview_image(self) -> None:
        if not self._preview_original_pixmap or self._preview_original_pixmap.isNull():
            return

        label_size = self.preview_image_label.size()
        target_width = max(100, label_size.width() - 16)
        target_height = max(100, label_size.height() - 16)

        scaled_pixmap = self._preview_original_pixmap.scaled(
            target_width,
            target_height,
            Qt.KeepAspectRatio,
            Qt.SmoothTransformation
        )
        self.preview_image_label.setPixmap(scaled_pixmap)

    def resizeEvent(self, event) -> None:
        super().resizeEvent(event)
        self.refresh_preview_image()

    def start_convert(self) -> None:
        output_dir = self.output_edit.text().strip()
        target_ext = normalize_ext(self.target_combo.currentText())
        selected_files = self.get_selected_files()

        if not selected_files:
            QMessageBox.warning(self, "提示", "请至少勾选一个需要转换的文件。")
            return

        if not output_dir:
            QMessageBox.warning(self, "提示", "请选择输出文件夹。")
            return

        if not target_ext:
            QMessageBox.warning(self, "提示", "当前没有可用的目标格式，请检查上传文件类型。")
            return

        output_path = Path(output_dir)
        if not output_path.exists() or not output_path.is_dir():
            QMessageBox.warning(self, "提示", "输出文件夹不存在或不是有效目录。")
            return

        for file_path_str in selected_files:
            file_path = Path(file_path_str)
            if file_path.suffix.lower() == f".{target_ext}":
                QMessageBox.warning(self, "提示", f"文件 {file_path.name} 已是目标格式，请选择其他格式。")
                return

        requires_ffmpeg = False
        requires_libreoffice = False
        requires_pdf2docx = False
        requires_ocrmypdf = False

        for file_path_str in selected_files:
            file_path = Path(file_path_str)
            category = detect_category_by_file(file_path)
            if category in {"audio", "video"}:
                requires_ffmpeg = True
            if file_path.suffix.lower() == ".pdf" and target_ext == "docx":
                requires_pdf2docx = True
                if uses_ocrmypdf_engine(file_path, target_ext):
                    requires_ocrmypdf = True
            elif uses_libreoffice_engine(file_path, target_ext):
                requires_libreoffice = True

        if requires_ffmpeg and not self.ffmpeg_bin:
            QMessageBox.warning(self, "提示", "当前选中文件包含音频或视频，但未检测到 ffmpeg。")
            return

        if requires_pdf2docx and not has_pdf2docx_engine():
            QMessageBox.warning(
                self,
                "提示",
                "当前转换包含 PDF 转 DOCX，但未安装 pdf2docx。请先执行：pip install pdf2docx",
            )
            return

        if requires_ocrmypdf and not has_ocrmypdf_engine():
            QMessageBox.warning(
                self,
                "提示",
                "当前转换包含扫描件 PDF 转 DOCX，但未安装 OCRmyPDF。请先执行：pip install ocrmypdf",
            )
            return

        if requires_libreoffice and not self.libreoffice_bin:
            QMessageBox.warning(
                self,
                "提示",
                "当前转换包含 LibreOffice 引擎负责的文档处理，或图片转 DOCX，但未检测到 LibreOffice。",
            )
            return

        self.log_edit.clear()
        self.progress_bar.setValue(0)

        self.worker = ConvertWorker(
            selected_files=selected_files,
            output_dir=output_dir,
            target_ext=target_ext,
            overwrite=self.overwrite_check.isChecked(),
            ffmpeg_bin=self.ffmpeg_bin,
            libreoffice_bin=self.libreoffice_bin,
            tesseract_bin=self.tesseract_bin,
            mode="convert",
        )
        self.worker.log_signal.connect(self.append_log)
        self.worker.progress_signal.connect(self.update_progress)
        self.worker.finished_signal.connect(self.on_finished)

        self.start_btn.setEnabled(False)
        self.stop_btn.setEnabled(True)
        self.worker.start()

    def stop_convert(self) -> None:
        if self.worker and self.worker.isRunning():
            self.worker.stop()
            self.append_log("正在请求停止当前转换任务……")

    def on_finished(self, success_count: int, fail_count: int, skip_count: int) -> None:
        self.start_btn.setEnabled(True)
        self.stop_btn.setEnabled(False)

        summary = (
            f"转换完成。\n"
            f"成功：{success_count}\n"
            f"失败：{fail_count}\n"
            f"跳过：{skip_count}"
        )

        self.append_log("")
        self.append_log(summary)
        QMessageBox.information(self, "完成", summary)


class ScanToDocPage(QWidget):
    back_signal = Signal()

    def __init__(
            self,
            tesseract_bin: Optional[str],
            libreoffice_bin: Optional[str],
    ):
        super().__init__()
        self.tesseract_bin = tesseract_bin
        self.libreoffice_bin = libreoffice_bin
        self.worker: Optional[ConvertWorker] = None
        self.selected_input_files: List[Path] = []

        self.init_ui()
        self.apply_styles()

    def init_ui(self) -> None:
        back_btn = QPushButton("← 返回首页")
        back_btn.setFixedWidth(120)
        back_btn.clicked.connect(lambda: self.back_signal.emit())
        back_btn.setStyleSheet("""
            QPushButton {
                background-color: #6b7280;
                color: white;
                border: none;
                border-radius: 8px;
                padding: 8px 14px;
                font-weight: 600;
            }
            QPushButton:hover {
                background-color: #4b5563;
            }
        """)

        page_title = QLabel("扫描图片转文档")
        page_title.setStyleSheet("""
            font-size: 20px;
            font-weight: bold;
            color: #1f2937;
        """)

        header_layout = QHBoxLayout()
        header_layout.addWidget(back_btn)
        header_layout.addStretch()
        header_layout.addWidget(page_title)
        header_layout.addStretch()

        desc_label = QLabel("常用流程：选择图片后会自动补齐输出目录，通常直接点开始即可。PDF 输出保留原图版式并叠加可检索文本层")
        desc_label.setStyleSheet("color: #6b7280; font-size: 13px; margin-bottom: 10px;")

        self.input_label = QLabel("选择图片文件")
        self.input_edit = QLineEdit()
        self.input_edit.setReadOnly(True)
        self.input_edit.setPlaceholderText("请选择需要识别的图片文件")
        self.input_browse_btn = QPushButton("选择文件")
        self.input_browse_btn.clicked.connect(self.choose_input_files)

        self.output_label = QLabel("输出文件夹")
        self.output_edit = QLineEdit()
        self.output_edit.setPlaceholderText("默认自动使用首个文件所在目录")
        self.output_edit.setReadOnly(True)
        self.output_browse_btn = QPushButton("浏览")
        self.output_browse_btn.clicked.connect(self.choose_output_dir)
        self.auto_output_check = QCheckBox("自动使用推荐输出目录")
        self.auto_output_check.setChecked(True)
        self.auto_output_check.stateChanged.connect(self.update_output_dir_mode)

        self.target_label = QLabel("输出格式")
        self.target_combo = QComboBox()
        self.target_combo.addItems(["pdf", "docx", "txt"])
        self.target_combo.setCurrentText("pdf")

        self.overwrite_check = QCheckBox("覆盖同名文件")
        self.overwrite_check.setChecked(False)

        tesseract_text = self.tesseract_bin if self.tesseract_bin else "未检测到"
        libreoffice_text = self.libreoffice_bin if self.libreoffice_bin else "未检测到"
        self.status_label = QLabel(f"Tesseract OCR：{tesseract_text}    |    LibreOffice：{libreoffice_text}（DOCX 需要）")
        self.status_label.setWordWrap(True)

        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)

        self.start_btn = QPushButton("开始转换")
        self.stop_btn = QPushButton("停止")
        self.stop_btn.setEnabled(False)
        self.start_btn.setDefault(True)

        self.start_btn.clicked.connect(self.start_convert)
        self.stop_btn.clicked.connect(self.stop_convert)

        self.file_list = QListWidget()
        self.file_list.setMinimumHeight(200)

        self.log_edit = QTextEdit()
        self.log_edit.setReadOnly(True)

        form_group = QGroupBox("转换设置")
        form_layout = QGridLayout()
        form_layout.addWidget(self.input_label, 0, 0)
        form_layout.addWidget(self.input_edit, 0, 1)
        form_layout.addWidget(self.input_browse_btn, 0, 2)

        form_layout.addWidget(self.output_label, 1, 0)
        form_layout.addWidget(self.output_edit, 1, 1)
        form_layout.addWidget(self.output_browse_btn, 1, 2)
        form_layout.addWidget(self.auto_output_check, 2, 1, 1, 2)

        form_layout.addWidget(self.target_label, 3, 0)
        form_layout.addWidget(self.target_combo, 3, 1)

        option_layout = QHBoxLayout()
        option_layout.addWidget(self.overwrite_check)
        option_layout.addStretch()
        form_layout.addLayout(option_layout, 4, 0, 1, 3)

        form_group.setLayout(form_layout)

        button_layout = QHBoxLayout()
        button_layout.addWidget(self.start_btn)
        button_layout.addWidget(self.stop_btn)
        button_layout.addStretch()

        log_group = QGroupBox("运行日志")
        log_layout = QVBoxLayout()
        log_layout.addWidget(self.log_edit)
        log_group.setLayout(log_layout)

        main_layout = QVBoxLayout()
        main_layout.addLayout(header_layout)
        main_layout.addWidget(desc_label)
        main_layout.addWidget(form_group)
        main_layout.addWidget(self.status_label)
        main_layout.addWidget(self.progress_bar)
        main_layout.addWidget(QLabel("待处理文件："))
        main_layout.addWidget(self.file_list)
        main_layout.addLayout(button_layout)
        main_layout.addWidget(log_group)

        self.setLayout(main_layout)

    def apply_styles(self) -> None:
        self.setStyleSheet("""
            QWidget {
                background-color: #f7f9fc;
                color: #1f2937;
                font-size: 14px;
                font-family: "Microsoft YaHei", "PingFang SC", "Segoe UI";
            }
            QGroupBox {
                border: 1px solid #d7deea;
                border-radius: 10px;
                margin-top: 12px;
                padding-top: 12px;
                background-color: #ffffff;
                font-weight: 600;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 12px;
                padding: 0 6px 0 6px;
            }
            QLineEdit, QComboBox, QTextEdit, QListWidget {
                border: 1px solid #cfd8e3;
                border-radius: 8px;
                padding: 6px 8px;
                background: #ffffff;
            }
            QPushButton {
                background-color: #2563eb;
                color: white;
                border: none;
                border-radius: 8px;
                padding: 8px 14px;
                font-weight: 600;
            }
            QPushButton:hover {
                background-color: #1d4ed8;
            }
            QPushButton:disabled {
                background-color: #94a3b8;
            }
            QProgressBar {
                border: 1px solid #d1d5db;
                border-radius: 8px;
                text-align: center;
                background: #ffffff;
                min-height: 18px;
            }
            QProgressBar::chunk {
                border-radius: 8px;
                background-color: #22c55e;
            }
            QLabel {
                background: transparent;
            }
        """)

    def choose_input_files(self) -> None:
        file_paths, _ = QFileDialog.getOpenFileNames(
            self,
            "选择图片文件",
            "",
            "图片文件 (*.jpg *.jpeg *.png *.bmp *.gif *.webp *.tif *.tiff)"
        )

        if not file_paths:
            return

        self.selected_input_files = [Path(item).expanduser().resolve() for item in file_paths]
        self.input_edit.setText(f"已选择 {len(self.selected_input_files)} 个文件")
        self.apply_recommended_output_dir()

        self.file_list.clear()
        for file_path in self.selected_input_files:
            self.file_list.addItem(file_path.name)

    def choose_output_dir(self) -> None:
        directory = QFileDialog.getExistingDirectory(self, "选择输出文件夹")
        if directory:
            self.output_edit.setText(directory)

    def update_output_dir_mode(self) -> None:
        use_auto = self.auto_output_check.isChecked()
        self.output_browse_btn.setEnabled(not use_auto)
        if use_auto:
            self.apply_recommended_output_dir()

    def apply_recommended_output_dir(self) -> None:
        recommended_dir = suggest_output_dir_from_files(self.selected_input_files)
        if recommended_dir:
            self.output_edit.setText(str(recommended_dir))
        elif self.auto_output_check.isChecked():
            self.output_edit.clear()

    def append_log(self, message: str) -> None:
        self.log_edit.append(message)

    def update_progress(self, current: int, total: int) -> None:
        if total <= 0:
            self.progress_bar.setValue(0)
            return
        percent = int(current * 100 / total)
        self.progress_bar.setValue(percent)

    def start_convert(self) -> None:
        output_dir = self.output_edit.text().strip()
        target_ext = self.target_combo.currentText()

        if not self.selected_input_files:
            QMessageBox.warning(self, "提示", "请选择需要转换的图片文件。")
            return

        if not output_dir:
            QMessageBox.warning(self, "提示", "请选择输出文件夹。")
            return

        output_path = Path(output_dir)
        if not output_path.exists() or not output_path.is_dir():
            QMessageBox.warning(self, "提示", "输出文件夹不存在或不是有效目录。")
            return

        if not self.tesseract_bin:
            QMessageBox.warning(self, "提示", "未检测到 Tesseract OCR，请先安装。")
            return

        if target_ext == "docx" and not self.libreoffice_bin:
            QMessageBox.warning(self, "提示", "OCR 转 DOCX 需要 LibreOffice，请先安装。")
            return

        self.log_edit.clear()
        self.progress_bar.setValue(0)

        self.worker = ConvertWorker(
            selected_files=[str(f) for f in self.selected_input_files],
            output_dir=output_dir,
            target_ext=target_ext,
            overwrite=self.overwrite_check.isChecked(),
            ffmpeg_bin=None,
            libreoffice_bin=self.libreoffice_bin,
            tesseract_bin=self.tesseract_bin,
            mode="scan",
        )
        self.worker.log_signal.connect(self.append_log)
        self.worker.progress_signal.connect(self.update_progress)
        self.worker.finished_signal.connect(self.on_finished)

        self.start_btn.setEnabled(False)
        self.stop_btn.setEnabled(True)
        self.worker.start()

    def stop_convert(self) -> None:
        if self.worker and self.worker.isRunning():
            self.worker.stop()
            self.append_log("正在请求停止当前任务……")

    def on_finished(self, success_count: int, fail_count: int, skip_count: int) -> None:
        self.start_btn.setEnabled(True)
        self.stop_btn.setEnabled(False)
        self.progress_bar.setRange(0, 100)
        if success_count > 0 or fail_count > 0 or skip_count > 0:
            self.progress_bar.setValue(100)
        else:
            self.progress_bar.setValue(0)

        summary = f"处理完成。成功：{success_count}，失败：{fail_count}，跳过：{skip_count}"
        self.append_log("")
        self.append_log(summary)
        QMessageBox.information(self, "完成", summary)


class CertificatePdfPage(QWidget):
    back_signal = Signal()

    def __init__(self, tesseract_bin: Optional[str]):
        super().__init__()
        self.tesseract_bin = tesseract_bin
        self.worker: Optional[CertificatePdfWorker] = None
        self.selected_input_files: List[Path] = []

        self.init_ui()
        self.apply_styles()
        self.update_output_dir_mode()

    def init_ui(self) -> None:
        back_btn = QPushButton("← 返回首页")
        back_btn.setFixedWidth(120)
        back_btn.clicked.connect(lambda: self.back_signal.emit())
        back_btn.setStyleSheet("""
            QPushButton {
                background-color: #6b7280;
                color: white;
                border: none;
                border-radius: 8px;
                padding: 8px 14px;
                font-weight: 600;
            }
            QPushButton:hover {
                background-color: #4b5563;
            }
        """)

        page_title = QLabel("证件识别整理")
        page_title.setStyleSheet("""
            font-size: 20px;
            font-weight: bold;
            color: #1f2937;
        """)

        header_layout = QHBoxLayout()
        header_layout.addWidget(back_btn)
        header_layout.addStretch()
        header_layout.addWidget(page_title)
        header_layout.addStretch()

        desc_label = QLabel("选择多张证件图片后，会自动识别证件边框、矫正角度、自动裁剪、抠掉背景，并直接合并成一个 PDF，同时导出识别文本。")
        desc_label.setStyleSheet("color: #6b7280; font-size: 13px; margin-bottom: 10px;")
        desc_label.setWordWrap(True)

        self.input_label = QLabel("选择证件图片")
        self.input_edit = QLineEdit()
        self.input_edit.setReadOnly(True)
        self.input_edit.setPlaceholderText("请选择一张或多张证件图片")
        self.input_browse_btn = QPushButton("选择文件")
        self.input_browse_btn.clicked.connect(self.choose_input_files)

        self.output_label = QLabel("输出文件夹")
        self.output_edit = QLineEdit()
        self.output_edit.setReadOnly(True)
        self.output_edit.setPlaceholderText("默认自动使用首个文件所在目录")
        self.output_browse_btn = QPushButton("浏览")
        self.output_browse_btn.clicked.connect(self.choose_output_dir)
        self.auto_output_check = QCheckBox("自动使用推荐输出目录")
        self.auto_output_check.setChecked(True)
        self.auto_output_check.stateChanged.connect(self.update_output_dir_mode)

        self.output_name_label = QLabel("输出 PDF 名称")
        self.output_name_edit = QLineEdit()
        self.output_name_edit.setPlaceholderText("默认按首张图片名自动生成")

        self.overwrite_check = QCheckBox("覆盖同名文件")
        self.overwrite_check.setChecked(False)

        tesseract_text = self.tesseract_bin if self.tesseract_bin else "未检测到"
        self.status_label = QLabel(f"Tesseract OCR：{tesseract_text}    |    依赖 OpenCV 做证件区域检测与矫正")
        self.status_label.setWordWrap(True)

        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)

        self.start_btn = QPushButton("开始整理")
        self.stop_btn = QPushButton("停止")
        self.stop_btn.setEnabled(False)
        self.start_btn.setDefault(True)
        self.start_btn.clicked.connect(self.start_process)
        self.stop_btn.clicked.connect(self.stop_process)

        self.file_list = QListWidget()
        self.file_list.setMinimumHeight(180)

        self.result_edit = QTextEdit()
        self.result_edit.setReadOnly(True)
        self.result_edit.setPlaceholderText("识别出的证件关键信息会显示在这里，并同步导出为同名 txt。")

        self.log_edit = QTextEdit()
        self.log_edit.setReadOnly(True)

        form_group = QGroupBox("处理设置")
        form_layout = QGridLayout()
        form_layout.addWidget(self.input_label, 0, 0)
        form_layout.addWidget(self.input_edit, 0, 1)
        form_layout.addWidget(self.input_browse_btn, 0, 2)

        form_layout.addWidget(self.output_label, 1, 0)
        form_layout.addWidget(self.output_edit, 1, 1)
        form_layout.addWidget(self.output_browse_btn, 1, 2)
        form_layout.addWidget(self.auto_output_check, 2, 1, 1, 2)

        form_layout.addWidget(self.output_name_label, 3, 0)
        form_layout.addWidget(self.output_name_edit, 3, 1, 1, 2)

        option_layout = QHBoxLayout()
        option_layout.addWidget(self.overwrite_check)
        option_layout.addStretch()
        form_layout.addLayout(option_layout, 4, 0, 1, 3)
        form_group.setLayout(form_layout)

        button_layout = QHBoxLayout()
        button_layout.addWidget(self.start_btn)
        button_layout.addWidget(self.stop_btn)
        button_layout.addStretch()

        result_group = QGroupBox("识别结果")
        result_layout = QVBoxLayout()
        result_layout.addWidget(self.result_edit)
        result_group.setLayout(result_layout)

        log_group = QGroupBox("运行日志")
        log_layout = QVBoxLayout()
        log_layout.addWidget(self.log_edit)
        log_group.setLayout(log_layout)

        main_layout = QVBoxLayout()
        main_layout.addLayout(header_layout)
        main_layout.addWidget(desc_label)
        main_layout.addWidget(form_group)
        main_layout.addWidget(self.status_label)
        main_layout.addWidget(self.progress_bar)
        main_layout.addWidget(QLabel("待处理文件："))
        main_layout.addWidget(self.file_list)
        main_layout.addLayout(button_layout)
        main_layout.addWidget(result_group)
        main_layout.addWidget(log_group)

        self.setLayout(main_layout)

    def apply_styles(self) -> None:
        self.setStyleSheet("""
            QWidget {
                background-color: #f7f9fc;
                color: #1f2937;
                font-size: 14px;
                font-family: "Microsoft YaHei", "PingFang SC", "Segoe UI";
            }
            QGroupBox {
                border: 1px solid #d7deea;
                border-radius: 10px;
                margin-top: 12px;
                padding-top: 12px;
                background-color: #ffffff;
                font-weight: 600;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 12px;
                padding: 0 6px 0 6px;
            }
            QLineEdit, QTextEdit, QListWidget {
                border: 1px solid #cfd8e3;
                border-radius: 8px;
                padding: 6px 8px;
                background: #ffffff;
            }
            QPushButton {
                background-color: #2563eb;
                color: white;
                border: none;
                border-radius: 8px;
                padding: 8px 14px;
                font-weight: 600;
            }
            QPushButton:hover {
                background-color: #1d4ed8;
            }
            QPushButton:disabled {
                background-color: #94a3b8;
            }
            QProgressBar {
                border: 1px solid #d1d5db;
                border-radius: 8px;
                text-align: center;
                background: #ffffff;
                min-height: 18px;
            }
            QProgressBar::chunk {
                border-radius: 8px;
                background-color: #22c55e;
            }
            QLabel {
                background: transparent;
            }
        """)

    def choose_input_files(self) -> None:
        file_paths, _ = QFileDialog.getOpenFileNames(
            self,
            "选择证件图片",
            "",
            "图片文件 (*.jpg *.jpeg *.png *.bmp *.webp *.tif *.tiff)"
        )
        if not file_paths:
            return

        self.selected_input_files = [Path(item).expanduser().resolve() for item in file_paths]
        self.input_edit.setText(f"已选择 {len(self.selected_input_files)} 个文件")
        self.apply_recommended_output_dir()
        self.apply_default_output_name()

        self.file_list.clear()
        for file_path in self.selected_input_files:
            self.file_list.addItem(file_path.name)

    def choose_output_dir(self) -> None:
        directory = QFileDialog.getExistingDirectory(self, "选择输出文件夹")
        if directory:
            self.output_edit.setText(directory)

    def update_output_dir_mode(self) -> None:
        use_auto = self.auto_output_check.isChecked()
        self.output_browse_btn.setEnabled(not use_auto)
        if use_auto:
            self.apply_recommended_output_dir()

    def apply_recommended_output_dir(self) -> None:
        recommended_dir = suggest_output_dir_from_files(self.selected_input_files)
        if recommended_dir:
            self.output_edit.setText(str(recommended_dir))
        elif self.auto_output_check.isChecked():
            self.output_edit.clear()

    def apply_default_output_name(self) -> None:
        if not self.selected_input_files:
            self.output_name_edit.clear()
            return
        first_file = self.selected_input_files[0]
        if len(self.selected_input_files) == 1:
            name = f"{first_file.stem}_证件整理.pdf"
        else:
            name = f"{first_file.stem}_等{len(self.selected_input_files)}张_证件整理.pdf"
        self.output_name_edit.setText(name)

    def append_log(self, message: str) -> None:
        self.log_edit.append(message)

    def append_summary(self, message: str) -> None:
        self.result_edit.append(message.rstrip())
        self.result_edit.append("")

    def update_progress(self, current: int, total: int) -> None:
        if total <= 0:
            self.progress_bar.setValue(0)
            return
        self.progress_bar.setValue(int(current * 100 / total))

    def start_process(self) -> None:
        output_dir = self.output_edit.text().strip()
        output_name = self.output_name_edit.text().strip()

        if not self.selected_input_files:
            QMessageBox.warning(self, "提示", "请选择证件图片。")
            return

        if not output_dir:
            QMessageBox.warning(self, "提示", "请选择输出文件夹。")
            return

        if not output_name:
            QMessageBox.warning(self, "提示", "请输入输出 PDF 文件名。")
            return

        output_path = Path(output_dir)
        if not output_path.exists() or not output_path.is_dir():
            QMessageBox.warning(self, "提示", "输出文件夹不存在或不是有效目录。")
            return

        if not self.tesseract_bin:
            QMessageBox.warning(self, "提示", "未检测到 Tesseract OCR，请先安装。")
            return

        self.log_edit.clear()
        self.result_edit.clear()
        self.progress_bar.setValue(0)

        self.worker = CertificatePdfWorker(
            selected_files=[str(item) for item in self.selected_input_files],
            output_dir=output_dir,
            output_name=output_name,
            overwrite=self.overwrite_check.isChecked(),
            tesseract_bin=self.tesseract_bin,
        )
        self.worker.log_signal.connect(self.append_log)
        self.worker.progress_signal.connect(self.update_progress)
        self.worker.summary_signal.connect(self.append_summary)
        self.worker.finished_signal.connect(self.on_finished)

        self.start_btn.setEnabled(False)
        self.stop_btn.setEnabled(True)
        self.worker.start()

    def stop_process(self) -> None:
        if self.worker and self.worker.isRunning():
            self.worker.stop()
            self.append_log("正在请求停止当前任务……")

    def on_finished(
            self,
            success_count: int,
            fail_count: int,
            skip_count: int,
            pdf_result_path: str,
            text_result_path: str,
    ) -> None:
        self.start_btn.setEnabled(True)
        self.stop_btn.setEnabled(False)
        self.progress_bar.setRange(0, 100)
        if success_count > 0 or fail_count > 0 or skip_count > 0:
            self.progress_bar.setValue(100)
        else:
            self.progress_bar.setValue(0)

        summary = f"处理完成。成功：{success_count}，失败：{fail_count}，跳过：{skip_count}"
        if pdf_result_path:
            summary += f"\nPDF：{pdf_result_path}"
        if text_result_path:
            summary += f"\n识别文本：{text_result_path}"

        self.append_log("")
        self.append_log(summary)
        QMessageBox.information(self, "完成", summary)


class PDFWatermarkRemovePage(QWidget):
    back_signal = Signal()

    def __init__(self):
        super().__init__()
        self.worker: Optional[ConvertWorker] = None
        self.selected_input_files: List[Path] = []

        self.init_ui()
        self.apply_styles()
        self.update_output_dir_mode()

    def init_ui(self) -> None:
        back_btn = QPushButton("← 返回首页")
        back_btn.setFixedWidth(120)
        back_btn.clicked.connect(lambda: self.back_signal.emit())
        back_btn.setStyleSheet("""
            QPushButton {
                background-color: #6b7280;
                color: white;
                border: none;
                border-radius: 8px;
                padding: 8px 14px;
                font-weight: 600;
            }
            QPushButton:hover {
                background-color: #4b5563;
            }
        """)

        page_title = QLabel("PDF去水印")
        page_title.setStyleSheet("""
            font-size: 20px;
            font-weight: bold;
            color: #1f2937;
        """)

        header_layout = QHBoxLayout()
        header_layout.addWidget(back_btn)
        header_layout.addStretch()
        header_layout.addWidget(page_title)
        header_layout.addStretch()

        desc_label = QLabel("常用流程：选择 PDF 后会自动补齐输出目录，通常可直接开始。仅移除注释型水印，未检测到会直接报错")
        desc_label.setStyleSheet("color: #6b7280; font-size: 13px; margin-bottom: 10px;")

        self.input_label = QLabel("选择PDF文件")
        self.input_edit = QLineEdit()
        self.input_edit.setReadOnly(True)
        self.input_edit.setPlaceholderText("请选择需要去水印的PDF文件")
        self.input_browse_btn = QPushButton("选择文件")
        self.input_browse_btn.clicked.connect(self.choose_input_files)

        self.output_label = QLabel("输出文件夹")
        self.output_edit = QLineEdit()
        self.output_edit.setPlaceholderText("默认自动使用首个文件所在目录")
        self.output_edit.setReadOnly(True)
        self.output_browse_btn = QPushButton("浏览")
        self.output_browse_btn.clicked.connect(self.choose_output_dir)
        self.auto_output_check = QCheckBox("自动使用推荐输出目录")
        self.auto_output_check.setChecked(True)
        self.auto_output_check.stateChanged.connect(self.update_output_dir_mode)

        self.overwrite_check = QCheckBox("覆盖同名文件")
        self.overwrite_check.setChecked(False)

        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)

        self.start_btn = QPushButton("开始处理")
        self.stop_btn = QPushButton("停止")
        self.stop_btn.setEnabled(False)
        self.start_btn.setDefault(True)

        self.start_btn.clicked.connect(self.start_convert)
        self.stop_btn.clicked.connect(self.stop_convert)

        self.file_list = QListWidget()
        self.file_list.setMinimumHeight(200)

        self.log_edit = QTextEdit()
        self.log_edit.setReadOnly(True)

        form_group = QGroupBox("处理设置")
        form_layout = QGridLayout()
        form_layout.addWidget(self.input_label, 0, 0)
        form_layout.addWidget(self.input_edit, 0, 1)
        form_layout.addWidget(self.input_browse_btn, 0, 2)

        form_layout.addWidget(self.output_label, 1, 0)
        form_layout.addWidget(self.output_edit, 1, 1)
        form_layout.addWidget(self.output_browse_btn, 1, 2)
        form_layout.addWidget(self.auto_output_check, 2, 1, 1, 2)

        option_layout = QHBoxLayout()
        option_layout.addWidget(self.overwrite_check)
        option_layout.addStretch()
        form_layout.addLayout(option_layout, 3, 0, 1, 3)

        form_group.setLayout(form_layout)

        button_layout = QHBoxLayout()
        button_layout.addWidget(self.start_btn)
        button_layout.addWidget(self.stop_btn)
        button_layout.addStretch()

        log_group = QGroupBox("运行日志")
        log_layout = QVBoxLayout()
        log_layout.addWidget(self.log_edit)
        log_group.setLayout(log_layout)

        main_layout = QVBoxLayout()
        main_layout.addLayout(header_layout)
        main_layout.addWidget(desc_label)
        main_layout.addWidget(form_group)
        main_layout.addWidget(self.progress_bar)
        main_layout.addWidget(QLabel("待处理文件："))
        main_layout.addWidget(self.file_list)
        main_layout.addLayout(button_layout)
        main_layout.addWidget(log_group)

        self.setLayout(main_layout)

    def apply_styles(self) -> None:
        self.setStyleSheet("""
            QWidget {
                background-color: #f7f9fc;
                color: #1f2937;
                font-size: 14px;
                font-family: "Microsoft YaHei", "PingFang SC", "Segoe UI";
            }
            QGroupBox {
                border: 1px solid #d7deea;
                border-radius: 10px;
                margin-top: 12px;
                padding-top: 12px;
                background-color: #ffffff;
                font-weight: 600;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 12px;
                padding: 0 6px 0 6px;
            }
            QLineEdit, QTextEdit, QListWidget {
                border: 1px solid #cfd8e3;
                border-radius: 8px;
                padding: 6px 8px;
                background: #ffffff;
            }
            QPushButton {
                background-color: #2563eb;
                color: white;
                border: none;
                border-radius: 8px;
                padding: 8px 14px;
                font-weight: 600;
            }
            QPushButton:hover {
                background-color: #1d4ed8;
            }
            QPushButton:disabled {
                background-color: #94a3b8;
            }
            QProgressBar {
                border: 1px solid #d1d5db;
                border-radius: 8px;
                text-align: center;
                background: #ffffff;
                min-height: 18px;
            }
            QProgressBar::chunk {
                border-radius: 8px;
                background-color: #22c55e;
            }
            QLabel {
                background: transparent;
            }
        """)

    def choose_input_files(self) -> None:
        file_paths, _ = QFileDialog.getOpenFileNames(
            self,
            "选择PDF文件",
            "",
            "PDF文件 (*.pdf)"
        )

        if not file_paths:
            return

        self.selected_input_files = [Path(item).expanduser().resolve() for item in file_paths]
        self.input_edit.setText(f"已选择 {len(self.selected_input_files)} 个文件")
        self.apply_recommended_output_dir()

        self.file_list.clear()
        for file_path in self.selected_input_files:
            self.file_list.addItem(file_path.name)

    def choose_output_dir(self) -> None:
        directory = QFileDialog.getExistingDirectory(self, "选择输出文件夹")
        if directory:
            self.output_edit.setText(directory)

    def update_output_dir_mode(self) -> None:
        use_auto = self.auto_output_check.isChecked()
        self.output_browse_btn.setEnabled(not use_auto)
        if use_auto:
            self.apply_recommended_output_dir()

    def apply_recommended_output_dir(self) -> None:
        recommended_dir = suggest_output_dir_from_files(self.selected_input_files)
        if recommended_dir:
            self.output_edit.setText(str(recommended_dir))
        elif self.auto_output_check.isChecked():
            self.output_edit.clear()

    def append_log(self, message: str) -> None:
        self.log_edit.append(message)

    def update_progress(self, current: int, total: int) -> None:
        if total <= 0:
            self.progress_bar.setValue(0)
            return
        percent = int(current * 100 / total)
        self.progress_bar.setValue(percent)

    def start_convert(self) -> None:
        output_dir = self.output_edit.text().strip()

        if not self.selected_input_files:
            QMessageBox.warning(self, "提示", "请选择需要处理的PDF文件。")
            return

        if not output_dir:
            QMessageBox.warning(self, "提示", "请选择输出文件夹。")
            return

        output_path = Path(output_dir)
        if not output_path.exists() or not output_path.is_dir():
            QMessageBox.warning(self, "提示", "输出文件夹不存在或不是有效目录。")
            return

        self.log_edit.clear()
        self.progress_bar.setRange(0, 0)
        self.append_log("压缩任务已开始，正在处理，请稍候……")

        self.worker = ConvertWorker(
            selected_files=[str(f) for f in self.selected_input_files],
            output_dir=output_dir,
            target_ext="pdf",
            overwrite=self.overwrite_check.isChecked(),
            ffmpeg_bin=None,
            libreoffice_bin=None,
            tesseract_bin=None,
            mode="watermark",
        )
        self.worker.log_signal.connect(self.append_log)
        self.worker.progress_signal.connect(self.update_progress)
        self.worker.finished_signal.connect(self.on_finished)

        self.start_btn.setEnabled(False)
        self.stop_btn.setEnabled(True)
        self.worker.start()

    def stop_convert(self) -> None:
        if self.worker and self.worker.isRunning():
            self.worker.stop()
            self.append_log("正在请求停止当前任务……")

    def on_finished(self, success_count: int, fail_count: int, skip_count: int) -> None:
        self.start_btn.setEnabled(True)
        self.stop_btn.setEnabled(False)
        self.progress_bar.setRange(0, 100)
        if success_count > 0 or fail_count > 0 or skip_count > 0:
            self.progress_bar.setValue(100)
        else:
            self.progress_bar.setValue(0)

        summary = f"处理完成。成功：{success_count}，失败：{fail_count}，跳过：{skip_count}"
        self.append_log("")
        self.append_log(summary)
        QMessageBox.information(self, "完成", summary)


class CompressPage(QWidget):
    back_signal = Signal()

    def __init__(self, ffmpeg_bin: Optional[str]):
        super().__init__()
        self.ffmpeg_bin = ffmpeg_bin
        self.worker: Optional[ConvertWorker] = None
        self.selected_input_files: List[Path] = []

        self.init_ui()
        self.apply_styles()
        self.update_output_dir_mode()

    def init_ui(self) -> None:
        back_btn = QPushButton("← 返回首页")
        back_btn.setFixedWidth(120)
        back_btn.clicked.connect(lambda: self.back_signal.emit())
        back_btn.setStyleSheet("""
            QPushButton {
                background-color: #6b7280;
                color: white;
                border: none;
                border-radius: 8px;
                padding: 8px 14px;
                font-weight: 600;
            }
            QPushButton:hover {
                background-color: #4b5563;
            }
        """)

        page_title = QLabel("图片/视频压缩")
        page_title.setStyleSheet("""
            font-size: 20px;
            font-weight: bold;
            color: #1f2937;
        """)

        header_layout = QHBoxLayout()
        header_layout.addWidget(back_btn)
        header_layout.addStretch()
        header_layout.addWidget(page_title)
        header_layout.addStretch()

        desc_label = QLabel("常用流程：选择文件后会自动补齐输出目录。压缩质量越高，速度越慢")
        desc_label.setStyleSheet("color: #6b7280; font-size: 13px; margin-bottom: 10px;")

        self.input_label = QLabel("选择文件")
        self.input_edit = QLineEdit()
        self.input_edit.setReadOnly(True)
        self.input_edit.setPlaceholderText("请选择图片或视频文件")
        self.input_browse_btn = QPushButton("选择文件")
        self.input_browse_btn.clicked.connect(self.choose_input_files)

        self.output_label = QLabel("输出文件夹")
        self.output_edit = QLineEdit()
        self.output_edit.setPlaceholderText("默认自动使用首个文件所在目录")
        self.output_edit.setReadOnly(True)
        self.output_browse_btn = QPushButton("浏览")
        self.output_browse_btn.clicked.connect(self.choose_output_dir)
        self.auto_output_check = QCheckBox("自动使用推荐输出目录")
        self.auto_output_check.setChecked(True)
        self.auto_output_check.stateChanged.connect(self.update_output_dir_mode)

        self.quality_label = QLabel("压缩质量（越高越慢）")
        self.quality_slider = QSlider(Qt.Horizontal)
        self.quality_slider.setMinimum(50)
        self.quality_slider.setMaximum(100)
        self.quality_slider.setValue(75)
        self.quality_slider.valueChanged.connect(self.update_quality_label)

        self.quality_value_label = QLabel("75")
        self.quality_value_label.setFixedWidth(40)

        self.overwrite_check = QCheckBox("覆盖同名文件")
        self.overwrite_check.setChecked(False)

        ffmpeg_text = self.ffmpeg_bin if self.ffmpeg_bin else "未检测到"
        self.status_label = QLabel(f"ffmpeg：{ffmpeg_text}（视频压缩需要）")
        self.status_label.setWordWrap(True)

        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)

        self.start_btn = QPushButton("开始压缩")
        self.stop_btn = QPushButton("停止")
        self.stop_btn.setEnabled(False)
        self.start_btn.setDefault(True)

        self.start_btn.clicked.connect(self.start_compress)
        self.stop_btn.clicked.connect(self.stop_compress)

        self.file_list = QListWidget()
        self.file_list.setMinimumHeight(200)

        self.log_edit = QTextEdit()
        self.log_edit.setReadOnly(True)

        form_group = QGroupBox("压缩设置")
        form_layout = QGridLayout()
        form_layout.addWidget(self.input_label, 0, 0)
        form_layout.addWidget(self.input_edit, 0, 1)
        form_layout.addWidget(self.input_browse_btn, 0, 2)

        form_layout.addWidget(self.output_label, 1, 0)
        form_layout.addWidget(self.output_edit, 1, 1)
        form_layout.addWidget(self.output_browse_btn, 1, 2)
        form_layout.addWidget(self.auto_output_check, 2, 1, 1, 2)

        quality_layout = QHBoxLayout()
        quality_layout.addWidget(self.quality_slider)
        quality_layout.addWidget(self.quality_value_label)
        form_layout.addWidget(self.quality_label, 3, 0)
        form_layout.addLayout(quality_layout, 3, 1)

        option_layout = QHBoxLayout()
        option_layout.addWidget(self.overwrite_check)
        option_layout.addStretch()
        form_layout.addLayout(option_layout, 4, 0, 1, 3)

        form_group.setLayout(form_layout)

        button_layout = QHBoxLayout()
        button_layout.addWidget(self.start_btn)
        button_layout.addWidget(self.stop_btn)
        button_layout.addStretch()

        log_group = QGroupBox("运行日志")
        log_layout = QVBoxLayout()
        log_layout.addWidget(self.log_edit)
        log_group.setLayout(log_layout)

        main_layout = QVBoxLayout()
        main_layout.addLayout(header_layout)
        main_layout.addWidget(desc_label)
        main_layout.addWidget(form_group)
        main_layout.addWidget(self.status_label)
        main_layout.addWidget(self.progress_bar)
        main_layout.addWidget(QLabel("待处理文件："))
        main_layout.addWidget(self.file_list)
        main_layout.addLayout(button_layout)
        main_layout.addWidget(log_group)

        self.setLayout(main_layout)

    def apply_styles(self) -> None:
        self.setStyleSheet("""
            QWidget {
                background-color: #f7f9fc;
                color: #1f2937;
                font-size: 14px;
                font-family: "Microsoft YaHei", "PingFang SC", "Segoe UI";
            }
            QGroupBox {
                border: 1px solid #d7deea;
                border-radius: 10px;
                margin-top: 12px;
                padding-top: 12px;
                background-color: #ffffff;
                font-weight: 600;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 12px;
                padding: 0 6px 0 6px;
            }
            QLineEdit, QTextEdit, QListWidget {
                border: 1px solid #cfd8e3;
                border-radius: 8px;
                padding: 6px 8px;
                background: #ffffff;
            }
            QPushButton {
                background-color: #2563eb;
                color: white;
                border: none;
                border-radius: 8px;
                padding: 8px 14px;
                font-weight: 600;
            }
            QPushButton:hover {
                background-color: #1d4ed8;
            }
            QPushButton:disabled {
                background-color: #94a3b8;
            }
            QProgressBar {
                border: 1px solid #d1d5db;
                border-radius: 8px;
                text-align: center;
                background: #ffffff;
                min-height: 18px;
            }
            QProgressBar::chunk {
                border-radius: 8px;
                background-color: #22c55e;
            }
            QSlider::groove:horizontal {
                border: 1px solid #cfd8e3;
                height: 8px;
                background: #ffffff;
                border-radius: 4px;
            }
            QSlider::handle:horizontal {
                background: #2563eb;
                border: none;
                width: 18px;
                margin: -5px 0;
                border-radius: 9px;
            }
            QLabel {
                background: transparent;
            }
        """)

    def update_quality_label(self, value: int) -> None:
        self.quality_value_label.setText(str(value))

    def choose_input_files(self) -> None:
        file_paths, _ = QFileDialog.getOpenFileNames(
            self,
            "选择图片或视频文件",
            "",
            "图片和视频 (*.jpg *.jpeg *.png *.bmp *.gif *.webp *.tif *.tiff *.mp4 *.avi *.mov *.mkv *.wmv *.flv *.webm)"
        )

        if not file_paths:
            return

        self.selected_input_files = [Path(item).expanduser().resolve() for item in file_paths]
        self.input_edit.setText(f"已选择 {len(self.selected_input_files)} 个文件")
        self.apply_recommended_output_dir()

        self.file_list.clear()
        for file_path in self.selected_input_files:
            self.file_list.addItem(file_path.name)

    def choose_output_dir(self) -> None:
        directory = QFileDialog.getExistingDirectory(self, "选择输出文件夹")
        if directory:
            self.output_edit.setText(directory)

    def update_output_dir_mode(self) -> None:
        use_auto = self.auto_output_check.isChecked()
        self.output_browse_btn.setEnabled(not use_auto)
        if use_auto:
            self.apply_recommended_output_dir()

    def apply_recommended_output_dir(self) -> None:
        recommended_dir = suggest_output_dir_from_files(self.selected_input_files)
        if recommended_dir:
            self.output_edit.setText(str(recommended_dir))
        elif self.auto_output_check.isChecked():
            self.output_edit.clear()

    def append_log(self, message: str) -> None:
        self.log_edit.append(message)

    def update_progress(self, current: int, total: int) -> None:
        if total <= 0:
            self.progress_bar.setValue(0)
            return
        percent = int(current * 100 / total)
        self.progress_bar.setValue(percent)

    def start_compress(self) -> None:
        output_dir = self.output_edit.text().strip()

        if not self.selected_input_files:
            QMessageBox.warning(self, "提示", "请选择需要压缩的文件。")
            return

        if not output_dir:
            QMessageBox.warning(self, "提示", "请选择输出文件夹。")
            return

        output_path = Path(output_dir)
        if not output_path.exists() or not output_path.is_dir():
            QMessageBox.warning(self, "提示", "输出文件夹不存在或不是有效目录。")
            return

        has_video = any(
            detect_category_by_file(f) == "video"
            for f in self.selected_input_files
        )
        if has_video and not self.ffmpeg_bin:
            QMessageBox.warning(self, "提示", "视频压缩需要 ffmpeg，但未检测到。")
            return

        self.log_edit.clear()
        self.progress_bar.setValue(0)

        self.worker = ConvertWorker(
            selected_files=[str(f) for f in self.selected_input_files],
            output_dir=output_dir,
            target_ext="",
            overwrite=self.overwrite_check.isChecked(),
            ffmpeg_bin=self.ffmpeg_bin,
            libreoffice_bin=None,
            tesseract_bin=None,
            mode="compress",
            compress_quality=self.quality_slider.value(),
        )
        self.worker.log_signal.connect(self.append_log)
        self.worker.progress_signal.connect(self.update_progress)
        self.worker.finished_signal.connect(self.on_finished)

        self.start_btn.setEnabled(False)
        self.stop_btn.setEnabled(True)
        self.worker.start()

    def stop_compress(self) -> None:
        if self.worker and self.worker.isRunning():
            self.worker.stop()
            self.append_log("正在请求停止当前任务……")

    def on_finished(self, success_count: int, fail_count: int, skip_count: int) -> None:
        self.start_btn.setEnabled(True)
        self.stop_btn.setEnabled(False)
        self.progress_bar.setRange(0, 100)
        if success_count > 0 or fail_count > 0 or skip_count > 0:
            self.progress_bar.setValue(100)
        else:
            self.progress_bar.setValue(0)

        summary = f"处理完成。成功：{success_count}，失败：{fail_count}，跳过：{skip_count}"
        self.append_log("")
        self.append_log(summary)
        QMessageBox.information(self, "完成", summary)


class MainWindow(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("文件格式转换工具")

        self.ffmpeg_bin = find_executable(FFMPEG_BINARY_CANDIDATES)
        self.libreoffice_bin = find_executable(LIBREOFFICE_BINARY_CANDIDATES)
        self.tesseract_bin = find_executable(TESSERACT_BINARY_CANDIDATES)
        self.yt_dlp_bin = find_executable(YT_DLP_BINARY_CANDIDATES)

        self.init_ui()
        self.configure_window_geometry()

    def init_ui(self) -> None:
        self.stacked_widget = QStackedWidget()

        self.home_page = HomePage()
        self.home_page.feature_selected.connect(self.show_feature_page)
        self.home_page_container = self.wrap_page(self.home_page)

        self.format_convert_page = FormatConvertPage(
            self.ffmpeg_bin,
            self.libreoffice_bin,
            self.tesseract_bin,
        )
        self.format_convert_page.back_signal.connect(self.show_home)
        self.format_convert_container = self.wrap_page(self.format_convert_page)

        self.scan_to_doc_page = ScanToDocPage(
            self.tesseract_bin,
            self.libreoffice_bin,
        )
        self.scan_to_doc_page.back_signal.connect(self.show_home)
        self.scan_to_doc_container = self.wrap_page(self.scan_to_doc_page)

        self.certificate_pdf_page = CertificatePdfPage(self.tesseract_bin)
        self.certificate_pdf_page.back_signal.connect(self.show_home)
        self.certificate_pdf_container = self.wrap_page(self.certificate_pdf_page)

        self.pdf_watermark_page = PDFWatermarkRemovePage()
        self.pdf_watermark_page.back_signal.connect(self.show_home)
        self.pdf_watermark_container = self.wrap_page(self.pdf_watermark_page)

        self.compress_page = CompressPage(self.ffmpeg_bin)
        self.compress_page.back_signal.connect(self.show_home)
        self.compress_container = self.wrap_page(self.compress_page)

        self.video_download_page = VideoDownloadPage(
            self.yt_dlp_bin,
            self.ffmpeg_bin,
        )
        self.video_download_page.back_signal.connect(self.show_home)
        self.video_download_container = self.wrap_page(self.video_download_page)

        self.stacked_widget.addWidget(self.home_page_container)
        self.stacked_widget.addWidget(self.format_convert_container)
        self.stacked_widget.addWidget(self.scan_to_doc_container)
        self.stacked_widget.addWidget(self.certificate_pdf_container)
        self.stacked_widget.addWidget(self.pdf_watermark_container)
        self.stacked_widget.addWidget(self.compress_container)
        self.stacked_widget.addWidget(self.video_download_container)

        main_layout = QVBoxLayout()
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.addWidget(self.stacked_widget)

        self.setLayout(main_layout)

    def wrap_page(self, page: QWidget) -> QScrollArea:
        container = QScrollArea()
        container.setWidgetResizable(True)
        container.setFrameShape(QFrame.NoFrame)
        container.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        container.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        container.setWidget(page)
        return container

    def configure_window_geometry(self) -> None:
        screen = self.screen() or QApplication.primaryScreen()
        if screen is None:
            self.resize(1200, 800)
            self.setMinimumSize(860, 620)
            return

        available = screen.availableGeometry()
        min_width = min(860, max(520, available.width() - 80), available.width())
        min_height = min(620, max(420, available.height() - 80), available.height())
        target_width = min(1280, max(min_width, available.width() - 120))
        target_height = min(860, max(min_height, available.height() - 120))

        self.setMinimumSize(min_width, min_height)
        self.resize(target_width, target_height)

    def show_feature_page(self, feature_key: str) -> None:
        if feature_key == "video_download":
            self.stacked_widget.setCurrentWidget(self.video_download_container)
        elif feature_key == "scan_to_doc":
            self.stacked_widget.setCurrentWidget(self.scan_to_doc_container)
        elif feature_key == "certificate_pdf":
            self.stacked_widget.setCurrentWidget(self.certificate_pdf_container)
        elif feature_key == "pdf_watermark":
            self.stacked_widget.setCurrentWidget(self.pdf_watermark_container)
        elif feature_key == "compress":
            self.stacked_widget.setCurrentWidget(self.compress_container)
        else:
            self.stacked_widget.setCurrentWidget(self.format_convert_container)

    def show_home(self) -> None:
        self.stacked_widget.setCurrentWidget(self.home_page_container)


def main() -> None:
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
